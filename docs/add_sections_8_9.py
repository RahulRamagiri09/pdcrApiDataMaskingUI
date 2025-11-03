from docx import Document
from docx.shared import Pt

# Open existing document
doc = Document(r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx")

def add_heading(text, level=1):
    doc.add_heading(text, level)

def add_text(text):
    return doc.add_paragraph(text)

def add_diagram(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

print("Adding Section 8: Security Architecture...")

doc.add_page_break()

# ==============================================================================
# SECTION 8: SECURITY ARCHITECTURE
# ==============================================================================

add_heading('8. SECURITY ARCHITECTURE', 1)

add_heading('8.1 Security Layers Overview', 2)

security_diagram = """
SECURITY ARCHITECTURE (Defense in Depth)

Layer 7: Physical Security
├─ Data center access controls
├─ Server room security
└─ Hardware security modules (HSM)
        ↓
Layer 6: Network Security
├─ Firewall rules (restrict database ports)
├─ VPN requirement for remote access
├─ Network segmentation (DMZ for web tier)
└─ DDoS protection
        ↓
Layer 5: Transport Security
├─ HTTPS/TLS 1.3 for all web traffic
├─ SSL/TLS for database connections
├─ Certificate validation
└─ Strong cipher suites only
        ↓
Layer 4: Application Authentication
├─ JWT token-based authentication
├─ Token expiration (24 hours)
├─ Automatic logout on expiration
├─ Password hashing (bcrypt, cost factor 12)
└─ No plaintext passwords stored or transmitted
        ↓
Layer 3: Authorization (RBAC)
├─ Role-based access control (Admin/User)
├─ Resource ownership validation
├─ Permission checks on every API endpoint
└─ Principle of least privilege
        ↓
Layer 2: Data Protection
├─ Database credentials encrypted at rest (AES-256)
├─ PII data masked during transformation
├─ No sensitive data in logs
└─ Encrypted backups
        ↓
Layer 1: Application Security
├─ SQL injection prevention (parameterized queries)
├─ XSS protection (input sanitization)
├─ CSRF protection (SameSite cookies)
├─ Input validation on all endpoints
├─ Output encoding
└─ Security headers (CSP, X-Frame-Options)
        ↓
Layer 0: Audit & Monitoring
├─ Complete audit trail (who, what, when)
├─ Failed login attempt logging
├─ Security event monitoring
├─ Alerting on suspicious activity
└─ Compliance reporting
"""
add_diagram(security_diagram)

add_heading('8.2 Authentication Architecture', 2)

jwt_diagram = """
JWT TOKEN AUTHENTICATION FLOW

User Credentials
      │
      ├─ Username: Plain text
      └─ Password: Plain text
      │
      ▼
HTTPS Transport (Encrypted in flight)
      │
      ▼
Backend Receives Login Request
      │
      ▼
Query users table WHERE username = ?
      │
      ▼
Compare Password Hash
bcrypt.compare(input_password, stored_hash)
      │
      ├─ MATCH ──────────────────────┐
      │                               │
      ▼                               ▼
Generate JWT Token            Return 401 Unauthorized
      │
      ├─ Header: {"alg": "HS256", "typ": "JWT"}
      ├─ Payload: {user_id, username, role, exp}
      └─ Signature: HMACSHA256(header + payload, SECRET_KEY)
      │
      ▼
Return Token to Frontend
      │
      ▼
Frontend Stores in localStorage
      │
      ▼
All Subsequent Requests Include:
Authorization: Bearer <JWT_TOKEN>
      │
      ▼
Backend JWT Middleware:
1. Extract token
2. Verify signature
3. Check expiration
4. Decode payload
5. Extract user context
      │
      ├─ VALID ──> Process Request
      └─ INVALID ─> Return 401
"""
add_diagram(jwt_diagram)

add_heading('8.3 Authentication Security Features', 2)

add_text('Password Security:')
add_text('- Passwords hashed using bcrypt with salt (cost factor: 12 rounds)')
add_text('- Never stored in plaintext')
add_text('- Never transmitted except during initial login (over HTTPS)')
add_text('- Password requirements: Minimum 8 characters (configurable)')
add_text('')

add_text('Token Security:')
add_text('- JWT signed with HMAC-SHA256 algorithm')
add_text('- Secret key stored in environment variables (never in code)')
add_text('- Different secret keys for dev/staging/production')
add_text('- Token expiration enforced (24-hour default)')
add_text('- Automatic logout on token expiration')
add_text('- Token includes user role for authorization')
add_text('')

add_text('Session Management:')
add_text('- Stateless authentication (no server-side sessions)')
add_text('- Token revocation via expiration (no refresh tokens in current version)')
add_text('- Frontend clears token on logout')
add_text('- Expired tokens automatically rejected by backend')

add_heading('8.4 Authorization (RBAC) Architecture', 2)

rbac_diagram = """
ROLE-BASED ACCESS CONTROL (RBAC) MODEL

                ┌──────────────┐
                │    Users     │
                └──────┬───────┘
                       │
            ┌──────────┴──────────┐
            │                     │
            ▼                     ▼
     ┌────────────┐        ┌────────────┐
     │ Admin Role │        │ User Role  │
     └─────┬──────┘        └─────┬──────┘
           │                     │
           ▼                     ▼
    All Permissions       Own Resources Only

PERMISSION MATRIX:

Resource / Action         | Admin | User | Guest
─────────────────────────┼───────┼──────┼───────
Login / Logout            │   ✓   │  ✓   │   ✓
View own profile          │   ✓   │  ✓   │   ✗
                          │       │      │
CONNECTIONS:              │       │      │
View all connections      │   ✓   │  ✓   │   ✗
Create connection         │   ✓   │  ✓   │   ✗
Edit own connection       │   ✓   │  ✓   │   ✗
Edit others' connection   │   ✓   │  ✗   │   ✗
Delete own connection     │   ✓   │  ✓   │   ✗
Delete others' connection │   ✓   │  ✗   │   ✗
Test connection           │   ✓   │  ✓   │   ✗
                          │       │      │
WORKFLOWS:                │       │      │
View all workflows        │   ✓   │  ✓   │   ✗
Create workflow           │   ✓   │  ✓   │   ✗
Edit own workflow         │   ✓   │  ✓   │   ✗
Edit others' workflow     │   ✓   │  ✗   │   ✗
Delete own workflow       │   ✓   │  ✓   │   ✗
Delete others' workflow   │   ✓   │  ✗   │   ✗
Execute own workflow      │   ✓   │  ✓   │   ✗
Execute others' workflow  │   ✓   │  ✗   │   ✗
                          │       │      │
EXECUTIONS:               │       │      │
View all execution history│   ✓   │  ✗   │   ✗
View own execution history│   ✓   │  ✓   │   ✗
View execution details    │   ✓   │  ✓*  │   ✗
                          │       │ *own │
SYSTEM:                   │       │      │
View system logs          │   ✓   │  ✗   │   ✗
Manage users              │   ✓   │  ✗   │   ✗
Change system settings    │   ✓   │  ✗   │   ✗
View audit trail          │   ✓   │  ✗   │   ✗
"""
add_diagram(rbac_diagram)

add_text('Benefits of RBAC:')
add_text('- Simplified permission management')
add_text('- Clear separation of admin and user capabilities')
add_text('- Principle of least privilege')
add_text('- Easy to add new roles in future')
add_text('- Audit trail shows who had what permissions')

add_heading('8.5 Data Encryption Architecture', 2)

encryption_diagram = """
ENCRYPTION AT REST

Database Connection Passwords (connections.password)
                │
                ▼
    Plaintext Password: "MySecurePass123!"
                │
                ▼
    AES-256-CBC Encryption
    Key: Environment variable SECRET_ENCRYPTION_KEY
    IV: Random 16 bytes
                │
                ▼
    Encrypted Ciphertext stored in database
    "U2FsdGVkX1+..."
                │
    WHEN NEEDED FOR CONNECTION:
                ▼
    AES-256-CBC Decryption
    Key: From environment, IV: From cipher
                │
                ▼
    Plaintext in Memory (used immediately, cleared after use)


PASSWORD HASHING (users.password_hash)

User Registration/Password Change
                │
                ▼
    User enters password: "UserPassword123"
                │
                ▼
    bcrypt Hash Generation
    Salt: Random (auto), Cost Factor: 12 rounds (2^12 iterations)
                │
                ▼
    Hash includes: Algorithm version, Cost factor, Salt, Hash output
    "$2b$12$abcd...xyz"
                │
                ▼
    Stored in Database users.password_hash

User Login
                │
                ▼
    User enters password: "UserPassword123"
                │
                ▼
    Retrieve stored hash from database
                │
                ▼
    bcrypt.compare()
    Hashes input password with same salt from stored hash
    Compares results
                │
         ┌──────┴──────┐
         │ MATCH       │ NO MATCH
         ▼             ▼
    Login Success   Login Failed


ENCRYPTION IN TRANSIT

Client                          Server
  │                               │
  │  HTTPS/TLS 1.3 Handshake     │
  │  1. Client Hello             │
  ├─────────────────────────────>│
  │  2. Server Hello             │
  │  3. Certificate (public key) │
  │<──────────────────────────────│
  │  4. Verify certificate       │
  │  5. Generate session keys    │
  │  6. Encrypted with public key│
  ├─────────────────────────────>│
  │  7. Derive symmetric keys    │
  │                               │
  │  ALL SUBSEQUENT DATA ENCRYPTED│
  │<─────────────────────────────>│
"""
add_diagram(encryption_diagram)

add_heading('8.6 Application Security Measures', 2)

add_text('SQL Injection Prevention: All database queries use parameterized statements. User input never concatenated into SQL strings. ORM/query builder enforces parameterization. Database permissions follow least privilege principle.')
add_text('')

add_text('Cross-Site Scripting (XSS) Prevention: React automatically escapes output (prevents XSS by default). User input sanitized before storage. Content Security Policy (CSP) headers. No use of dangerous functions (dangerouslySetInnerHTML).')
add_text('')

add_text('Cross-Site Request Forgery (CSRF) Prevention: SameSite cookie attribute set to Strict. JWT tokens in Authorization header (not cookies). State-changing operations require valid JWT. No GET requests for state changes.')
add_text('')

add_text('Input Validation: Frontend validation for user experience. Backend validation for security (never trust client). Type checking with Pydantic models. Length limits on all text fields. Whitelist validation for enums.')
add_text('')

add_text('Security Headers: X-Content-Type-Options: nosniff, X-Frame-Options: DENY, X-XSS-Protection: 1; mode=block, Strict-Transport-Security: max-age=31536000, Content-Security-Policy: default-src "self".')

add_heading('8.7 Audit Trail', 2)

audit_diagram = """
AUDIT TRAIL SYSTEM

Every User Action Logged:

┌──────────────────────────────────────────────────────┐
│ AUDIT LOG ENTRY                                      │
├──────────────────────────────────────────────────────┤
│ • Timestamp: 2025-01-15 14:30:45.123                │
│ • User ID: 123                                       │
│ • Username: john.doe                                 │
│ • Role: user                                         │
│ • Action: EXECUTE_WORKFLOW                           │
│ • Resource: workflow_id=456                          │
│ • IP Address: 192.168.1.100                          │
│ • User Agent: Chrome/120.0 Windows                   │
│ • Status: SUCCESS                                    │
│ • Details: Processed 1000 rows, masked 1000 rows    │
│ • Duration: 150 seconds                              │
└──────────────────────────────────────────────────────┘

Logged Actions:
✓ Login attempts (success and failure)
✓ Logout events
✓ Connection creation/modification/deletion
✓ Connection test attempts
✓ Workflow creation/modification/deletion
✓ Workflow executions (with full metrics)
✓ Schema/table browsing
✓ Configuration changes
✓ Permission denials (403 errors)
✓ Authentication failures (401 errors)
✓ API errors (500 errors)

Compliance Benefits:
• GDPR Article 30: Records of processing activities
• HIPAA Audit Controls: Complete audit trail
• SOC 2: User access and change logs
• ISO 27001: Security event logging
• PCI DSS: Track and monitor all access to data

Retention Policy:
• Execution logs: 90 days (configurable)
• Security events: 1 year
• Compliance logs: 7 years (if required)
• Regular archival to cold storage
"""
add_diagram(audit_diagram)

doc.add_page_break()

# ==============================================================================
# SECTION 9: DEPLOYMENT ARCHITECTURE
# ==============================================================================

print("Adding Section 9: Deployment Architecture...")

add_heading('9. DEPLOYMENT ARCHITECTURE', 1)

add_heading('9.1 Development Environment', 2)

dev_diagram = """
DEVELOPMENT ENVIRONMENT

Developer Workstation (Windows/Mac/Linux)
┌────────────────────────────────────────────────────┐
│                                                    │
│  ┌─────────────────────────────────────────────┐  │
│  │        Frontend (Port 3000)                 │  │
│  │  ──────────────────────────────             │  │
│  │  React Development Server                   │  │
│  │  • npm start                                │  │
│  │  • Hot Module Replacement (HMR)             │  │
│  │  • Source maps for debugging                │  │
│  │  • React Developer Tools                    │  │
│  │  • Auto-reload on file changes              │  │
│  │  • URL: http://localhost:3000               │  │
│  └─────────────────────────────────────────────┘  │
│                      │                             │
│                      │ API Calls                   │
│                      ▼                             │
│  ┌─────────────────────────────────────────────┐  │
│  │        Backend (Port 8000)                  │  │
│  │  ──────────────────────────────             │  │
│  │  FastAPI with Uvicorn                       │  │
│  │  • uvicorn main:app --reload                │  │
│  │  • Auto-reload on code changes              │  │
│  │  • OpenAPI docs: /docs                      │  │
│  │  • Debug logging enabled                    │  │
│  │  • Single worker (no concurrency needed)    │  │
│  │  • URL: http://localhost:8000               │  │
│  └─────────────────────────────────────────────┘  │
│                      │                             │
│                      │ SQL Queries                 │
│                      ▼                             │
│  ┌─────────────────────────────────────────────┐  │
│  │      SQL Server Database                    │  │
│  │  ──────────────────────────────             │  │
│  │  Local Instance or Remote Dev Server        │  │
│  │  • Application DB: pii_tool_dev             │  │
│  │  • Test source/target databases             │  │
│  │  • Sample data for testing                  │  │
│  │  • Regular backups                          │  │
│  │  • Connection: localhost:1433 or dev-server │  │
│  └─────────────────────────────────────────────┘  │
│                                                    │
│  Configuration (.env file):                        │
│  • SECRET_KEY=dev-secret-key-12345                 │
│  • DATABASE_URL=localhost:1433                     │
│  • ENVIRONMENT=development                         │
│  • DEBUG=True                                      │
│  • LOG_LEVEL=DEBUG                                 │
│  • CORS_ORIGINS=http://localhost:3000              │
│                                                    │
└────────────────────────────────────────────────────┘

Development Tools:
• Git for version control
• VS Code / PyCharm for IDE
• Postman for API testing
• SQL Server Management Studio (SSMS) for database
• Chrome DevTools for frontend debugging
• Python debugger (pdb) for backend debugging
"""
add_diagram(dev_diagram)

add_heading('9.2 Production Environment', 2)

prod_diagram = """
PRODUCTION ENVIRONMENT (High Availability)

                    Internet
                       │
                       │ HTTPS
                       ▼
┌──────────────────────────────────────────────────┐
│         LAYER 1: LOAD BALANCER                   │
│      (Nginx / HAProxy / Cloud LB)                │
├──────────────────────────────────────────────────┤
│ • SSL/TLS termination (HTTPS certificates)       │
│ • Health checks (ping endpoints every 30s)       │
│ • Request routing: /api/* → Backend              │
│ •                  /* → Frontend                 │
│ • Rate limiting (1000 req/min per IP)            │
│ • DDoS protection                                │
│ • Gzip compression                               │
│ • Security headers injection                     │
│ • Session affinity (sticky sessions)             │
└────────┬─────────────────────────────────────────┘
         │
   ┌─────┴─────┐
   │           │
   ▼           ▼
┌────────┐  ┌─────────────────────────────────────┐
│Frontend│  │   LAYER 2: BACKEND API CLUSTER      │
│Server  │  │                                     │
│        │  │  ┌───────────┐  ┌───────────┐      │
│Nginx/  │  │  │API Server1│  │API Server2│      │
│Apache  │  │  │Gunicorn   │  │Gunicorn   │      │
│        │  │  │+Uvicorn   │  │+Uvicorn   │      │
│Serves: │  │  │4 workers  │  │4 workers  │      │
│• index │  │  └───────────┘  └───────────┘      │
│  .html │  │          │              │           │
│• bundle│  │  ┌───────────┐  ┌───────────┐      │
│  .js   │  │  │API Server3│  │API Server4│      │
│• static│  │  │Gunicorn   │  │Gunicorn   │      │
│  assets│  │  │+Uvicorn   │  │+Uvicorn   │      │
│        │  │  │4 workers  │  │4 workers  │      │
│Caching:│  │  └───────────┘  └───────────┘      │
│• JS/CSS│  └──────────┬────────────────────────┘
│  1 year│             │
│• Images│             │ SQL Queries
│  30d   │             │ (Connection Pool)
└────────┘             ▼
         ┌──────────────────────────────────────────┐
         │   LAYER 3: DATABASE TIER (SQL Server)    │
         ├──────────────────────────────────────────┤
         │  ┌────────────────────────────────────┐  │
         │  │   PRIMARY DATABASE SERVER          │  │
         │  │ • Application DB (pii_tool_prod)   │  │
         │  │ • Always-On Availability Groups    │  │
         │  │ • Automatic failover to secondary  │  │
         │  │ • Full backups: Daily at 2 AM      │  │
         │  │ • Differential backups: Every 6hrs │  │
         │  │ • Transaction log: Every 15 mins   │  │
         │  │ • Point-in-time recovery enabled   │  │
         │  └──────┬─────────────────────────────┘  │
         │         │ Synchronous Replication        │
         │         ▼                                 │
         │  ┌────────────────────────────────────┐  │
         │  │   SECONDARY DATABASE (Standby)     │  │
         │  │ • Synchronized replica             │  │
         │  │ • Automatic failover on failure    │  │
         │  │ • Can serve read-only queries      │  │
         │  └────────────────────────────────────┘  │
         │                                           │
         │  ┌────────────────────────────────────┐  │
         │  │   READ REPLICA (For Reporting)     │  │
         │  │ • Asynchronous replication         │  │
         │  │ • Read-only access                 │  │
         │  │ • Used for analytics               │  │
         │  │ • Offloads read queries            │  │
         │  └────────────────────────────────────┘  │
         └──────────────────────────────────────────┘

MONITORING & LOGGING INFRASTRUCTURE
┌──────────────────────────────────────────────────┐
│ Application Logs:                                │
│ ├─ ELK Stack (Elasticsearch, Logstash, Kibana)  │
│ ├─ Centralized log aggregation                  │
│ ├─ Structured JSON logs                          │
│ └─ Log retention: 90 days                        │
│                                                  │
│ Metrics & Monitoring:                            │
│ ├─ Prometheus (metrics collection)               │
│ ├─ Grafana (visualization dashboards)            │
│ ├─ Metrics tracked:                              │
│ │   • API response times                         │
│ │   • Error rates (4xx, 5xx)                     │
│ │   • Database query performance                 │
│ │   • CPU/Memory/Disk usage                      │
│ │   • Active user sessions                       │
│ │   • Workflow execution metrics                 │
│ └─ Data retention: 1 year                         │
│                                                  │
│ Alerting:                                        │
│ ├─ PagerDuty / Opsgenie                          │
│ ├─ Alert conditions:                             │
│ │   • Error rate > 5% (5 minutes)               │
│ │   • Response time > 5 seconds (p95)           │
│ │   • CPU usage > 80% (10 minutes)              │
│ │   • Disk space < 20%                          │
│ │   • Database connection failures              │
│ │   • SSL certificate expiring (30 days)        │
│ └─ Escalation: Email → SMS → Phone call          │
│                                                  │
│ Uptime Monitoring:                               │
│ ├─ External monitoring (Pingdom/UptimeRobot)     │
│ ├─ Health check endpoints every 1 minute         │
│ └─ SLA target: 99.9% uptime                       │
└──────────────────────────────────────────────────┘
"""
add_diagram(prod_diagram)

add_heading('9.3 Deployment Process', 2)

cicd_diagram = """
CI/CD PIPELINE (Continuous Deployment)

Developer                                  Production
   │                                           │
   │  1. Code Changes (Feature/Bug Fix)       │
   ├──────────┐                                │
   │          │                                │
   │  2. Commit to Git (feature branch)       │
   ▼          │                                │
Git Repo      │                                │
(GitHub/      │                                │
 GitLab)      │                                │
   │          │                                │
   │ 3. Create Pull Request                    │
   ▼          │                                │
Code Review   │                                │
• Peer review│                                 │
• Arch review│                                 │
• Security   │                                 │
   │          │                                │
   │ 4. Merge to main branch                   │
   ▼          │                                │
CI/CD Pipeline│                                │
(GitHub Actions/Jenkins/GitLab CI)             │
   │          │                                │
   │ 5. Automated Tests                        │
   ▼          │                                │
Test Suite    │                                │
• Unit tests  │                                │
• Integration │                                │
• E2E tests   │                                │
• Security    │                                │
• Linting     │                                │
   │          │                                │
   ├─ FAIL ──> Notify Developer                │
   │           Stop Pipeline                    │
   │          │                                │
   ▼ PASS     │                                │
Build Stage   │                                │
Frontend:     │                                │
• npm build   │                                │
• Optimize    │                                │
Backend:      │                                │
• Install deps│                                │
• Create      │                                │
  artifacts   │                                │
   │          │                                │
   │ 6. Deploy to Staging                      │
   ▼          │                                │
Staging Env   │                                │
• Identical   │                                │
  to prod     │                                │
• Run smoke   │                                │
  tests       │                                │
• QA verify   │                                │
• Load testing│                                │
   │          │                                │
   │ 7. Manual approval (optional)             │
   ▼          │                                │
Deployment    │                                │
Gate          │                                │
• QA sign-off │                                │
• Product     │                                │
  sign-off    │                                │
   │          │                                │
   │ 8. Deploy to Production                   │
   ▼          │                                │
Production    │                                │
Deploy        │                                │
Blue-Green:   │                                │
1. Deploy to  │                                │
   "green"    │                                │
2. Health     │                                │
   checks     │                                │
3. Switch     │                                │
   traffic    │                                │
4. Keep "blue"│                                │
   for rollback                                │
   │          │                                │
   │ 9. Post-deployment                        │
   ▼          │                                │
Verification  │                                │
• Smoke tests │                                │
• Monitor     │                                │
  errors      │                                │
• Check       │                                │
  metrics     │                                │
• Validate    │                                │
  uptime      │                                │
   │          │                                │
   ├─ SUCCESS ──> Update docs                  │
   │              Notify team                  │
   │              Close tickets                │
   │          │                                │
   └─ FAILURE ──> Automatic rollback           │
                  to previous version          │
                  Alert on-call engineer       │
                                              │
                                              ▼
                                  Production Running
                                  New Version
"""
add_diagram(cicd_diagram)

add_text('Deployment Strategy:')
add_text('- Blue-Green Deployment: Zero-downtime deployments')
add_text('- Rollback Plan: Previous version kept running for quick rollback')
add_text('- Database Migrations: Run before application deployment with backward compatibility')
add_text('- Feature Flags: Enable/disable features without redeployment')

add_heading('9.4 Environment Configuration Matrix', 2)

config_table = """
ENVIRONMENT CONFIGURATION MATRIX

Configuration       | Development      | Staging               | Production
───────────────────┼─────────────────┼──────────────────────┼────────────────
Frontend URL        | localhost:3000   | staging.company.com   | pii-tool.company.com
Backend URL         | localhost:8000   | api-staging.company..| api.company.com
Database            | Local/Dev Server | Staging DB            | Production DB Cluster
HTTPS               | No (HTTP)        | Yes (Let's Encrypt)   | Yes (Commercial SSL)
JWT Expiration      | 7 days           | 24 hours              | 24 hours
Logging Level       | DEBUG            | INFO                  | WARNING
Error Details       | Full stack traces| Limited details       | Generic messages only
CORS                | localhost:3000   | staging domain only   | production domain only
Rate Limiting       | Disabled         | Enabled (lenient)     | Enabled (strict)
Backups             | Manual           | Daily                 | Hourly + daily + weekly
Monitoring          | Local only       | Basic monitoring      | Full monitoring + alerts
Workers             | 1                | 2                     | 4-8 (auto-scale)
"""
add_diagram(config_table)

# Save progress
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print(f"Sections 8-9 added successfully!")
print(f"Document saved: {output_path}")
