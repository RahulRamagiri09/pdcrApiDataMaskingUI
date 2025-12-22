from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open existing document
doc = Document(r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx")

def add_heading(text, level=1):
    doc.add_heading(text, level)

def add_text(text):
    return doc.add_paragraph(text)

def add_diagram(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

print("Adding Section 10: Integration Architecture...")

doc.add_page_break()

# ==============================================================================
# SECTION 10: INTEGRATION ARCHITECTURE
# ==============================================================================

add_heading('10. INTEGRATION ARCHITECTURE', 1)

add_heading('10.1 API Integration Pattern', 2)

integration_diagram = """
API INTEGRATION ARCHITECTURE

            ┌────────────────┐
            │  External      │
            │  Systems       │
            │  (Future)      │
            └────────┬───────┘
                     │
                     │ REST API
                     ▼
┌────────────────────────────────────────────────────┐
│         API GATEWAY (Future)                       │
│  ├─ Rate limiting                                  │
│  ├─ API key management                             │
│  ├─ Request transformation                         │
│  └─ Response caching                               │
└────────────┬───────────────────────────────────────┘
             │
             ▼
┌────────────────────────────────────────────────────┐
│     PII MASKING TOOL BACKEND API                   │
│        (FastAPI Application)                       │
├────────────────────────────────────────────────────┤
│                                                    │
│  API Endpoints (RESTful):                          │
│                                                    │
│  Authentication:                                   │
│  POST   /api/auth/login                            │
│  GET    /api/auth/me                               │
│                                                    │
│  Connections:                                      │
│  GET    /api/connections                           │
│  POST   /api/connections                           │
│  GET    /api/connections/{id}                      │
│  PUT    /api/connections/{id}                      │
│  DELETE /api/connections/{id}                      │
│  POST   /api/connections/{id}/test                 │
│                                                    │
│  Schema Discovery:                                 │
│  GET    /api/connections/{id}/schemas              │
│  GET    /api/connections/{id}/schemas/{schema}/    │
│         tables                                     │
│  GET    /api/connections/{id}/schemas/{schema}/    │
│         tables/{table}/columns                     │
│                                                    │
│  Workflows:                                        │
│  GET    /api/workflows                             │
│  POST   /api/workflows                             │
│  GET    /api/workflows/{id}                        │
│  PUT    /api/workflows/{id}                        │
│  DELETE /api/workflows/{id}                        │
│  POST   /api/workflows/{id}/execute                │
│                                                    │
│  Executions:                                       │
│  GET    /api/executions                            │
│  GET    /api/executions/{id}                       │
│                                                    │
│  PII Attributes:                                   │
│  GET    /api/pii-attributes                        │
│                                                    │
│  Health Check:                                     │
│  GET    /health                                    │
│  GET    /api/health/db                             │
│                                                    │
└────────────────────────────────────────────────────┘
"""
add_diagram(integration_diagram)

add_heading('10.2 API Request/Response Format', 2)

add_text('Standard Request Format:')
req_format = """
Method: POST
URL: https://api.company.com/api/workflows
Headers:
  Authorization: Bearer <JWT_TOKEN>
  Content-Type: application/json
  Accept: application/json

Body (JSON):
{
  "name": "Mask Employee Data",
  "description": "Mask PII in employee table",
  "source_connection_id": 1,
  "target_connection_id": 2,
  "status": "active",
  "mappings": [
    {
      "source_schema": "dbo",
      "source_table": "employees",
      "target_schema": "dbo",
      "target_table": "masked_employees",
      "column_mappings": [
        {
          "source_column": "first_name",
          "pii_attribute": "first_name",
          "target_column": "first_name"
        }
      ]
    }
  ]
}
"""
add_diagram(req_format)

add_text('Standard Response Format (Success):')
success_format = """
Status: 201 Created
Headers:
  Content-Type: application/json

Body (JSON):
{
  "success": true,
  "data": {
    "id": 123,
    "name": "Mask Employee Data",
    "description": "Mask PII in employee table",
    "status": "active",
    "created_at": "2025-01-15T14:30:00Z"
  },
  "message": "Workflow created successfully"
}
"""
add_diagram(success_format)

add_text('Standard Response Format (Error):')
error_format = """
Status: 400 Bad Request
Headers:
  Content-Type: application/json

Body (JSON):
{
  "success": false,
  "error": {
    "code": "VALIDATION_ERROR",
    "message": "Invalid workflow configuration",
    "details": [
      {
        "field": "source_connection_id",
        "issue": "Connection with ID 1 does not exist"
      }
    ]
  }
}
"""
add_diagram(error_format)

add_heading('10.3 Integration with External Systems (Future)', 2)

add_text('Potential Integrations:')
add_text('')
add_text('CI/CD Pipelines:')
add_text('- Trigger workflow execution as part of deployment')
add_text('- Mask test data before running integration tests')
add_text('- API endpoint: POST /api/workflows/{id}/execute')
add_text('')
add_text('Data Warehouse ETL:')
add_text('- Mask PII before loading into data warehouse')
add_text('- Schedule recurring executions via API')
add_text('- Monitor execution status programmatically')
add_text('')
add_text('Compliance Tools:')
add_text('- Export audit logs for compliance reporting')
add_text('- Provide API for security information and event management (SIEM)')
add_text('- Data lineage tracking')
add_text('')
add_text('Monitoring Systems:')
add_text('- Health check endpoints for uptime monitoring')
add_text('- Metrics API for dashboard integration')
add_text('- Alert webhook for execution failures')

add_heading('10.4 API Endpoint Reference', 2)

api_ref = """
COMPLETE API ENDPOINT REFERENCE

┌─────────────────────────────────────────────────────────┐
│ AUTHENTICATION ENDPOINTS                                │
├─────────────────────────────────────────────────────────┤
│ POST   /api/auth/login                                  │
│        Description: Authenticate user, receive JWT      │
│        Auth Required: No                                │
│        Request: {username, password}                    │
│        Response: {access_token, user}                   │
│                                                         │
│ GET    /api/auth/me                                     │
│        Description: Get current user profile            │
│        Auth Required: Yes                               │
│        Response: {id, username, email, role}            │
└─────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────┐
│ CONNECTION ENDPOINTS                                    │
├─────────────────────────────────────────────────────────┤
│ GET    /api/connections                                 │
│        Description: List all connections                │
│        Auth Required: Yes                               │
│        Query Params: type (optional: source/target)     │
│                                                         │
│ POST   /api/connections                                 │
│        Description: Create new connection               │
│        Auth Required: Yes                               │
│        Request: {name, type, server, database, ...}     │
│                                                         │
│ POST   /api/connections/{id}/test                       │
│        Description: Test connection validity            │
│        Auth Required: Yes                               │
│        Response: {success, message}                     │
│                                                         │
│ GET    /api/connections/{id}/schemas                    │
│        Description: Get list of schemas                 │
│        Auth Required: Yes                               │
│        Response: ["dbo", "schema1", ...]                │
│                                                         │
│ GET    /api/connections/{id}/schemas/{schema}/tables    │
│        Description: Get list of tables in schema        │
│        Auth Required: Yes                               │
│        Response: ["table1", "table2", ...]              │
│                                                         │
│ GET    /api/connections/{id}/schemas/{schema}/          │
│        source_tables/{table}/columns                    │
│        Description: Get column metadata                 │
│        Auth Required: Yes                               │
│        Response: [{name, data_type, is_nullable}, ...]  │
└─────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────┐
│ WORKFLOW ENDPOINTS                                      │
├─────────────────────────────────────────────────────────┤
│ GET    /api/workflows                                   │
│        Description: List all workflows                  │
│        Auth Required: Yes                               │
│                                                         │
│ POST   /api/workflows                                   │
│        Description: Create new workflow                 │
│        Auth Required: Yes                               │
│        Request: {name, description, connections, ...}   │
│                                                         │
│ GET    /api/workflows/{id}                              │
│        Description: Get workflow details                │
│        Auth Required: Yes                               │
│                                                         │
│ PUT    /api/workflows/{id}                              │
│        Description: Update workflow                     │
│        Auth Required: Yes                               │
│                                                         │
│ DELETE /api/workflows/{id}                              │
│        Description: Delete workflow                     │
│        Auth Required: Yes                               │
│                                                         │
│ POST   /api/workflows/{id}/execute                      │
│        Description: Execute workflow                    │
│        Auth Required: Yes                               │
│        Response: {execution_id, status}                 │
└─────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────┐
│ EXECUTION ENDPOINTS                                     │
├─────────────────────────────────────────────────────────┤
│ GET    /api/executions                                  │
│        Description: List execution history              │
│        Auth Required: Yes                               │
│        Query Params: workflow_id, status                │
│                                                         │
│ GET    /api/executions/{id}                             │
│        Description: Get execution details               │
│        Auth Required: Yes                               │
│        Response: {id, status, rows_processed, ...}      │
└─────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────┐
│ PII ATTRIBUTES ENDPOINT                                 │
├─────────────────────────────────────────────────────────┤
│ GET    /api/pii-attributes                              │
│        Description: Get categorized PII attributes      │
│        Auth Required: Yes                               │
│        Response: {string: [...], numeric: [...], ...}   │
└─────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────┐
│ HEALTH CHECK ENDPOINTS                                  │
├─────────────────────────────────────────────────────────┤
│ GET    /health                                          │
│        Description: Application health check            │
│        Auth Required: No                                │
│        Response: {status: "healthy"}                    │
│                                                         │
│ GET    /api/health/db                                   │
│        Description: Database connectivity check         │
│        Auth Required: Yes                               │
│        Response: {status: "connected"}                  │
└─────────────────────────────────────────────────────────┘
"""
add_diagram(api_ref)

# ==============================================================================
# DOCUMENT CONCLUSION
# ==============================================================================

doc.add_page_break()

add_heading('DOCUMENT END', 1)

conclusion = doc.add_paragraph()
conclusion_run = conclusion.add_run('This technical architecture document provides a comprehensive overview of the PII Masking Tool system design, focusing on:')
conclusion.add_run('\n\n• 3-tier architecture with clear separation of concerns')
conclusion.add_run('\n• JWT-based authentication for stateless security')
conclusion.add_run('\n• Smart PII filtering based on SQL data types')
conclusion.add_run('\n• ETL-style masking execution with transaction safety')
conclusion.add_run('\n• Role-based access control for authorization')
conclusion.add_run('\n• Production-ready deployment patterns')
conclusion.add_run('\n\n')
conclusion.add_run('The architecture prioritizes:')
conclusion.add_run('\n\n• Data Integrity: Constraint validation before insertion')
conclusion.add_run('\n• Security: Multiple layers of protection (network, application, data)')
conclusion.add_run('\n• Maintainability: Clear component boundaries and documentation')
conclusion.add_run('\n• Scalability: Horizontal scaling support for frontend and backend')
conclusion.add_run('\n• User Experience: Intelligent UI that prevents configuration errors')
conclusion.add_run('\n\n')
conclusion.add_run('All components work together to provide a robust, secure, and user-friendly solution for PII data masking across databases.')

doc.add_paragraph()
doc.add_paragraph()

footer_info = doc.add_paragraph()
footer_info.add_run('Document Version: 1.0\n').bold = True
footer_info.add_run('Last Updated: January 2025\n').bold = True
footer_info.add_run('Author: Technical Architecture Team\n').bold = True
footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save final document
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print(f"\nSection 10 added successfully!")
print(f"="*60)
print(f"DOCUMENT COMPLETED!")
print(f"="*60)
print(f"File saved: {output_path}")
print(f"\nAll 10 sections have been added to the document:")
print(f"  1. Executive Summary")
print(f"  2. System Overview")
print(f"  3. High-Level Architecture")
print(f"  4. Component Architecture")
print(f"  5. Data Flow Diagrams")
print(f"  6. Technology Stack")
print(f"  7. Database Design")
print(f"  8. Security Architecture")
print(f"  9. Deployment Architecture")
print(f"  10. Integration Architecture")
print(f"\nYou can now open the document in Microsoft Word!")
