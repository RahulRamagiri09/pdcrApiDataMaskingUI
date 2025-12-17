from docx import Document
from docx.shared import Pt

# Open existing document
doc = Document(r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx")

def add_heading(text, level=1):
    doc.add_heading(text, level)

def add_text(text):
    return doc.add_paragraph(text)

def add_diagram(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

print("Adding Section 6: Technology Stack...")

# ==============================================================================
# SECTION 6: TECHNOLOGY STACK
# ==============================================================================

add_heading('6. TECHNOLOGY STACK', 1)

add_heading('6.1 Technology Architecture', 2)

tech_diagram = """
TECHNOLOGY LAYERS

┌─────────────────────────────────────────────────────────┐
│                  PRESENTATION LAYER                     │
├─────────────────────────────────────────────────────────┤
│  React 18.3.1           Material-UI v5                  │
│  JavaScript Library     UI Component Library            │
│  • Virtual DOM          • Pre-built components          │
│  • Component-based      • Material Design               │
│  • Hooks for state      • Responsive layouts            │
│                                                         │
│  React Router v6        Axios                           │
│  Client Routing         HTTP Client                     │
│  • SPA navigation       • Promise-based                 │
│  • No page reloads      • Request/response interceptors │
└─────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────┐
│                  APPLICATION LAYER                      │
├─────────────────────────────────────────────────────────┤
│  FastAPI (Python)       Pydantic                        │
│  Web Framework          Data Validation                 │
│  • High performance     • Type checking                 │
│  • Auto OpenAPI docs    • Request validation            │
│  • Async support        • Schema generation             │
│                                                         │
│  PyJWT                  pyodbc                          │
│  Authentication         Database Driver                 │
│  • Token generation     • SQL Server connection         │
│  • Token validation     • ODBC protocol                 │
│                                                         │
│  Uvicorn               Gunicorn (Production)            │
│  ASGI Server           Process Manager                  │
│  • Async handling      • Multi-worker                   │
│  • Fast I/O            • Auto-restart                   │
└─────────────────────────────────────────────────────────┘

┌─────────────────────────────────────────────────────────┐
│                      DATA LAYER                         │
├─────────────────────────────────────────────────────────┤
│  Microsoft SQL Server 2016+                             │
│  • ACID transactions    • Query optimization            │
│  • Referential integrity• Indexes for performance       │
│  • Constraint enforcement                               │
│  • Backup and recovery  • High availability (Always-On) │
│                                                         │
│  ODBC Driver 17 for SQL Server                          │
│  • Native SQL Server protocol                           │
│  • SSL/TLS encryption   • Connection pooling            │
└─────────────────────────────────────────────────────────┘
"""
add_diagram(tech_diagram)

add_heading('6.2 Technology Selection Rationale', 2)

add_heading('Frontend Technologies', 3)

add_text('React 18.3.1:')
add_text('Why Chosen: Industry-standard JavaScript library for building user interfaces')
add_text('Benefits: Component reusability, virtual DOM for fast updates, large ecosystem, strong community support, easy to find developers, hooks simplify state management')
add_text('Use Cases: Interactive forms (workflow wizard), managing complex UI state (column mappings), real-time updates (execution monitoring)')
add_text('')

add_text('Material-UI (MUI) v5:')
add_text('Why Chosen: Comprehensive React component library implementing Material Design')
add_text('Benefits: Professional consistent UI out of the box, responsive components for mobile and desktop, accessible WCAG-compliant components, customizable theming, pre-built complex components (tables, dialogs, steppers), reduces custom CSS development')
add_text('Use Cases: Data tables for workflows and executions, multi-step wizard, dialogs for connection testing, navigation components')
add_text('')

add_text('React Router v6:')
add_text('Why Chosen: Standard routing library for React single-page applications')
add_text('Benefits: Client-side navigation without page reloads (fast UX), dynamic route parameters, route protection with authentication guards, nested routing for complex layouts, browser history management')
add_text('Use Cases: Protecting routes requiring authentication, dynamic workflow detail pages, multi-step form navigation')
add_text('')

add_text('Axios:')
add_text('Why Chosen: Popular promise-based HTTP client for browsers')
add_text('Benefits: Request/response interceptors (auto-add JWT tokens), automatic JSON transformation, error handling with detailed messages, request cancellation support, better API than native fetch()')
add_text('Use Cases: All API communication with backend, automatic JWT token injection, global error handling (401 auto-logout)')

add_heading('Backend Technologies', 3)

add_text('FastAPI (Python):')
add_text('Why Chosen: Modern, high-performance Python web framework')
add_text('Benefits: Automatic OpenAPI documentation (interactive API testing), type hints for better code quality, async support for high concurrency, fast performance comparable to Node.js and Go, built-in request validation with Pydantic, easy dependency injection, Python rich ecosystem for data processing')
add_text('Use Cases: REST API endpoints for frontend, PII masking algorithm implementation, database connection orchestration, workflow execution engine')
add_text('')

add_text('Pydantic:')
add_text('Why Chosen: Data validation using Python type annotations')
add_text('Benefits: Automatic request validation, clear error messages for invalid data, schema generation for OpenAPI docs, type safety reduces bugs, easy to define complex data models')
add_text('Use Cases: Validate workflow creation payloads, ensure correct data types in API requests, generate API documentation schemas')
add_text('')

add_text('PyJWT:')
add_text('Why Chosen: Industry-standard JWT implementation for Python')
add_text('Benefits: Stateless authentication (no session storage), cryptographically signed tokens, configurable expiration, small payload size, cross-domain compatible')
add_text('Use Cases: Generate JWT tokens on login, validate tokens on every API request, extract user information from tokens')
add_text('')

add_text('pyodbc:')
add_text('Why Chosen: Python database driver for ODBC connections')
add_text('Benefits: Native SQL Server support, parameterized queries prevent SQL injection, connection pooling for performance, support for all SQL Server features, platform independent (Windows, Linux, macOS)')
add_text('Use Cases: Connect to application database, connect to user source/target databases, execute queries for schema discovery, execute masking transformations')

add_heading('Database Technology', 3)

add_text('Microsoft SQL Server 2016+:')
add_text('Why Chosen: Enterprise-grade relational database management system')
add_text('Benefits: ACID compliance ensures data integrity, robust transaction support with rollback on failure, advanced query optimizer, rich constraint system (NOT NULL, UNIQUE, FK), system catalog for schema discovery, high availability with Always-On, backup and point-in-time recovery, enterprise support and tooling')
add_text('Use Cases: Store application metadata (users, workflows, connections), audit trail (execution history), source databases (production data with PII), target databases (masked data for non-production)')

doc.add_page_break()

# ==============================================================================
# SECTION 7: DATABASE DESIGN
# ==============================================================================

print("Adding Section 7: Database Design...")

add_heading('7. DATABASE DESIGN', 1)

add_heading('7.1 Entity Relationship Diagram', 2)

er_diagram = """
APPLICATION DATABASE ER DIAGRAM

                ┌──────────────────┐
                │      users       │
                ├──────────────────┤
                │ PK  id           │
                │     username     │
                │     password_hash│
                │     email        │
                │     role         │
                │     created_at   │
                └────────┬─────────┘
                         │
                         │ 1:N (creates)
                         │
        ┌────────────────┼────────────────┐
        │                │                │
        ▼                ▼                ▼
┌──────────────┐  ┌──────────────┐  ┌──────────────┐
│ connections  │  │  workflows   │  │  executions  │
├──────────────┤  ├──────────────┤  ├──────────────┤
│ PK  id       │  │ PK  id       │  │ PK  id       │
│     name     │◄─┤ FK  source_  │  │ FK  workflow_│
│     type     │  │     conn_id  │  │     id       │
│     server   │  │ FK  target_  │  │     status   │
│     database │  │     conn_id  │──┤     rows_    │
│     username │  │     status   │  │     processed│
│     password │  │ FK  created_ │  │     rows_    │
│ FK  created_ │  │     by       │  │     masked   │
│     by       │  │     created_ │  │     error_   │
│     created_ │  │     at       │  │     message  │
│     at       │  └──────┬───────┘  │     started_ │
└──────────────┘         │          │     at       │
                         │ 1:N      │     completed│
                         │          │     _at      │
                         ▼          │ FK  executed_│
              ┌──────────────────┐  │     by       │
              │ workflow_mappings│  └──────────────┘
              ├──────────────────┤
              │ PK  id           │
              │ FK  workflow_id  │
              │     source_      │
              │     schema       │
              │     source_table │
              │     target_      │
              │     schema       │
              │     target_table │
              │     column_      │
              │     mappings     │
              │     (JSON)       │
              │     created_at   │
              └──────────────────┘

RELATIONSHIP CARDINALITY:
• users (1) ───< (N) connections
  One user can create multiple database connections

• users (1) ───< (N) workflows
  One user can create multiple workflows

• users (1) ───< (N) executions
  One user can trigger multiple workflow executions

• connections (1) ───< (N) workflows (as source)
  One connection can be source for multiple workflows

• connections (1) ───< (N) workflows (as target)
  One connection can be target for multiple workflows

• workflows (1) ───< (N) workflow_mappings
  One workflow can have multiple table mappings

• workflows (1) ───< (N) executions
  One workflow can have multiple execution history records
"""
add_diagram(er_diagram)

add_heading('7.2 Database Schema Descriptions', 2)

add_heading('Table: users', 3)
add_text('Purpose: Store user account information for authentication and authorization')
add_text('')
add_text('Columns:')
add_text('- id (PK, INT, Auto-increment): Unique user identifier')
add_text('- username (VARCHAR(50), UNIQUE, NOT NULL): Login username')
add_text('- password_hash (VARCHAR(255), NOT NULL): Bcrypt hashed password (never plaintext)')
add_text('- email (VARCHAR(100), UNIQUE): User email address for notifications')
add_text('- role (VARCHAR(20), NOT NULL, DEFAULT "user"): Permission level ("admin" or "user")')
add_text('- created_at (DATETIME, NOT NULL, DEFAULT CURRENT_TIMESTAMP): Account creation timestamp')
add_text('')
add_text('Constraints: Username must be unique, role must be either "admin" or "user", password hash must use secure hashing (bcrypt with cost factor 12)')
add_text('')
add_text('Indexes: Primary key on id, unique index on username for fast login lookups, unique index on email')

add_heading('Table: connections', 3)
add_text('Purpose: Store database connection configurations for both source and target databases')
add_text('')
add_text('Columns:')
add_text('- id (PK, INT, Auto-increment): Unique connection identifier')
add_text('- name (VARCHAR(100), UNIQUE, NOT NULL): User-friendly connection name')
add_text('- connection_type (VARCHAR(10), NOT NULL): "source" (read from) or "target" (write to)')
add_text('- server (VARCHAR(255), NOT NULL): SQL Server hostname or IP address')
add_text('- database_name (VARCHAR(128), NOT NULL): Database name on the server')
add_text('- username (VARCHAR(128), NOT NULL): Database authentication username')
add_text('- password (VARCHAR(255), NOT NULL): Encrypted database password (AES-256)')
add_text('- created_by (FK → users.id, NOT NULL): User who created this connection')
add_text('- created_at (DATETIME, NOT NULL, DEFAULT CURRENT_TIMESTAMP): Creation timestamp')
add_text('')
add_text('Constraints: Connection name must be unique, connection type must be "source" or "target", foreign key to users table, password encrypted at rest (never stored in plaintext)')
add_text('')
add_text('Security Notes: Passwords are encrypted using AES-256 with encryption key stored in environment variables. Passwords never returned in API responses (masked with "******").')

add_heading('Table: workflows', 3)
add_text('Purpose: Store workflow configuration metadata (high-level workflow information)')
add_text('')
add_text('Columns:')
add_text('- id (PK, INT, Auto-increment): Unique workflow identifier')
add_text('- name (VARCHAR(100), UNIQUE, NOT NULL): User-friendly workflow name')
add_text('- description (TEXT): Optional description of workflow purpose')
add_text('- status (VARCHAR(20), NOT NULL, DEFAULT "active"): "active" or "inactive"')
add_text('- source_connection_id (FK → connections.id, NOT NULL): Source database connection')
add_text('- target_connection_id (FK → connections.id, NOT NULL): Target database connection')
add_text('- created_by (FK → users.id, NOT NULL): User who created workflow')
add_text('- created_at (DATETIME, NOT NULL, DEFAULT CURRENT_TIMESTAMP): Creation timestamp')
add_text('- updated_at (DATETIME): Last modification timestamp')
add_text('')
add_text('Constraints: Workflow name must be unique, status must be "active" or "inactive", source and target connections must be different (CHECK constraint), foreign keys to connections and users tables')
add_text('')
add_text('Business Rules: Only active workflows can be executed, source and target connections cannot be the same database, workflow name must be descriptive and unique')

add_heading('Table: workflow_mappings', 3)
add_text('Purpose: Store detailed table and column mapping rules for workflows')
add_text('')
add_text('Columns:')
add_text('- id (PK, INT, Auto-increment): Unique mapping identifier')
add_text('- workflow_id (FK → workflows.id, NOT NULL): Parent workflow')
add_text('- source_schema (VARCHAR(128), NOT NULL): Source database schema name (e.g., "dbo")')
add_text('- source_table (VARCHAR(128), NOT NULL): Source table name')
add_text('- target_schema (VARCHAR(128), NOT NULL): Target database schema name')
add_text('- target_table (VARCHAR(128), NOT NULL): Target table name')
add_text('- column_mappings (NVARCHAR(MAX), NOT NULL): JSON array of column mapping rules')
add_text('- created_at (DATETIME, NOT NULL, DEFAULT CURRENT_TIMESTAMP): Creation timestamp')
add_text('')
add_text('Constraints: Foreign key to workflows table with CASCADE DELETE, column_mappings must be valid JSON (CHECK constraint: ISJSON(column_mappings) = 1), one workflow can have multiple mappings for different tables')
add_text('')
add_text('JSON Structure for column_mappings:')
add_text('[{"source_column": "first_name", "pii_attribute": "first_name", "target_column": "first_name", "data_type": "varchar"}, {"source_column": "age", "pii_attribute": "random_number", "target_column": "age", "data_type": "int"}]')
add_text('')
add_text('Why JSON Column: Flexible for different numbers of columns across tables, no schema changes needed when adding new PII attributes, easy to query and parse in application code, SQL Server provides native JSON functions')

add_heading('Table: executions', 3)
add_text('Purpose: Store workflow execution history and audit trail')
add_text('')
add_text('Columns:')
add_text('- id (PK, INT, Auto-increment): Unique execution identifier')
add_text('- workflow_id (FK → workflows.id, NOT NULL): Workflow that was executed')
add_text('- status (VARCHAR(20), NOT NULL, DEFAULT "running"): "running", "success", or "failed"')
add_text('- rows_processed (INT, DEFAULT 0): Total rows read from source database')
add_text('- rows_masked (INT, DEFAULT 0): Total rows successfully masked and written')
add_text('- error_message (TEXT): Error details if status is "failed"')
add_text('- started_at (DATETIME, NOT NULL, DEFAULT CURRENT_TIMESTAMP): Execution start time')
add_text('- completed_at (DATETIME): Execution completion time (NULL while running)')
add_text('- executed_by (FK → users.id, NOT NULL): User who triggered execution')
add_text('')
add_text('Constraints: Foreign keys to workflows and users tables with CASCADE DELETE, status must be "running", "success", or "failed", rows_masked cannot exceed rows_processed (CHECK constraint)')
add_text('')
add_text('Audit Trail: Every execution creates a permanent record, complete history preserved even if workflow is modified, user attribution for compliance and accountability, error messages stored for troubleshooting')

add_heading('7.3 Database Relationships Description', 2)

add_text('Parent-Child Relationships:')
add_text('')
add_text('users → connections (1:N): A user can create multiple database connections. Each connection belongs to exactly one user. If user is deleted, connections remain (NO ACTION constraint for data integrity).')
add_text('')
add_text('users → workflows (1:N): A user can create multiple workflows. Each workflow belongs to exactly one user. User attribution for ownership and access control.')
add_text('')
add_text('users → executions (1:N): A user can execute workflows multiple times. Each execution is attributed to the user who triggered it. Audit trail shows who executed what and when.')
add_text('')
add_text('connections → workflows (Source) (1:N): A connection can be the source for multiple workflows. Each workflow has exactly one source connection. Allows reusing database connections across workflows.')
add_text('')
add_text('connections → workflows (Target) (1:N): A connection can be the target for multiple workflows. Each workflow has exactly one target connection. Same connection can be source in one workflow, target in another.')
add_text('')
add_text('workflows → workflow_mappings (1:N): A workflow can have multiple table mappings. Allows masking data from multiple tables in one workflow. Each mapping belongs to exactly one workflow. CASCADE DELETE: When workflow deleted, all mappings are deleted.')
add_text('')
add_text('workflows → executions (1:N): A workflow can be executed many times. Each execution records one run of the workflow. Complete history preserved for auditing. CASCADE DELETE: When workflow deleted, execution history is deleted.')

add_heading('7.4 Data Integrity Rules', 2)

add_text('Referential Integrity: All foreign keys enforced at database level. Prevents orphaned records (e.g., execution without workflow). Cascade delete for dependent data (mappings, executions).')
add_text('')
add_text('Constraint Validation: NOT NULL constraints ensure required data is always present. UNIQUE constraints prevent duplicate names/usernames. CHECK constraints validate enum values (status, role, connection_type). CHECK constraints ensure logical consistency (e.g., rows_masked <= rows_processed).')
add_text('')
add_text('Data Quality Rules: Passwords must be hashed before storage. Connection credentials encrypted at rest. JSON columns validated for correct structure. Timestamps automatically managed by database.')

# Save progress
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print(f"Sections 6-7 added successfully!")
print(f"Document saved: {output_path}")
