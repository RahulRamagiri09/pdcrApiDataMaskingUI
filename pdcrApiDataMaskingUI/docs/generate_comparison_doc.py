from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

# Create a new document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def add_heading(text, level=1):
    """Add a heading with specified level"""
    heading = doc.add_heading(text, level)
    return heading

def add_text(text):
    """Add a normal paragraph"""
    return doc.add_paragraph(text)

def add_bullet(text):
    """Add a bullet point"""
    return doc.add_paragraph(text, style='List Bullet')

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_table(headers, rows, header_color="003366"):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def add_colored_box(text, color="E8F4FD", border_color="3498DB"):
    """Add a colored box with text"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    # Note: python-docx doesn't directly support background colors on paragraphs
    # We'll use a table with one cell instead
    return p

def create_flow_table(steps, color="27AE60"):
    """Create a horizontal flow diagram using a table"""
    table = doc.add_table(rows=1, cols=len(steps))
    table.style = 'Table Grid'

    for i, step in enumerate(steps):
        cell = table.rows[0].cells[i]
        cell.text = step
        set_cell_shading(cell, color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    return table

# ==============================================================================
# TITLE PAGE
# ==============================================================================
print("Creating title page...")

title = doc.add_paragraph()
title_run = title.add_run('MasterCraft vs PII Masking Tool')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Comprehensive Comparison Document')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(5):
    doc.add_paragraph()

info = doc.add_paragraph()
info.add_run('Document Version: 1.0\n\n').bold = True
info.add_run('Date: January 2026\n\n').bold = True
info.add_run('Purpose: Compare MasterCraft and PII Masking Tool approaches to data masking\n\n').bold = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================
print("Adding Table of Contents...")

add_heading('Table of Contents', 1)
toc_items = [
    'Part 1: MasterCraft',
    '    - What is MasterCraft?',
    '    - Key Characteristics',
    '    - 4-Job Pipeline',
    '    - Job Flow Details',
    '    - Large Volume Handling',
    'Part 2: PII Masking Tool',
    '    - What is PII Masking Tool?',
    '    - Key Characteristics',
    '    - Complete User Journey',
    '    - Batch Execution System',
    '    - Execution State Machine',
    'Part 3: Side-by-Side Comparison',
    '    - Architecture Comparison',
    '    - Feature Comparison Table',
    '    - Use Case Comparison',
    '    - Key Advantages',
]
for item in toc_items:
    add_text(item)

doc.add_page_break()

# ==============================================================================
# PART 1: MASTERCRAFT
# ==============================================================================
print("Creating Part 1: MasterCraft...")

# Title
part1_title = doc.add_paragraph()
part1_run = part1_title.add_run('PART 1: MASTERCRAFT')
part1_run.font.size = Pt(24)
part1_run.font.bold = True
part1_run.font.color.rgb = RGBColor(192, 57, 43)  # Red color for MasterCraft
part1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# What is MasterCraft?
add_heading('What is MasterCraft?', 1)

desc = doc.add_paragraph()
desc.add_run('MasterCraft is an ').font.size = Pt(11)
bold_run = desc.add_run('enterprise data replication platform')
bold_run.bold = True
bold_run.font.size = Pt(11)
desc.add_run(' that copies production data to non-production environments while applying data masking. It uses a ').font.size = Pt(11)
bold_run2 = desc.add_run('4-job sequential pipeline')
bold_run2.bold = True
bold_run2.font.size = Pt(11)
desc.add_run(' to move, transform, and load data between different database environments.').font.size = Pt(11)

doc.add_paragraph()

# Core Purpose Box
purpose_para = doc.add_paragraph()
purpose_para.add_run('Core Purpose: ').bold = True
purpose_para.add_run('"Copy masked production data from source database to a separate non-production database"')

doc.add_paragraph()

# Key Characteristics
add_heading('Key Characteristics', 2)
mc_chars = [
    ('Data Movement', 'Copies data from one database to another'),
    ('File-Based', 'Creates intermediate .dat files on server'),
    ('Sequential Jobs', '4 jobs must run in order'),
    ('Manual Intervention', 'Requires DBA approval via email'),
    ('Environment Sync', 'Keeps non-prod in sync with production'),
]
for char, desc in mc_chars:
    p = doc.add_paragraph()
    p.add_run(f'{char}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# 4-Job Pipeline
add_heading('MasterCraft 4-Job Pipeline', 2)

add_text('MasterCraft uses a sequential 4-job pipeline to process data:')
doc.add_paragraph()

# Create pipeline flow table
pipeline_steps = ['JOB 1\nSchema Sync', 'JOB 2\nBackup', 'JOB 3\nMasking', 'JOB 4\nLoad Data']
create_flow_table(pipeline_steps, "C0392B")

doc.add_paragraph()
doc.add_paragraph()

# Job Details Table
add_heading('Job Details', 3)
job_headers = ['Job', 'Name', 'Purpose', 'Output', 'Bottleneck']
job_rows = [
    ['1', 'Schema Sync', 'Synchronize source and target schemas', 'Schema metadata, Project ID', 'No'],
    ['2', 'Backup', 'Extract data to portable format', '.dat backup files on server', 'No'],
    ['3', 'Masking', 'Apply masking rules to sensitive data', 'Masked .dat files', 'Yes - DBA Email'],
    ['4', 'Load Data', 'Load masked data to non-production', 'Populated target database', 'Yes - DBA Email'],
]
create_table(job_headers, job_rows, "C0392B")

doc.add_paragraph()

# Warning about DBA
warning = doc.add_paragraph()
warning.add_run('WARNING: ').bold = True
warning.add_run('Jobs 3 and 4 require DBA email approval before proceeding. This creates significant delays in the workflow.')

doc.add_page_break()

# Job Flow Details
add_heading('Job Flow Details', 2)

# Job 1 Details
add_heading('Job 1: Schema Sync', 3)
j1_steps = ['Fetch Source\nSchema', 'Compare\nSchemas', 'Update Target\nSchema', 'Store\nMetadata']
create_flow_table(j1_steps, "E74C3C")
doc.add_paragraph()
add_bullet('Purpose: Ensure source and target databases have matching schema')
add_bullet('Output: Schema synchronized, Project ID created')

doc.add_paragraph()

# Job 2 Details
add_heading('Job 2: Backup', 3)
j2_steps = ['Extract Data\nfrom DB', 'Convert to\n.dat Format', 'Store on\nFile Server']
create_flow_table(j2_steps, "E74C3C")
doc.add_paragraph()
add_bullet('Purpose: Create backup of source data in portable format')
add_bullet('Output: .dat backup files stored on file server')

doc.add_paragraph()

# Job 3 Details
add_heading('Job 3: Masking', 3)
j3_steps = ['Create\nProject ID', 'EMAIL TO DBA\n(WAIT)', 'Apply Masking\nRules', 'Store Masked\n.dat Files']
create_flow_table(j3_steps, "F39C12")
doc.add_paragraph()
add_bullet('Purpose: Apply masking rules to sensitive data')
add_bullet('Output: Masked .dat files stored on server')
warning3 = doc.add_paragraph()
warning3.add_run('BOTTLENECK: Requires DBA email approval to disable constraints before masking').bold = True

doc.add_paragraph()

# Job 4 Details
add_heading('Job 4: Load Data', 3)
j4_steps = ['EMAIL TO DBA\n(WAIT)', 'Connect to\nNon-Prod', 'Load .dat\nto Tables', 'Validate &\nEnable Constr.']
create_flow_table(j4_steps, "F39C12")
doc.add_paragraph()
add_bullet('Purpose: Load masked data into non-production database')
add_bullet('Output: Non-prod database populated with masked data')
warning4 = doc.add_paragraph()
warning4.add_run('BOTTLENECK: Requires DBA email approval + manual issue resolution').bold = True

doc.add_page_break()

# Large Volume Handling
add_heading('Large Volume Handling', 2)

add_text('When large data volume is detected, MasterCraft uses a PARALLEL SUBSET flow:')
doc.add_paragraph()

lv_steps = [
    '1. Large Volume Detected',
    '2. Enable Batch Setup (Manual configuration required)',
    '3. Split Records into Batches (Batch 1, Batch 2, ... Batch N)',
    '4. Create Subset Jobs (Each batch becomes a separate subset job)',
    '5. Load -> Mask -> Load (Per Subset)',
    '6. Complete',
]
for step in lv_steps:
    add_bullet(step)

doc.add_paragraph()
note = doc.add_paragraph()
note.add_run('Note: ').bold = True
note.add_run('Batch setup in MasterCraft requires MANUAL configuration and creates separate subset jobs for each batch.')

doc.add_paragraph()

# MasterCraft Summary
add_heading('MasterCraft Summary', 2)
mc_summary_headers = ['Aspect', 'Description']
mc_summary_rows = [
    ['Type', 'Data replication & migration platform'],
    ['Data Flow', 'Source DB -> .dat files -> Target DB'],
    ['Jobs Required', '4 sequential jobs'],
    ['DBA Dependency', 'Required (email approval for constraints)'],
    ['Execution Control', 'Start only (no pause/resume)'],
    ['Storage Required', 'Yes (.dat files on server)'],
    ['Best For', 'Full database refresh to non-prod environments'],
]
create_table(mc_summary_headers, mc_summary_rows, "C0392B")

doc.add_page_break()

# ==============================================================================
# PART 2: PII MASKING TOOL
# ==============================================================================
print("Creating Part 2: PII Masking Tool...")

# Title
part2_title = doc.add_paragraph()
part2_run = part2_title.add_run('PART 2: PII MASKING TOOL')
part2_run.font.size = Pt(24)
part2_run.font.bold = True
part2_run.font.color.rgb = RGBColor(39, 174, 96)  # Green color for PII Tool
part2_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# What is PII Masking Tool?
add_heading('What is PII Masking Tool?', 1)

desc2 = doc.add_paragraph()
desc2.add_run('PII Masking Tool is a ').font.size = Pt(11)
bold_run3 = desc2.add_run('specialized data privacy platform')
bold_run3.bold = True
bold_run3.font.size = Pt(11)
desc2.add_run(' designed for ').font.size = Pt(11)
bold_run4 = desc2.add_run('in-place masking')
bold_run4.bold = True
bold_run4.font.size = Pt(11)
desc2.add_run(' of Personally Identifiable Information (PII) directly within databases. It transforms sensitive data without moving it to another location.').font.size = Pt(11)

doc.add_paragraph()

# Core Purpose Box
purpose_para2 = doc.add_paragraph()
purpose_para2.add_run('Core Purpose: ').bold = True
purpose_para2.add_run('"Fast, secure, in-place PII masking with zero data movement for compliance and privacy"')

doc.add_paragraph()

# Key Characteristics
add_heading('Key Characteristics', 2)
pii_chars = [
    ('In-Place Masking', 'Data stays in the same database'),
    ('No File Storage', 'Direct database updates, no .dat files'),
    ('Single Workflow', 'One workflow executes masking'),
    ('Fully Automated', 'No DBA email approvals needed'),
    ('Batch Execution', 'Built-in batch processing with pause/resume'),
]
for char, desc in pii_chars:
    p = doc.add_paragraph()
    check_mark = p.add_run('* ')
    check_mark.font.color.rgb = RGBColor(39, 174, 96)
    p.add_run(f'{char}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# Architecture
add_heading('PII Masking Tool Architecture', 2)

arch_text = doc.add_paragraph()
arch_text.add_run('In-Place Transformation: ').bold = True
arch_text.add_run('Data never leaves the database')

doc.add_paragraph()

# Data transformation example
add_heading('Data Transformation Example', 3)
transform_headers = ['Original Data', 'Masked Data']
transform_rows = [
    ['John Smith', 'XXXX XXXXX'],
    ['john@email.com', 'XXX@XXXXX.com'],
    ['123-45-6789', 'XXX-XX-XXXX'],
    ['555-123-4567', 'XXX-XXX-XXXX'],
]
create_table(transform_headers, transform_rows, "27AE60")

doc.add_paragraph()

# Key Points
key_points = doc.add_paragraph()
key_points.add_run('Key Points:').bold = True
add_bullet('No files created during masking process')
add_bullet('No DBA approval required')
add_bullet('Real-time control (pause/resume/stop)')
add_bullet('Data masked in-place within same database')

doc.add_page_break()

# Complete User Journey
add_heading('Complete User Journey', 2)

journey_headers = ['Step', 'Name', 'Description']
journey_rows = [
    ['1', 'Authentication', 'JWT authentication, Role-based access (Admin/Privilege/General/Support)'],
    ['2', 'Connection Setup', 'Add DB details, Test connection, Save config. Supports: Azure SQL, SQL Server, PostgreSQL, Oracle'],
    ['3', 'Workflow Creation', '4-step wizard: Basic Info -> Select Table -> Map Columns -> Review & Save'],
    ['4', 'Constraint Validation', 'Automatic pre-execution checks: Primary Keys, Foreign Keys, Unique Constraints, Check Constraints, Triggers, Indexes'],
    ['5', 'Preview Masking', 'Optional: View original vs masked data before execution'],
    ['6', 'Batch Execution', 'Execute with real-time control (Start/Pause/Resume/Stop)'],
    ['7', 'Audit & Logging', 'Full audit trail: Who executed what, when, how many records affected'],
]
create_table(journey_headers, journey_rows, "27AE60")

doc.add_paragraph()

# Important Notes
important = doc.add_paragraph()
important.add_run('Important: ').bold = True
important.add_run('NO DBA EMAIL REQUIRED - All validation happens automatically. Issues shown in UI before execution.')

doc.add_paragraph()

# Batch Execution System
add_heading('Batch Execution System', 2)

add_text('The PII Masking Tool includes automatic batch processing:')
doc.add_paragraph()

batch_headers = ['Feature', 'Description']
batch_rows = [
    ['Automatic Batching', 'Records automatically split into configurable batch sizes'],
    ['Example', 'Total: 1,000,000 records | Batch Size: 10,000 | Total Batches: 100'],
    ['Independent Commits', 'Each batch committed independently'],
    ['Failure Handling', 'Failed batch rolls back, previous batches preserved'],
    ['Progress Tracking', 'Real-time progress: Records processed, percentage, current batch, duration'],
]
create_table(batch_headers, batch_rows, "27AE60")

doc.add_paragraph()

# Execution Controls
add_heading('Execution Control Options', 3)
control_headers = ['Control', 'Action']
control_rows = [
    ['START', 'Begin execution from batch 1'],
    ['PAUSE', 'Complete current batch, then pause. "Execution paused successfully at batch 45"'],
    ['RESUME', 'Continue from next batch after pause point. "Execution resumed successfully from batch 46"'],
    ['STOP', 'Cancel execution immediately. "Execution stopped. 45 batches completed."'],
]
create_table(control_headers, control_rows, "3498DB")

doc.add_page_break()

# Execution State Machine
add_heading('Execution State Machine', 2)

add_text('The execution follows a state machine with the following states:')
doc.add_paragraph()

state_headers = ['State', 'Description', 'Transitions']
state_rows = [
    ['QUEUED', 'Waiting to start', 'Start -> RUNNING'],
    ['RUNNING', 'Processing batches', 'Pause -> PAUSED, Complete -> COMPLETED, Error -> FAILED'],
    ['PAUSED', 'Waiting after pause', 'Resume -> RUNNING, Stop -> STOPPED'],
    ['COMPLETED', 'Successfully finished', 'Terminal state'],
    ['FAILED', 'Error occurred', 'Terminal state'],
    ['STOPPED', 'Cancelled by user', 'Terminal state'],
]
create_table(state_headers, state_rows, "27AE60")

doc.add_paragraph()

# PII Masking Tool Summary
add_heading('PII Masking Tool Summary', 2)
pii_summary_headers = ['Aspect', 'Description']
pii_summary_rows = [
    ['Type', 'In-place data masking platform'],
    ['Data Flow', 'Same DB -> Mask -> Same DB (no movement)'],
    ['Workflows', 'Single workflow execution'],
    ['DBA Dependency', 'None (fully automated)'],
    ['Execution Control', 'Start, Pause, Resume, Stop'],
    ['Storage Required', 'None (no intermediate files)'],
    ['Batch Processing', 'Built-in automatic batching'],
    ['Best For', 'PII compliance, data privacy, GDPR/CCPA/HIPAA'],
]
create_table(pii_summary_headers, pii_summary_rows, "27AE60")

doc.add_page_break()

# ==============================================================================
# PART 3: SIDE-BY-SIDE COMPARISON
# ==============================================================================
print("Creating Part 3: Side-by-Side Comparison...")

# Title
part3_title = doc.add_paragraph()
part3_run = part3_title.add_run('PART 3: SIDE-BY-SIDE COMPARISON')
part3_run.font.size = Pt(24)
part3_run.font.bold = True
part3_run.font.color.rgb = RGBColor(142, 68, 173)  # Purple color for Comparison
part3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Architecture Comparison
add_heading('Architecture Comparison', 1)

# MasterCraft Architecture
add_heading('MasterCraft Architecture', 2)
mc_arch_steps = ['Source DB', 'JOB 1', 'JOB 2', '.dat Files', 'JOB 3', 'JOB 4', 'Target DB']
create_flow_table(mc_arch_steps, "C0392B")
doc.add_paragraph()
add_bullet('4 sequential jobs required')
add_bullet('File storage required (.dat files)')
add_bullet('DBA email approvals needed')
add_bullet('No pause/resume capability')

doc.add_paragraph()

# PII Tool Architecture
add_heading('PII Masking Tool Architecture', 2)
pii_arch_steps = ['Single DB', 'Original Data', 'In-Place Masking', 'Masked Data', 'Same DB']
create_flow_table(pii_arch_steps, "27AE60")
doc.add_paragraph()
add_bullet('Single workflow execution')
add_bullet('NO file storage needed')
add_bullet('NO DBA approval required')
add_bullet('Built-in pause/resume/stop')

doc.add_page_break()

# Feature Comparison Table
add_heading('Feature Comparison Table', 1)

# Custom table with 3 columns
feature_table = doc.add_table(rows=1, cols=3)
feature_table.style = 'Table Grid'

# Header row
headers = ['Feature', 'MasterCraft', 'PII Masking Tool']
header_colors = ['003366', 'C0392B', '27AE60']
header_cells = feature_table.rows[0].cells
for i, (header, color) in enumerate(zip(headers, header_colors)):
    header_cells[i].text = header
    set_cell_shading(header_cells[i], color)
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

# Data rows
feature_rows = [
    ['Data Movement', 'Source -> Target', 'None (In-Place)'],
    ['Jobs Required', '4 sequential', '1 workflow'],
    ['File Storage', '.dat files required', 'None required'],
    ['DBA Approval', 'Required (Email)', 'Not required'],
    ['Batch Processing', 'Manual setup', 'Automatic'],
    ['Pause Execution', 'No', 'Yes (at batch boundary)'],
    ['Resume Execution', 'No', 'Yes (from last batch)'],
    ['Stop Execution', 'Limited', 'Yes (immediate)'],
    ['Preview Masking', 'No', 'Yes'],
    ['Constraint Validation', 'Post-load', 'Pre-execution'],
    ['Smart PII Filtering', 'No', 'Yes (by data type)'],
    ['RBAC', 'Not specified', '4-tier (Admin/Privilege/General/Support)'],
    ['Audit Trail', 'Basic', 'Comprehensive'],
    ['Execution Progress', 'Job-level', 'Batch-level + percentage'],
]

for row_data in feature_rows:
    row_cells = feature_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

doc.add_paragraph()
doc.add_page_break()

# Use Case Comparison
add_heading('Use Case Comparison', 1)

# When to use MasterCraft
add_heading('When to Use MasterCraft', 2)
mc_use_cases = [
    'Need full database copy to non-production',
    'Need schema synchronization between environments',
    'Need .dat file backups',
    'Complex multi-table orchestration',
    'Separate source and target databases required',
]
for uc in mc_use_cases:
    add_bullet(uc)

doc.add_paragraph()

# When to use PII Masking Tool
add_heading('When to Use PII Masking Tool', 2)
pii_use_cases = [
    'GDPR/CCPA/HIPAA compliance',
    'Quick PII data sanitization',
    'In-place masking without data movement',
    'Need execution control (pause/resume/stop)',
    'No DBA availability for approvals',
    'Need preview before execution',
    'Role-based team access required',
    'Complete audit trail needed',
]
for uc in pii_use_cases:
    add_bullet(uc)

doc.add_paragraph()

# One-Line Summary
add_heading('One-Line Summary', 1)

summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = 'Table Grid'

# Headers
sum_headers = ['Tool', 'Definition']
sum_header_cells = summary_table.rows[0].cells
for i, header in enumerate(sum_headers):
    sum_header_cells[i].text = header
    set_cell_shading(sum_header_cells[i], '003366')
    for paragraph in sum_header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

# Data
summary_data = [
    ['MasterCraft', '"Copy masked production data to non-production through 4-job ETL pipeline"'],
    ['PII Masking Tool', '"Fast, secure, in-place PII masking with batch execution and zero data movement"'],
]
for row_data in summary_data:
    row_cells = summary_table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

doc.add_page_break()

# Key Advantages of PII Masking Tool
add_heading('Key Advantages of PII Masking Tool', 1)

advantages = [
    ('1. IN-PLACE MASKING', 'Data stays in same database, no movement required'),
    ('2. AUTOMATIC BATCH EXECUTION', 'Built-in batching with configurable batch size. Each batch committed independently. Failed batch rolls back, previous batches preserved.'),
    ('3. REAL-TIME EXECUTION CONTROL', 'PAUSE: Stop at batch boundary, resume later. RESUME: Continue from last completed batch. STOP: Cancel execution immediately.'),
    ('4. ZERO DBA DEPENDENCY', 'No email approvals required. Automatic constraint validation before execution.'),
    ('5. SMART PII FILTERING', 'Shows only valid masking options based on column data type'),
    ('6. PREVIEW MASKING', 'See original vs masked data before execution'),
    ('7. COMPREHENSIVE RBAC', '4-tier permission system (Admin/Privilege/General/Support)'),
    ('8. COMPLETE AUDIT TRAIL', 'Full execution history with user attribution'),
]

for title, desc in advantages:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run('\n' + desc)
    doc.add_paragraph()

doc.add_paragraph()

# Flow Summary
add_heading('Flow Summary', 1)

add_heading('MasterCraft Flow:', 2)
mc_flow = doc.add_paragraph()
mc_flow.add_run('Ticket -> Connect -> Schema -> Backup -> Mask -> Load')
mc_flow2 = doc.add_paragraph()
mc_flow2.add_run('         |         |        |        |       |')
mc_flow3 = doc.add_paragraph()
mc_flow3.add_run('      DBA Setup  DBA Sync  Files   DBA    DBA Fix')
mc_flow4 = doc.add_paragraph()
mc_flow4.add_run('                          Store   Email')

doc.add_paragraph()

add_heading('PII Masking Tool Flow:', 2)
pii_flow = doc.add_paragraph()
pii_flow.add_run('Login -> Connect -> Workflow -> Validate -> Batch Execute -> Complete')
pii_flow2 = doc.add_paragraph()
pii_flow2.add_run('                                              |')
pii_flow3 = doc.add_paragraph()
pii_flow3.add_run('                                    [Pause] [Resume] [Stop]')

doc.add_paragraph()
doc.add_paragraph()

# Document End
end_text = doc.add_paragraph()
end_text.add_run('Document End').bold = True
end_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

footer = doc.add_paragraph()
footer.add_run('Document prepared for comparison analysis between MasterCraft and PII Masking Tool')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_date = doc.add_paragraph()
footer_date.add_run('January 2026')
footer_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
output_path = r"c:\Users\rahul.ramagiri\Documents\git\pdcrApiDataMaskingUI\pdcrApiDataMaskingUI\docs\MasterCraft_vs_PII_Masking_Tool_Comparison.docx"
doc.save(output_path)

print(f"\nDocument saved successfully: {output_path}")
print("MasterCraft vs PII Masking Tool Comparison DOCX created!")
