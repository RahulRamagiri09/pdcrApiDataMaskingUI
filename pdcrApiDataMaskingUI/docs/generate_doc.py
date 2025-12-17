from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

# Create document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Configure styles
styles = doc.styles

# Title style
title_style = styles['Title']
title_font = title_style.font
title_font.name = 'Calibri'
title_font.size = Pt(28)
title_font.bold = True
title_font.color.rgb = RGBColor(0, 51, 102)

# Heading 1 style
heading1_style = styles['Heading 1']
heading1_font = heading1_style.font
heading1_font.name = 'Calibri'
heading1_font.size = Pt(18)
heading1_font.bold = True
heading1_font.color.rgb = RGBColor(0, 51, 102)

# Heading 2 style
heading2_style = styles['Heading 2']
heading2_font = heading2_style.font
heading2_font.name = 'Calibri'
heading2_font.size = Pt(14)
heading2_font.bold = True
heading2_font.color.rgb = RGBColor(0, 102, 204)

# Heading 3 style
heading3_style = styles['Heading 3']
heading3_font = heading3_style.font
heading3_font.name = 'Calibri'
heading3_font.size = Pt(12)
heading3_font.bold = True

# ==============================================================================
# TITLE PAGE
# ==============================================================================

# Add title
title = doc.add_paragraph()
title_run = title.add_run('PII MASKING TOOL')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Add subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('ARCHITECTURE & DESIGN DOCUMENT')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Add document info
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.add_run('Version: 1.0\n\n').bold = True
info.add_run(f'Date: {datetime.datetime.now().strftime("%B %Y")}\n\n').bold = True
info.add_run('Document Type: Technical Architecture\n\n').bold = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Page break
doc.add_page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================

doc.add_heading('TABLE OF CONTENTS', 0)
doc.add_paragraph()

toc_entries = [
    ('1. Executive Summary', ''),
    ('2. System Overview', ''),
    ('3. High-Level Architecture', ''),
    ('4. Component Architecture', ''),
    ('5. Data Flow Diagrams', ''),
    ('6. Technology Stack', ''),
    ('7. Database Design', ''),
    ('8. Security Architecture', ''),
    ('9. Deployment Architecture', ''),
    ('10. Integration Architecture', ''),
]

for entry, page in toc_entries:
    toc_p = doc.add_paragraph(entry, style='List Number')
    toc_p.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

print("Creating document sections...")

# Save the document
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print(f"✓ Document created successfully: {output_path}")
