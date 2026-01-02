from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

# Create a new document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(text, level=1):
    """Add a heading with specified level"""
    heading = doc.add_heading(text, level)
    return heading

def add_text(text):
    """Add a normal paragraph"""
    return doc.add_paragraph(text)

def add_bullet(text):
    """Add a bullet point"""
    return doc.add_paragraph(text, style='List Bullet')

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_table(headers, rows, header_color="003366"):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

# ==============================================================================
# TITLE PAGE
# ==============================================================================
print("Creating title page...")

title = doc.add_paragraph()
title_run = title.add_run('RBAC Access Matrix')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('PII Masking Tool')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(5):
    doc.add_paragraph()

info = doc.add_paragraph()
info.add_run('Document Version: 4.0\n\n').bold = True
info.add_run('Date: January 2026\n\n').bold = True
info.add_run('Purpose: Define role-based access control for single database in-place masking\n\n').bold = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==============================================================================
# ROLE DEFINITIONS
# ==============================================================================
print("Adding Role Definitions...")

add_heading('Role Definitions', 1)

role_headers = ['Role', 'Description', 'Typical Users']
role_rows = [
    ['Admin', 'Full system administration', 'System admins, Database admins'],
    ['Privilege', 'Execute workflows and manage executions', 'Data engineers, Operations team'],
    ['General', 'Read-only monitoring', 'Business analysts, Auditors'],
    ['Support', 'Technical support (read-only)', 'Help desk, Support team'],
]
create_table(role_headers, role_rows)

doc.add_paragraph()
doc.add_page_break()

# ==============================================================================
# MODULE ACCESS MATRICES
# ==============================================================================
print("Adding Module Access Matrices...")

add_heading('Module Access Matrix', 1)

# 1. Server Connections Module
add_heading('1. Server Connections Module', 2)
conn_headers = ['Feature', 'Admin', 'Privilege', 'General', 'Support']
conn_rows = [
    ['View connection list', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Create new connection', 'Yes', 'No', 'No', 'No'],
    ['Edit connection details', 'Yes', 'No', 'No', 'No'],
    ['Delete connection', 'Yes', 'No', 'No', 'No'],
    ['Test connection', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View schemas/tables', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View table columns', 'Yes', 'Yes', 'Yes', 'Yes'],
]
create_table(conn_headers, conn_rows)
doc.add_paragraph()

# 2. Workflows Module
add_heading('2. Workflows Module', 2)
wf_headers = ['Feature', 'Admin', 'Privilege', 'General', 'Support']
wf_rows = [
    ['View workflow list', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Create new workflow', 'Yes', 'No', 'No', 'No'],
    ['Edit workflow', 'Yes', 'No', 'No', 'No'],
    ['Delete workflow', 'Yes', 'No', 'No', 'No'],
    ['Execute workflow', 'Yes', 'Yes', 'No', 'No'],
    ['View column mappings', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Edit column mappings', 'Yes', 'No', 'No', 'No'],
    ['View PII attributes', 'Yes', 'Yes', 'Yes', 'Yes'],
]
create_table(wf_headers, wf_rows)
doc.add_paragraph()

# 3. Execution Module
add_heading('3. Execution Module', 2)
exec_headers = ['Feature', 'Admin', 'Privilege', 'General', 'Support']
exec_rows = [
    ['View execution history', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Start execution', 'Yes', 'Yes', 'No', 'No'],
    ['Stop execution', 'Yes', 'Yes', 'No', 'No'],
    ['Pause execution', 'Yes', 'Yes', 'No', 'No'],
    ['Resume execution', 'Yes', 'Yes', 'No', 'No'],
    ['View execution status', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View execution logs', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View execution progress', 'Yes', 'Yes', 'Yes', 'Yes'],
]
create_table(exec_headers, exec_rows)
doc.add_paragraph()

doc.add_page_break()

# 4. Preview & Validation Module
add_heading('4. Preview & Validation Module', 2)
prev_headers = ['Feature', 'Admin', 'Privilege', 'General', 'Support']
prev_rows = [
    ['View masking preview', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View constraint checks', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View primary keys', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View foreign keys', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View unique constraints', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View check constraints', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View triggers', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View indexes', 'Yes', 'Yes', 'Yes', 'Yes'],
]
create_table(prev_headers, prev_rows)
doc.add_paragraph()

# 5. User Management Module
add_heading('5. User Management Module', 2)
user_headers = ['Feature', 'Admin', 'Privilege', 'General', 'Support']
user_rows = [
    ['View users', 'Yes', 'No', 'No', 'No'],
    ['Create users', 'Yes', 'No', 'No', 'No'],
    ['Assign roles', 'Yes', 'No', 'No', 'No'],
    ['View roles', 'Yes', 'No', 'No', 'No'],
    ['Create roles', 'Yes', 'No', 'No', 'No'],
]
create_table(user_headers, user_rows)
doc.add_paragraph()

# 6. Dashboard Module
add_heading('6. Dashboard Module', 2)
dash_headers = ['Feature', 'Admin', 'Privilege', 'General', 'Support']
dash_rows = [
    ['View dashboard', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View statistics', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['View recent activity', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Quick actions', 'Yes', 'Yes', 'No', 'No'],
]
create_table(dash_headers, dash_rows)
doc.add_paragraph()

doc.add_page_break()

# ==============================================================================
# SUMMARY TABLE
# ==============================================================================
print("Adding Summary Table...")

add_heading('Summary Table', 1)

summary_headers = ['Module', 'Admin Access', 'Privilege Access', 'General Access', 'Support Access']
summary_rows = [
    ['Server Connections', 'Full (CRUD + Test)', 'Read + Test', 'Read + Test', 'Read + Test'],
    ['Workflows', 'Full (CRUD + Execute)', 'Read + Execute', 'Read Only', 'Read Only'],
    ['Executions', 'Full (Start + Stop + Pause + Resume)', 'Start + Stop + Pause + Resume', 'Read Only', 'Read Only'],
    ['Preview & Validation', 'Full Read', 'Full Read', 'Full Read', 'Full Read'],
    ['User Management', 'Full (CRUD)', 'No Access', 'No Access', 'No Access'],
    ['Dashboard', 'Full Access', 'Full Access', 'View Only', 'View Only'],
]
create_table(summary_headers, summary_rows)

doc.add_paragraph()
doc.add_page_break()

# ==============================================================================
# PERMISSION DEFINITIONS
# ==============================================================================
print("Adding Permission Definitions...")

add_heading('Permission Definitions (Code Reference)', 1)

add_text('The following permissions are defined in src/utils/rbac.js:')
doc.add_paragraph()

add_heading('Admin Permissions', 2)
admin_perms = doc.add_paragraph()
admin_perms.add_run('''admin.access
connection.view, connection.create, connection.update, connection.delete, connection.test
workflow.view, workflow.create, workflow.update, workflow.delete, workflow.execute
execution.start, execution.view, execution.stop, execution.pause, execution.resume
preview.view, masking.view, columnMapping.view, constraint.view''')

add_heading('Privilege Permissions', 2)
priv_perms = doc.add_paragraph()
priv_perms.add_run('''connection.view, connection.test
workflow.view, workflow.execute
execution.start, execution.view, execution.stop, execution.pause, execution.resume
preview.view, columnMapping.view, constraint.view''')

add_heading('General Permissions', 2)
gen_perms = doc.add_paragraph()
gen_perms.add_run('''connection.view, connection.test
workflow.view
execution.view
preview.view, columnMapping.view, constraint.view''')

add_heading('Support Permissions', 2)
sup_perms = doc.add_paragraph()
sup_perms.add_run('''connection.view, connection.test
workflow.view
execution.view
preview.view, columnMapping.view, constraint.view''')

doc.add_page_break()

# ==============================================================================
# ROUTE PROTECTION
# ==============================================================================
print("Adding Route Protection...")

add_heading('Route Protection', 1)

add_text('The following routes are protected by RBAC:')
doc.add_paragraph()

route_headers = ['Route', 'Required Permission', 'Redirect on Deny']
route_rows = [
    ['/registeruser', 'admin.access', '/datamasking/dashboard'],
    ['/registerrole', 'admin.access', '/datamasking/dashboard'],
    ['/datamasking/workflows/create', 'workflow.create', '/datamasking/workflows'],
    ['/datamasking/workflows/:id/edit', 'workflow.update', '/datamasking/workflows'],
]
create_table(route_headers, route_rows)

doc.add_paragraph()

# ==============================================================================
# KEY DIFFERENCES
# ==============================================================================
print("Adding Key Differences...")

add_heading('Key Differences Between Roles', 1)

add_heading('Admin vs Privilege', 2)
add_bullet('Admin: Can CREATE, UPDATE, and DELETE workflows and connections')
add_bullet('Privilege: Can only VIEW and EXECUTE pre-configured workflows')

add_heading('Privilege vs General', 2)
add_bullet('Privilege: Can EXECUTE workflows and control executions (start/stop/pause/resume)')
add_bullet('General: Can only VIEW workflows and executions (cannot execute or control)')

add_heading('General vs Support', 2)
add_bullet('Identical permissions (both are read-only roles)')
add_bullet('Different organizational purpose (auditing vs troubleshooting)')

doc.add_paragraph()

# ==============================================================================
# IMPLEMENTATION REFERENCE
# ==============================================================================
print("Adding Implementation Reference...")

add_heading('Implementation Reference', 1)

add_heading('Frontend Files:', 2)
add_bullet('RBAC Logic: src/utils/rbac.js')
add_bullet('Protected UI Component: src/components/common/ProtectedAction.jsx')
add_bullet('Permission Hooks: src/hooks/usePermission.js')
add_bullet('Route Protection: src/components/ProtectedRoute/ProtectedRoute.jsx')
add_bullet('Auth Utilities: src/utils/auth.js')
add_bullet('Encryption Utilities: src/utils/encryption.js')

add_heading('Key Functions:', 2)
add_bullet('getUserRole() - Get current user\'s role from localStorage')
add_bullet('canPerformAction(action) - Check if user has specific permission')
add_bullet('isAdmin() - Check if user is admin')
add_bullet('getUserPermissions() - Get all permissions for current role')
add_bullet('usePermission(action) - React hook for permission checks')

doc.add_paragraph()
doc.add_page_break()

# ==============================================================================
# CHANGES FROM V3
# ==============================================================================
print("Adding Changes from V3...")

add_heading('Changes from Version 3.0', 1)

changes_headers = ['Change', 'v3.0', 'v4.0']
changes_rows = [
    ['File extensions', '.js', '.jsx for JSX files'],
    ['Document references', 'References v3 docs', 'Updated to v4'],
    ['Protected UI Component', 'ProtectedAction.js', 'ProtectedAction.jsx'],
    ['Route Protection', 'ProtectedRoute.js', 'ProtectedRoute.jsx'],
]
create_table(changes_headers, changes_rows)

doc.add_paragraph()
doc.add_paragraph()

# Add document end
end_text = doc.add_paragraph()
end_text.add_run('Document End').bold = True
end_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
output_path = r"c:\Users\rahul.ramagiri\Documents\git\pdcrApiDataMaskingUI\pdcrApiDataMaskingUI\docs\RBAC_Access_Matrix_v4.docx"
doc.save(output_path)

print(f"\nDocument saved successfully: {output_path}")
print("RBAC Access Matrix v4 DOCX created!")
