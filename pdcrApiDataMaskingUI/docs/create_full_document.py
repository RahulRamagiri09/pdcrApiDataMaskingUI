from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# Create a new document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Configure styles
styles = doc.styles
title_style = styles['Title']
title_font = title_style.font
title_font.name = 'Calibri'
title_font.size = Pt(28)
title_font.bold = True
title_font.color.rgb = RGBColor(0, 51, 102)

def add_heading(text, level=1):
    """Add a heading with specified level"""
    doc.add_heading(text, level)

def add_text(text):
    """Add a normal paragraph"""
    return doc.add_paragraph(text)

def add_diagram(text):
    """Add ASCII diagram in monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

# ==============================================================================
# TITLE PAGE
# ==============================================================================
print("Creating title page...")

title = doc.add_paragraph()
title_run = title.add_run('PII MASKING TOOL')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('ARCHITECTURE & DESIGN DOCUMENT')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(5):
    doc.add_paragraph()

info = doc.add_paragraph()
info.add_run('Version: 1.0\n\n').bold = True
info.add_run(f'Date: {datetime.datetime.now().strftime("%B %Y")}\n\n').bold = True
info.add_run('Document Type: Technical Architecture\n\n').bold = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================
print("Creating table of contents...")

doc.add_heading('TABLE OF CONTENTS', 0)
doc.add_paragraph()

toc = [
    '1. Executive Summary',
    '2. System Overview',
    '3. High-Level Architecture',
    '4. Component Architecture',
    '5. Data Flow Diagrams',
    '6. Technology Stack',
    '7. Database Design',
    '8. Security Architecture',
    '9. Deployment Architecture',
    '10. Integration Architecture'
]

for item in toc:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ==============================================================================
# SECTION 1: EXECUTIVE SUMMARY
# ==============================================================================
print("Adding Section 1: Executive Summary...")

add_heading('1. EXECUTIVE SUMMARY', 1)

add_heading('1.1 Project Overview', 2)
add_text('The PII Masking Tool is an enterprise web application designed to protect Personally Identifiable Information (PII) by masking sensitive data when copying from production databases to non-production environments (UAT, Development, Testing).')

add_heading('1.2 Purpose', 2)
add_text('This document provides a comprehensive technical architecture overview of the PII Masking Tool, including:')
add_text('- System architecture and component relationships\n- Data flow through the application\n- Technology choices and their justification\n- Security measures and deployment strategy')

add_heading('1.3 Key Features', 2)
add_text('- Database Connection Management: Configure and test connections to multiple SQL Server databases')
add_text('- Schema Discovery: Browse database schemas, tables, and column metadata')
add_text('- Intelligent Workflow Creation: Map source columns to appropriate PII masking techniques based on data types')
add_text('- Automated PII Masking: Execute data transformation with constraint validation')
add_text('- Execution Monitoring: Track workflow execution history with success/failure metrics')
add_text('- Role-Based Access Control: Admin and User roles with different permission levels')
add_text('- Audit Trail: Complete history of who executed what and when')

add_heading('1.4 Target Users', 2)
add_text('- Database Administrators: Manage database connections and schema exploration')
add_text('- Data Privacy Officers: Configure PII masking workflows')
add_text('- QA Engineers: Execute workflows to create test datasets')
add_text('- Development Teams: Access masked data for development purposes')

doc.add_page_break()

# ==============================================================================
# SECTION 2: SYSTEM OVERVIEW
# ==============================================================================
print("Adding Section 2: System Overview...")

add_heading('2. SYSTEM OVERVIEW', 1)

add_heading('2.1 Business Context', 2)
add_text('Organizations need to comply with data privacy regulations (GDPR, CCPA, HIPAA) while providing realistic test data to non-production environments. The PII Masking Tool automates the process of:')
add_text('1. Connecting to production (source) and non-production (target) databases')
add_text('2. Identifying columns containing PII')
add_text('3. Applying appropriate masking transformations')
add_text('4. Writing masked data to target databases')
add_text('5. Maintaining data integrity and referential constraints')

add_heading('2.2 System Capabilities', 2)

capabilities_diagram = """
SYSTEM CAPABILITIES

CONNECTION MANAGEMENT
├─ Register source and target database connections
├─ Test connection validity before saving
├─ Secure credential storage with encryption
└─ Support for multiple SQL Server instances

SCHEMA EXPLORATION
├─ Browse available schemas in connected databases
├─ View tables within each schema
├─ Inspect column metadata (name, type, nullability)
└─ Automatic data type detection

WORKFLOW CONFIGURATION
├─ Create named workflows with descriptions
├─ Select source and target connections
├─ Map source tables to target tables
├─ Configure column-level PII masking rules
├─ Smart filtering: Only show compatible masking options
└─ Save/edit/delete workflow configurations

PII MASKING EXECUTION
├─ Extract data from source database
├─ Apply masking transformations based on rules
├─ Validate data integrity constraints
├─ Load masked data into target database
├─ Transaction management (rollback on failure)
└─ Performance metrics (rows processed, duration)

MONITORING & REPORTING
├─ View workflow execution history
├─ Track success/failure status
├─ Display error messages for failed executions
├─ Show row counts (processed vs masked)
├─ Manual refresh of execution status
└─ User attribution (who executed what)

SECURITY & COMPLIANCE
├─ JWT-based authentication
├─ Role-based authorization (Admin/User)
├─ Encrypted database credentials
├─ Audit trail of all operations
└─ Secure password storage (hashed)
"""
add_diagram(capabilities_diagram)

doc.add_page_break()

# ==============================================================================
# SECTION 3: HIGH-LEVEL ARCHITECTURE
# ==============================================================================
print("Adding Section 3: High-Level Architecture...")

add_heading('3. HIGH-LEVEL ARCHITECTURE', 1)

add_heading('3.1 System Architecture Diagram', 2)

arch_diagram = """
3-TIER ARCHITECTURE

┌─────────────────────────────────────────────────────────┐
│                    CLIENT TIER                          │
│                 (Presentation Layer)                    │
│  ┌───────────────────────────────────────────────────┐ │
│  │        React Single Page Application              │ │
│  │                                                   │ │
│  │  Pages         Components      API Services      │ │
│  │  • Login       • Tables        • authAPI         │ │
│  │  • Dashboard   • Forms         • connectionsAPI  │ │
│  │  • Workflows   • Dialogs       • workflowsAPI    │ │
│  │  • Connections • Material-UI   • executionsAPI   │ │
│  └───────────────────────────────────────────────────┘ │
└────────────────────┬────────────────────────────────────┘
                     │
                     │ HTTPS/REST API (JSON + JWT)
                     ▼
┌─────────────────────────────────────────────────────────┐
│                APPLICATION TIER                         │
│               (Business Logic Layer)                    │
│  ┌───────────────────────────────────────────────────┐ │
│  │            FastAPI Backend Server                 │ │
│  │                                                   │ │
│  │  Middleware      Routers        Services         │ │
│  │  • CORS          • /auth        • DB Manager     │ │
│  │  • JWT Auth      • /connections • Workflow Exec  │ │
│  │  • Logging       • /workflows   • Masking Engine │ │
│  │  • Error Handler • /executions  • Validator      │ │
│  └───────────────────────────────────────────────────┘ │
└────────────────────┬────────────────────────────────────┘
                     │
                     │ SQL Queries (ODBC)
                     ▼
┌─────────────────────────────────────────────────────────┐
│                     DATA TIER                           │
│                (Persistence Layer)                      │
│  ┌───────────────────────────────────────────────────┐ │
│  │         Microsoft SQL Server Database             │ │
│  │                                                   │ │
│  │  Application DB    │    User Databases           │ │
│  │  • users           │    • Source (Production)    │ │
│  │  • connections     │    • Target (UAT/Dev/Test)  │ │
│  │  • workflows       │                             │ │
│  │  • mappings        │                             │ │
│  │  • executions      │                             │ │
│  └───────────────────────────────────────────────────┘ │
└─────────────────────────────────────────────────────────┘
"""
add_diagram(arch_diagram)

add_heading('3.2 Architecture Pattern', 2)
add_text('Pattern Type: 3-Tier Layered Architecture')
add_text('')
add_text('Why This Pattern:')
add_text('- Separation of Concerns: Each tier has a distinct responsibility')
add_text('- Independent Scaling: Scale presentation, logic, and data layers independently')
add_text('- Technology Flexibility: Replace or upgrade individual tiers without affecting others')
add_text('- Security: Multiple security checkpoints at each layer')
add_text('- Maintainability: Changes to one layer minimally impact other layers')

add_heading('3.3 Tier Responsibilities', 2)

add_heading('Client Tier (Presentation)', 3)
add_text('Responsibility: User interface and user experience')
add_text('')
add_text('Key Functions:')
add_text('- Render user interfaces with Material-UI components')
add_text('- Handle user interactions (clicks, form submissions)')
add_text('- Client-side routing between pages')
add_text('- Display data fetched from backend')
add_text('- Form validation and error messages')
add_text('- JWT token storage in browser localStorage')
add_text('')
add_text('Technology: React JavaScript framework running in web browser')
add_text('Communication: Makes HTTPS REST API calls to application tier')

add_heading('Application Tier (Business Logic)', 3)
add_text('Responsibility: Business rules, authentication, and workflow orchestration')
add_text('')
add_text('Key Functions:')
add_text('- Authenticate users and issue JWT tokens')
add_text('- Validate API requests and authorize access')
add_text('- Implement PII masking algorithms')
add_text('- Orchestrate ETL process (Extract, Transform, Load)')
add_text('- Enforce business rules and constraints')
add_text('- Manage database connections to user databases')
add_text('- Log operations for audit trail')
add_text('')
add_text('Technology: Python FastAPI framework running on application server')
add_text('Communication: Receives HTTPS requests from client tier, makes SQL queries to data tier')

add_heading('Data Tier (Persistence)', 3)
add_text('Responsibility: Data storage, integrity, and retrieval')
add_text('')
add_text('Key Functions:')
add_text('- Store application metadata (users, workflows, connections)')
add_text('- Store execution history and audit logs')
add_text('- Connect to user databases (source/target)')
add_text('- Enforce database constraints (NOT NULL, UNIQUE, FK)')
add_text('- Transaction management (ACID compliance)')
add_text('- Query optimization and indexing')
add_text('')
add_text('Technology: Microsoft SQL Server database')
add_text('Communication: Receives SQL queries from application tier')

doc.add_page_break()

# Save document
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print(f"Document saved successfully: {output_path}")
print("Document created with all sections!")
