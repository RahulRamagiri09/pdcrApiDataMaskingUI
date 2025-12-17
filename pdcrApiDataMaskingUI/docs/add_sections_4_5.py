from docx import Document
from docx.shared import Pt

# Open existing document
doc = Document(r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx")

def add_heading(text, level=1):
    doc.add_heading(text, level)

def add_text(text):
    return doc.add_paragraph(text)

def add_diagram(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

print("Adding Section 4: Component Architecture...")

# ==============================================================================
# SECTION 4: COMPONENT ARCHITECTURE
# ==============================================================================

add_heading('4. COMPONENT ARCHITECTURE', 1)

add_heading('4.1 Frontend Component Structure', 2)

frontend_diagram = """
FRONTEND ARCHITECTURE

                    App.js (Root Component)
                            |
            +---------------+---------------+
            |               |               |
      Authentication    Routing         Layout
       Management    (React Router)     Wrapper
            |               |               |
            |       +-------+-------+       |
            |       |       |       |       |
      Public Route  |   Protected  |   Shared UI
            |       |    Routes     |       |
            v       v       v       v       v
        Login   Dashboard Workflows Executions AppBar
         Page     Page      Page      Page    Drawer

WORKFLOWS PAGE HIERARCHY:
    WorkflowsPage.js
            |
            +-- CreateWorkflowPage.js (Multi-Step Wizard)
            |       |
            |       +-- Step 1: Basic Information
            |       +-- Step 2: Source Selection
            |       +-- Step 3: Column Mapping (Smart Filtering)
            |       +-- Step 4: Target Configuration
            |
            +-- WorkflowDetailPage.js
                    |
                    +-- Workflow Header
                    +-- Configuration Details
                    +-- Column Mapping Table
                    +-- Execution History
                    +-- Manual Refresh Button
"""
add_diagram(frontend_diagram)

add_heading('4.2 Frontend Component Descriptions', 2)

add_heading('Login Page', 3)
add_text('Entry point for unauthenticated users. Handles credential submission via authAPI.login(), stores JWT token in localStorage upon successful authentication, redirects to dashboard after login, and displays error messages for invalid credentials.')

add_heading('Dashboard Page', 3)
add_text('Landing page after authentication. Displays aggregate statistics (total connections, workflows, executions), shows recent activity feed, provides quick navigation to main features, and fetches real-time data on component mount.')

add_heading('ConnectionsPage', 3)
add_text('Manages database connections (both source and target). Provides CRUD operations for connections, connection testing functionality before saving, displays connection details (server, database, type), and includes filters by connection type (source/target).')

add_heading('CreateWorkflowPage (Core Component)', 3)
add_text('Multi-step wizard for workflow creation. Handles both create and edit modes, implements smart PII attribute filtering by data type, manages complex state (connections, schemas, tables, columns, mappings), validates user input at each step, and makes API calls for fetching metadata and saving workflow.')

add_text('Step 1 - Basic Information: User fills in workflow name, description, and selects source/target connections.')
add_text('Step 2 - Source Selection: Cascading dropdowns for schema, table, and column selection with data type display.')
add_text('Step 3 - Column Mapping (CRITICAL): Smart filtering shows only compatible PII attributes based on column data types. For example, integer columns show only numeric masking options, varchar columns show only string masking options.')
add_text('Step 4 - Target Configuration: Select target schema and table, review all mappings before submission.')

add_heading('WorkflowDetailPage', 3)
add_text('Displays comprehensive workflow information including configuration details, column mapping rules, and execution history. Features a back button with circular hover effect, run workflow button for on-demand execution, and manual refresh button (no automatic polling) to update execution history.')

add_heading('4.3 Backend Component Structure', 2)

backend_diagram = """
BACKEND ARCHITECTURE

            FastAPI Application
                    |
        +-----------+-----------+
        |           |           |
   Middleware    Routers    Business
     Layer        Layer     Services
        |           |           |
        |           |           |
    +---+---+   +---+---+   +---+---+
    |       |   |       |   |       |
  CORS   JWT  Auth  Workflows DB   Masking
         Verify     Connections  Manager Engine
         Error      Executions
         Handler    Schemas
                              Validator

DATA ACCESS LAYER
        |
        +-- Application DB (Metadata)
        +-- Source DB (Read PII)
        +-- Target DB (Write Masked)
"""
add_diagram(backend_diagram)

add_heading('4.4 Backend Component Descriptions', 2)

add_heading('Middleware Layer', 3)
add_text('Intercepts all incoming requests before reaching API routers:')
add_text('- CORS Middleware: Validates request origin, allows only frontend URL')
add_text('- JWT Authentication Middleware: Decodes and validates JWT token, extracts user information')
add_text('- Error Handler Middleware: Catches exceptions and formats error responses consistently')
add_text('- Logging Middleware: Logs all API requests with timestamp, user, and endpoint')

add_heading('API Routers', 3)
add_text('Handle HTTP requests and route to appropriate business logic:')
add_text('- Auth Router (/api/auth): Login, user profile retrieval')
add_text('- Connections Router (/api/connections): Connection CRUD, testing, schema/table discovery')
add_text('- Workflows Router (/api/workflows): Workflow CRUD, execution triggering')
add_text('- Executions Router (/api/executions): Execution history and status monitoring')
add_text('- PII Attributes Router (/api/pii-attributes): Return categorized masking techniques')

add_heading('Business Services Layer', 3)
add_text('Database Manager Service: Establishes connections to SQL Server databases, manages connection pooling, executes parameterized queries, and handles connection errors with retries.')
add_text('')
add_text('Workflow Executor Service: Orchestrates the ETL process - loads configuration, connects to databases, extracts source data, applies masking transformations, validates constraints, loads to target, and logs results.')
add_text('')
add_text('Masking Engine Service: Implements PII masking algorithms including string masking (fake names, emails, phones), numeric masking (random numbers, ranges), date/datetime masking (shifting), and boolean masking.')
add_text('')
add_text('Validator Service: Validates data integrity before insertion including NOT NULL checks, data type compatibility, UNIQUE constraints, and referential integrity.')

doc.add_page_break()

# ==============================================================================
# SECTION 5: DATA FLOW DIAGRAMS
# ==============================================================================

print("Adding Section 5: Data Flow Diagrams...")

add_heading('5. DATA FLOW DIAGRAMS', 1)

add_heading('5.1 User Authentication Flow', 2)

auth_flow = """
USER AUTHENTICATION FLOW

User                Frontend            Backend           Database
 |                     |                   |                  |
 | 1. Enter           |                   |                  |
 |    Credentials     |                   |                  |
 +------------------->|                   |                  |
                      | 2. POST /api/auth/|                  |
                      |    login          |                  |
                      +------------------>|                  |
                                          | 3. Query users   |
                                          |    table         |
                                          +----------------->|
                                          | 4. User record   |
                                          |<-----------------+
                                          | 5. Validate      |
                                          |    password hash |
                                          |                  |
                                          | IF VALID:        |
                                          | 6. Generate JWT  |
                                          |    Token with:   |
                                          |    - user_id     |
                                          |    - username    |
                                          |    - role        |
                                          |    - expiration  |
                      | 7. Return Token   |                  |
                      |<------------------+                  |
                      | 8. Store in       |                  |
                      |    localStorage   |                  |
 | 9. Redirect to     |                   |                  |
 |    Dashboard       |                   |                  |
 |<-------------------+                   |                  |

SUBSEQUENT REQUESTS:
                      | 10. API Request   |                  |
                      |     Authorization:|                  |
                      |     Bearer <token>|                  |
                      +------------------>|                  |
                                          | 11. Verify JWT   |
                                          |     signature    |
                                          |     & expiration |
                                          | 12. Extract      |
                                          |     user context |
                      | 13. Authorized    |                  |
                      |     Response      |                  |
                      |<------------------+                  |
"""
add_diagram(auth_flow)

add_heading('5.2 Authentication Flow Description', 2)
add_text('The authentication flow implements JWT-based stateless authentication:')
add_text('')
add_text('Steps 1-2: User enters credentials in login form. Frontend sends POST request to /api/auth/login.')
add_text('')
add_text('Steps 3-5: Backend queries users table, retrieves user record, and compares password hash using bcrypt.')
add_text('')
add_text('Steps 6-9: If valid, backend generates JWT token with user claims and 24-hour expiration. Frontend stores token in localStorage and redirects to dashboard.')
add_text('')
add_text('Steps 10-13: All subsequent API requests include JWT token in Authorization header. Backend validates token on every request and extracts user context for authorization.')
add_text('')
add_text('Security Benefits: Stateless authentication (no server-side sessions), tamper-proof signature, self-contained token, automatic expiration, and independent request validation.')

add_heading('5.3 Workflow Creation Flow', 2)

workflow_flow = """
WORKFLOW CREATION FLOW

User          Frontend         Backend        App DB      Source DB
 |                |                |              |            |
 | Navigate to    |                |              |            |
 | Create Page    |                |              |            |
 +--------------->|                |              |            |
                  | Load Initial   |              |            |
                  | Data           |              |            |
                  |                |              |            |
                  | GET /connections|             |            |
                  +--------------->|              |            |
                  |                | Query        |            |
                  |                | connections  |            |
                  |                +------------->|            |
                  | Connection     | List         |            |
                  | List           |<-------------+            |
                  |<---------------+              |            |
                  |                |              |            |
                  | GET /pii-      |              |            |
                  | attributes     |              |            |
                  +--------------->|              |            |
                  | Categorized    |              |            |
                  | PII List       |              |            |
                  |<---------------+              |            |
                  |                |              |            |
 | STEP 1:        |                |              |            |
 | Fill Basic Info|                |              |            |
 +--------------->|                |              |            |
                  | Validate &     |              |            |
                  | Move to Step 2 |              |            |
                  |                |              |            |
 | STEP 2:        |                |              |            |
 | Select Schema  |                |              |            |
 +--------------->|                |              |            |
                  | GET /schemas   |              |            |
                  +--------------->|              |            |
                  |                | Connect to   |            |
                  |                | source DB    |            |
                  |                +------------------------>  |
                  |                | Schema list  |            |
                  | Schema List    |<-------------------------+
                  |<---------------+              |            |
                  |                |              |            |
 | Select Table   |                |              |            |
 +--------------->|                |              |            |
                  | GET /tables    |              |            |
                  +--------------->+------------->+----------->|
                  | Table List     |              |            |
                  |<---------------+<-------------+<-----------+
                  |                |              |            |
 | STEP 3:        |                |              |            |
 | Column Mapping |                |              |            |
 | with Smart     |                |              |            |
 | Filtering      |                |              |            |
 +--------------->|                |              |            |
                  | For each column|              |            |
                  | filter PII by  |              |            |
                  | data type      |              |            |
                  |                |              |            |
 | Select PII for |                |              |            |
 | each column    |                |              |            |
 +--------------->|                |              |            |
                  |                |              |            |
 | STEP 4:        |                |              |            |
 | Target Config  |                |              |            |
 +--------------->|                |              |            |
                  |                |              |            |
 | Submit         |                |              |            |
 +--------------->|                |              |            |
                  | POST /workflows|              |            |
                  +--------------->|              |            |
                  |                | INSERT       |            |
                  |                | workflows &  |            |
                  |                | mappings     |            |
                  |                +------------->|            |
                  | Success        | Workflow ID  |            |
                  |<---------------+<-------------+            |
 | Navigate to    |                |              |            |
 | Workflows List |                |              |            |
 |<---------------+                |              |            |
"""
add_diagram(workflow_flow)

add_heading('5.4 Workflow Creation Flow Description', 2)
add_text('The workflow creation is a multi-step wizard process:')
add_text('')
add_text('Initial Load: Frontend fetches connections and categorized PII attributes from backend.')
add_text('')
add_text('Step 1 - Basic Information: User provides workflow name, description, and selects source/target connections.')
add_text('')
add_text('Step 2 - Source Selection: Cascading API calls load schemas, then tables, then column metadata. Each selection triggers the next level of data fetching.')
add_text('')
add_text('Step 3 - Column Mapping (Smart Filtering): For each selected column, frontend identifies data type, maps it to PII category (string/numeric/date/datetime/boolean), and filters PII dropdown to show only compatible masking options. This prevents invalid mappings.')
add_text('')
add_text('Step 4 - Target Configuration: User selects destination schema and table, reviews all mappings.')
add_text('')
add_text('Submission: Frontend sends complete workflow configuration to backend. Backend creates workflow and mapping records in a transaction. Success confirmation navigates user to workflows list.')

doc.add_page_break()

add_heading('5.5 Workflow Execution Flow', 2)

exec_flow = """
WORKFLOW EXECUTION FLOW

User   Frontend   Backend     App DB    Source DB  Target DB
 |        |          |            |          |          |
 | Click  |          |            |          |          |
 | Run    |          |            |          |          |
 +------->|          |            |          |          |
          | POST     |            |          |          |
          | /execute |            |          |          |
          +--------->|            |          |          |
                     |            |          |          |
                     | PHASE 1: INITIALIZATION           |
                     +----------->|          |          |
                     | Load       |          |          |
                     | workflow   |          |          |
                     | config     |          |          |
                     | Create exec|          |          |
                     | record     |          |          |
          | exec_id  |<-----------+          |          |
          |<---------+            |          |          |
 | Show   |          |            |          |          |
 | Started|          |            |          |          |
 |<-------+          |            |          |          |
          |          |            |          |          |
          |          | PHASE 2: DATA EXTRACTION          |
          |          +----------->+--------->|          |
          |          | SELECT *   |          |          |
          |          | FROM source|          |          |
          |          | Result set |          |          |
          |          |<-----------+<---------+          |
          |          |            |          |          |
          |          | PHASE 3: TRANSFORMATION           |
          |          | For each row:        |          |
          |          | - Apply PII masking  |          |
          |          | - first_name: John   |          |
          |          |   -> Michael         |          |
          |          | - age: 30 -> 42      |          |
          |          | - birth_date shifted |          |
          |          |            |          |          |
          |          | PHASE 4: VALIDATION  |          |
          |          | - NOT NULL checks    |          |
          |          | - Data type checks   |          |
          |          | - UNIQUE checks      |          |
          |          |            |          |          |
          |          | PHASE 5: DATA LOADING            |
          |          +----------->+--------->+--------->|
          |          | INSERT     |          |          |
          |          | masked data|          |          |
          |          | (batches)  |          |          |
          |          | COMMIT     |          |          |
          |          |<-----------+<---------+<---------+
          |          |            |          |          |
          |          | PHASE 6: COMPLETION  |          |
          |          +----------->|          |          |
          |          | Update exec|          |          |
          |          | status:    |          |          |
          |          | success    |          |          |
          |          |<-----------+          |          |
          |          |            |          |          |
 | Click  |          |            |          |          |
 | Refresh|          |            |          |          |
 +------->|          |            |          |          |
          | GET      |            |          |          |
          | /executions           |          |          |
          +--------->+----------->|          |          |
          | Updated  | Query      |          |          |
          | history  | executions |          |          |
          |<---------+<-----------+          |          |
 | Display|          |            |          |          |
 | Success|          |            |          |          |
 |<-------+          |            |          |          |
"""
add_diagram(exec_flow)

add_heading('5.6 Workflow Execution Flow Description', 2)
add_text('The workflow execution orchestrates the complete ETL (Extract, Transform, Load) process:')
add_text('')
add_text('Phase 1 - Initialization: Backend loads workflow configuration, retrieves connection credentials, creates execution record with status "running", and returns execution_id to frontend immediately.')
add_text('')
add_text('Phase 2 - Data Extraction: Backend connects to source database, executes SELECT query, fetches all rows, and counts rows_processed.')
add_text('')
add_text('Phase 3 - Data Transformation: For each row, backend applies PII masking based on configured rules. Examples: first_name "John" becomes random name "Michael", age 30 becomes random number 42, birth_date shifted by random days.')
add_text('')
add_text('Phase 4 - Constraint Validation: Backend validates NOT NULL constraints, data type compatibility, UNIQUE constraints. If any validation fails, execution stops and status marked as "failed".')
add_text('')
add_text('Phase 5 - Data Loading: Backend connects to target database, begins transaction, inserts masked data in batches (100 rows per batch), commits transaction if successful, or rolls back on error.')
add_text('')
add_text('Phase 6 - Completion: Backend updates execution record with final status, row counts, and completion time. User clicks manual refresh to see updated history.')
add_text('')
add_text('Error Handling: If any phase fails, backend logs error, updates execution status to "failed", rolls back database transaction, and no partial data is written (atomic operation).')

doc.add_page_break()

add_heading('5.7 Smart PII Filtering Flow', 2)

filter_flow = """
SMART PII ATTRIBUTE FILTERING

User                Frontend Logic           PII Categories
 |                         |                        |
 | Load Create             |                        |
 | Workflow Page           |                        |
 +------------------------>| Fetch categorized      |
                           | PII attributes         |
                           +----------------------->|
                           | Return:                |
                           | {                      |
                           |   string: [first_name, |
                           |            last_name]  |
                           |   numeric: [random_    |
                           |             number]    |
                           |   date: [date_shift]   |
                           |   datetime: [datetime_ |
                           |              shift]    |
                           |   boolean: [random_    |
                           |             boolean]   |
                           | }                      |
                           |<-----------------------+
                           | Store in state         |
                           |                        |
 | User selects            |                        |
 | table columns           |                        |
 +------------------------>|                        |
                           | Backend returns        |
                           | column metadata:       |
                           | {                      |
                           |   name: "age",         |
                           |   data_type: "int"     |
                           | }                      |
                           |                        |
 | SCENARIO: User checks   |                        |
 | column "age" (int)      |                        |
 +------------------------>|                        |
                           | Smart Filter:          |
                           | 1. Get column info     |
                           | 2. Map data type:      |
                           |    "int" -> "numeric"  |
                           | 3. Filter PII:         |
                           |    Return numeric      |
                           |    attributes only     |
                           |                        |
 | Show PII dropdown:      |                        |
 | [v] Select PII:         |                        |
 |   - random_number       |                        |
 |   - range_number        |                        |
 |   - fixed_number        |                        |
 | NOT showing:            |                        |
 |   X first_name          |                        |
 |   X date_shift          |                        |
 |<------------------------+                        |

DATA TYPE MAPPING:
varchar/text      -> string    -> first_name, email_mask
int/numeric       -> numeric   -> random_number
date              -> date      -> date_shift
datetime          -> datetime  -> datetime_shift
bit/boolean       -> boolean   -> random_boolean
"""
add_diagram(filter_flow)

add_heading('5.8 Smart PII Filtering Description', 2)
add_text('Smart PII filtering is the core innovation that prevents data type mismatches:')
add_text('')
add_text('Problem: Without filtering, users could create invalid mappings like applying "first_name" (string) to an integer column.')
add_text('')
add_text('Solution: Frontend automatically filters PII options based on column data type.')
add_text('')
add_text('Process:')
add_text('1. Column Selection: User checks checkbox for a column (e.g., "age")')
add_text('2. Metadata Lookup: Frontend retrieves data type from column metadata (data_type: "int")')
add_text('3. Category Mapping: Frontend maps SQL type to category using rules: int->numeric, varchar->string, date->date, datetime->datetime, bit->boolean')
add_text('4. PII Filtering: Frontend returns only PII attributes matching the category')
add_text('5. UI Rendering: Dropdown shows only compatible options')
add_text('')
add_text('Benefits:')
add_text('- Prevents configuration errors at design time, not runtime')
add_text('- Improves user experience by reducing irrelevant choices')
add_text('- Ensures type compatibility before execution')
add_text('- Maintains data integrity')

# Save progress
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print(f"Sections 4-5 added successfully!")
print(f"Document saved: {output_path}")
