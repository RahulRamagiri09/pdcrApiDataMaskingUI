from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open the existing document
doc = Document(r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx")

def add_section(title, level=1):
    """Add a heading section"""
    doc.add_heading(title, level)

def add_paragraph_text(text):
    """Add a normal paragraph"""
    p = doc.add_paragraph(text)
    return p

def add_diagram(diagram_text):
    """Add ASCII diagram in monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(diagram_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

# Add all content sections
print("Adding Section 1: Executive Summary...")
add_section('1. EXECUTIVE SUMMARY', 1)

add_section('1.1 Project Overview', 2)
add_paragraph_text('The PII Masking Tool is an enterprise web application designed to protect Personally Identifiable Information (PII) by masking sensitive data when copying from production databases to non-production environments (UAT, Development, Testing).')

add_section('1.2 Purpose', 2)
add_paragraph_text('This document provides a comprehensive technical architecture overview of the PII Masking Tool, including system architecture and component relationships, data flow through the application, technology choices and their justification, and security measures and deployment strategy.')

add_section('1.3 Key Features', 2)
features = """• Database Connection Management: Configure and test connections to multiple SQL Server databases
• Schema Discovery: Browse database schemas, tables, and column metadata
• Intelligent Workflow Creation: Map source columns to appropriate PII masking techniques based on data types
• Automated PII Masking: Execute data transformation with constraint validation
• Execution Monitoring: Track workflow execution history with success/failure metrics
• Role-Based Access Control: Admin and User roles with different permission levels
• Audit Trail: Complete history of who executed what and when"""
add_paragraph_text(features)

add_section('1.4 Target Users', 2)
users = """• Database Administrators: Manage database connections and schema exploration
• Data Privacy Officers: Configure PII masking workflows
• QA Engineers: Execute workflows to create test datasets
• Development Teams: Access masked data for development purposes"""
add_paragraph_text(users)

print("Adding Section 2: System Overview...")
add_section('2. SYSTEM OVERVIEW', 1)

add_section('2.1 Business Context', 2)
add_paragraph_text('Organizations need to comply with data privacy regulations (GDPR, CCPA, HIPAA) while providing realistic test data to non-production environments. The PII Masking Tool automates the process of connecting to production and non-production databases, identifying columns containing PII, applying appropriate masking transformations, writing masked data to target databases, and maintaining data integrity and referential constraints.')

add_section('2.2 System Capabilities', 2)
capabilities = """CONNECTION MANAGEMENT:
• Register source and target database connections
• Test connection validity before saving
• Secure credential storage with encryption
• Support for multiple SQL Server instances

SCHEMA EXPLORATION:
• Browse available schemas in connected databases
• View tables within each schema
• Inspect column metadata (name, type, nullability)
• Automatic data type detection

WORKFLOW CONFIGURATION:
• Create named workflows with descriptions
• Select source and target connections
• Map source tables to target tables
• Configure column-level PII masking rules
• Smart filtering: Only show compatible masking options
• Save/edit/delete workflow configurations

PII MASKING EXECUTION:
• Extract data from source database
• Apply masking transformations based on rules
• Validate data integrity constraints
• Load masked data into target database
• Transaction management (rollback on failure)
• Performance metrics (rows processed, duration)

MONITORING & REPORTING:
• View workflow execution history
• Track success/failure status
• Display error messages for failed executions
• Show row counts (processed vs masked)
• Manual refresh of execution status
• User attribution (who executed what)

SECURITY & COMPLIANCE:
• JWT-based authentication
• Role-based authorization (Admin/User)
• Encrypted database credentials
• Audit trail of all operations
• Secure password storage (hashed)"""
add_paragraph_text(capabilities)

doc.add_page_break()

print("Adding Section 3: High-Level Architecture...")
add_section('3. HIGH-LEVEL ARCHITECTURE', 1)

add_section('3.1 System Architecture Diagram', 2)
arch_diagram = """
                         CLIENT TIER
                      (Presentation Layer)

               React Single Page Application

        User Interface  |  Components  |  HTTP Client
            Pages       |  Material-UI |  (Axios)
        • Login         |  • Tables    |  • JWT Storage
        • Dashboard     |  • Forms     |  • Auto Auth
        • Workflows     |  • Dialogs   |  • Error Handle
        • Connections   |              |
                              |
                              | HTTPS/REST API
                              | (JSON + JWT)
                              ▼
                      APPLICATION TIER
                   (Business Logic Layer)

                    FastAPI Backend Server

        Middleware    |    API Routers    |  Business Logic
        • CORS        |    • /auth        |  • DB Management
        • JWT Verify  |    • /connections |  • Workflow Exec
        • Error       |    • /workflows   |  • PII Masking
        • Logging     |    • /executions  |  • Validation
                              |
                              | SQL Queries
                              ▼
                          DATA TIER
                      (Persistence Layer)

                  Microsoft SQL Server Database

        Application DB        |        User Databases
        (Metadata Storage)    |        (Source/Target Data)
        • users               |        • Production DB
        • connections         |        • UAT DB
        • workflows           |        • Dev DB
        • mappings            |        • Test DB
        • executions          |
"""
add_diagram(arch_diagram)

add_section('3.2 Architecture Pattern', 2)
add_paragraph_text('Pattern Type: 3-Tier Layered Architecture')
add_paragraph_text('Why This Pattern: Separation of concerns ensures each tier has a distinct responsibility, independent scaling allows scaling presentation, logic, and data layers independently, technology flexibility enables replacing or upgrading individual tiers without affecting others, multiple security checkpoints exist at each layer, and changes to one layer minimally impact other layers.')

add_section('3.3 Tier Responsibilities', 2)

add_section('Client Tier (Presentation)', 3)
add_paragraph_text('Responsibility: User interface and user experience. Key Functions: Render user interfaces with Material-UI components, handle user interactions, client-side routing between pages, display data fetched from backend, form validation and error messages, JWT token storage in browser localStorage. Technology: React JavaScript framework running in web browser. Communication: Makes HTTPS REST API calls to application tier.')

add_section('Application Tier (Business Logic)', 3)
add_paragraph_text('Responsibility: Business rules, authentication, and workflow orchestration. Key Functions: Authenticate users and issue JWT tokens, validate API requests and authorize access, implement PII masking algorithms, orchestrate ETL process (Extract, Transform, Load), enforce business rules and constraints, manage database connections to user databases, log operations for audit trail. Technology: Python FastAPI framework running on application server. Communication: Receives HTTPS requests from client tier and makes SQL queries to data tier.')

add_section('Data Tier (Persistence)', 3)
add_paragraph_text('Responsibility: Data storage, integrity, and retrieval. Key Functions: Store application metadata (users, workflows, connections), store execution history and audit logs, connect to user databases (source/target), enforce database constraints (NOT NULL, UNIQUE, FK), transaction management (ACID compliance), query optimization and indexing. Technology: Microsoft SQL Server database. Communication: Receives SQL queries from application tier.')

doc.add_page_break()

print("Adding remaining sections...")

# Save the updated document
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print("Document updated with content sections")
print(f"File saved: {output_path}")
