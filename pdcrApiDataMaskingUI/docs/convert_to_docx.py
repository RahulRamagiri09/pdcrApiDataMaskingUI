"""
Convert PII_Masking_Tool_How_It_Works.md to DOCX format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_borders(table):
    """Add borders to table"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX"""

    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Modify Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    # Create Code style for diagrams
    if 'Code Block' not in [s.name for s in styles]:
        code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)

    # Split content into lines
    lines = content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Handle horizontal rules
        if line.strip() == '---':
            # Add a thin horizontal line
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        # Handle headings
        if line.startswith('#'):
            heading_match = re.match(r'^(#+)\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)

                if level == 1:
                    heading = doc.add_heading(text, level=0)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif level <= 4:
                    doc.add_heading(text, level=level)
                else:
                    para = doc.add_paragraph(text)
                    para.runs[0].bold = True

                i += 1
                continue

        # Handle code blocks (ASCII diagrams)
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```

            # Add code block with monospace font
            if code_lines:
                for code_line in code_lines:
                    para = doc.add_paragraph(code_line)
                    para.style = 'Code Block'
                    # Ensure monospace font
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
            continue

        # Handle tables
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not lines[i].strip().replace('|', '').replace('-', '').replace(':', '').strip() == '':
                    table_lines.append(lines[i])
                i += 1

            if table_lines:
                # Parse table
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip().split('|')[1:-1]]
                    if cells:
                        rows.append(cells)

                if rows:
                    # Create table
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    add_table_borders(table)

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.cell(row_idx, col_idx)
                                # Remove markdown bold markers
                                clean_text = cell_text.replace('**', '')
                                cell.text = clean_text

                                # Style header row
                                if row_idx == 0:
                                    set_cell_shading(cell, 'D9E2F3')
                                    cell.paragraphs[0].runs[0].bold = True

                    # Add some spacing after table
                    doc.add_paragraph()
            continue

        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown bold markers
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            para = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Handle numbered lists
        numbered_match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if numbered_match:
            text = numbered_match.group(2)
            # Remove markdown bold markers
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            para = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Handle regular paragraphs
        text = line.strip()
        if text:
            # Remove markdown bold markers but preserve the text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            para = doc.add_paragraph(text)

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"Successfully created: {docx_file}")

if __name__ == '__main__':
    # Get the directory of this script
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Define input and output files
    md_file = os.path.join(script_dir, 'PII_Masking_Tool_How_It_Works.md')
    docx_file = os.path.join(script_dir, 'PII_Masking_Tool_How_It_Works.docx')

    # Convert
    convert_md_to_docx(md_file, docx_file)
