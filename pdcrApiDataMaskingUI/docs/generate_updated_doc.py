from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# Create a new document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Configure styles
styles = doc.styles
title_style = styles['Title']
title_font = title_style.font
title_font.name = 'Calibri'
title_font.size = Pt(28)
title_font.bold = True
title_font.color.rgb = RGBColor(0, 51, 102)

def add_heading(text, level=1):
    """Add a heading with specified level"""
    doc.add_heading(text, level)

def add_text(text):
    """Add a normal paragraph"""
    return doc.add_paragraph(text)

def add_bullet(text):
    """Add a bullet point"""
    return doc.add_paragraph(text, style='List Bullet')

def add_diagram(text):
    """Add ASCII diagram in monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

# ==============================================================================
# TITLE PAGE
# ==============================================================================
print("Creating title page...")

title = doc.add_paragraph()
title_run = title.add_run('PII MASKING TOOL')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('ARCHITECTURE & DESIGN DOCUMENT')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(5):
    doc.add_paragraph()

info = doc.add_paragraph()
info.add_run('Version: 1.0\n\n').bold = True
info.add_run('Date: January 2025\n\n').bold = True
info.add_run('Document Type: Technical Architecture\n\n').bold = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================
print("Creating table of contents...")

doc.add_heading('TABLE OF CONTENTS', 0)
doc.add_paragraph()

toc = [
    '1. Executive Summary',
    '2. System Overview',
    '3. High-Level Architecture',
    '4. Component Architecture',
    '5. Data Flow Diagrams',
    '6. Security Architecture',
    '7. Deployment Architecture',
    '8. Integration Architecture'
]

for item in toc:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ==============================================================================
# SECTION 1: EXECUTIVE SUMMARY
# ==============================================================================
print("Adding Section 1: Executive Summary...")

add_heading('1. EXECUTIVE SUMMARY', 1)

add_heading('1.1 Project Overview', 2)
add_text('The PII Masking Tool is an enterprise web application designed to protect Personally Identifiable Information (PII) by performing in-place masking of sensitive data within a single database. The tool masks PII columns directly in the same table, schema, and database where the data resides.')

add_heading('1.2 Purpose', 2)
add_text('This document provides a comprehensive technical architecture overview of the PII Masking Tool, including:')
add_bullet('System architecture and component relationships')
add_bullet('Data flow through the application')
add_bullet('Technology choices and their justification')
add_bullet('Security measures and deployment strategy')

add_heading('1.3 Key Features', 2)
add_bullet('Server Connection Management: Configure and test connections to SQL Server databases (Azure SQL, PostgreSQL, Oracle, SQL Server)')
add_bullet('Schema Discovery: Browse database schemas, tables, and column metadata')
add_bullet('Intelligent Workflow Creation: Select tables and map columns to appropriate PII masking techniques based on data types')
add_bullet('In-Place PII Masking: Execute data transformation directly on the same table with constraint validation')
add_bullet('Preview Masking: Preview how masking will affect sample records before execution')
add_bullet('Constraint Checking: Validate primary keys, foreign keys, unique constraints, check constraints, triggers, and indexes before masking')
add_bullet('Execution Monitoring: Track workflow execution history with success/failure metrics')
add_bullet('Role-Based Access Control: Admin and User roles with different permission levels')
add_bullet('Audit Trail: Complete history of who executed what and when')

add_heading('1.4 Target Users', 2)
add_bullet('Database Administrators: Manage database connections and schema exploration')
add_bullet('Data Privacy Officers: Configure PII masking workflows')
add_bullet('Data Engineers: Execute workflows to mask sensitive data in-place')
add_bullet('Compliance Teams: Monitor masking operations and audit trails')

doc.add_page_break()

# ==============================================================================
# SECTION 2: SYSTEM OVERVIEW
# ==============================================================================
print("Adding Section 2: System Overview...")

add_heading('2. SYSTEM OVERVIEW', 1)

add_heading('2.1 Business Context', 2)
add_text('Organizations need to comply with data privacy regulations (GDPR, CCPA, HIPAA) by protecting sensitive PII data within their databases. The PII Masking Tool automates the process of:')
add_bullet('Connecting to database servers (Azure SQL, PostgreSQL, Oracle, SQL Server)')
add_bullet('Discovering schemas, tables, and columns containing PII')
add_bullet('Configuring column-level masking rules based on data types')
add_bullet('Validating database constraints before masking (PKs, FKs, unique, check, triggers, indexes)')
add_bullet('Executing in-place masking transformations directly on the same table')
add_bullet('Maintaining data integrity and referential constraints')

add_heading('2.2 System Capabilities', 2)

capabilities_diagram = """
SYSTEM CAPABILITIES

SERVER CONNECTION MANAGEMENT
|- Register database server connections
|- Support multiple database types (Azure SQL, PostgreSQL, Oracle, SQL Server)
|- Test connection validity before saving
|- Secure credential storage with encryption
|- Connection name validation (cannot match server name)

SCHEMA EXPLORATION
|- Browse available schemas in connected databases
|- View tables within each schema
|- Inspect column metadata (name, type, nullability)
|- Automatic data type detection

WORKFLOW CONFIGURATION
|- Create named workflows with descriptions
|- Select single connection for in-place masking
|- Select schema and table for masking
|- Configure column-level PII masking rules
|- Smart filtering: Only show compatible masking options
|- Save/edit/delete workflow configurations

CONSTRAINT CHECKING
|- Check primary key constraints
|- Check foreign key constraints
|- Check unique constraints
|- Check check constraints
|- Check triggers on table
|- Check indexes on table

PREVIEW MASKING
|- Preview masking on sample records
|- Configure number of preview records
|- View original vs masked data side-by-side
|- Validate masking before execution

IN-PLACE PII MASKING EXECUTION
|- Execute masking directly on same table
|- Apply masking transformations based on rules
|- Transaction management (rollback on failure)
|- Performance metrics (rows processed, duration)
|- Detailed execution logs

MONITORING & REPORTING
|- View workflow execution history
|- Track success/failure status
|- Display error messages for failed executions
|- Show row counts (processed vs masked)
|- View execution logs
|- User attribution (who executed what)

SECURITY & COMPLIANCE
|- JWT-based authentication
|- Role-based authorization (Admin/User)
|- Encrypted database credentials
|- Audit trail of all operations
|- Secure password storage (hashed)
"""
add_diagram(capabilities_diagram)

add_heading('2.3 User Journey', 2)

user_journey = """
User Login
    |
Dashboard (Overview of connections, workflows, recent executions)
    |
    |-> Manage Server Connections
    |       |
    |   Add Connection (Connection Details -> Test Connection -> Save)
    |       |
    |   View/Delete existing connections
    |
    |-> Create Workflow
    |       |
    |   Step 1: Basic Info (name, description)
    |       |
    |   Step 2: Select connection, schema, table
    |       |
    |   Step 3: Configure column mappings with PII attributes
    |       |
    |   Save Workflow
    |
    |-> View Workflow Details
    |       |
    |   Overview Tab: Workflow configuration
    |       |
    |   Execution History Tab: Past executions
    |       |
    |   Preview Masking Tab:
    |       |-> Column Mapping: View configured mappings
    |       |-> Constraint Checks: Validate PKs, FKs, etc.
    |       |-> Preview Masking: See sample masked data
    |
    |-> Execute Workflow
    |       |
    |   Trigger in-place masking
    |       |
    |   Monitor progress
    |       |
    |   View execution logs and results
    |
    |-> View Execution History
            |
        Audit trail and reporting
"""
add_diagram(user_journey)

doc.add_page_break()

# ==============================================================================
# SECTION 3: HIGH-LEVEL ARCHITECTURE
# ==============================================================================
print("Adding Section 3: High-Level Architecture...")

add_heading('3. HIGH-LEVEL ARCHITECTURE', 1)

add_heading('3.1 System Architecture Diagram', 2)

arch_diagram = """
CLIENT TIER (Presentation Layer)
+-------------------------------------------------------------------+
|                                                                   |
|               React Single Page Application                       |
|                                                                   |
|  +----------------+  +----------------+  +--------------------+   |
|  |                |  |                |  |                    |   |
|  |   User         |  |  Reusable      |  |   HTTP Client      |   |
|  |   Interface    |  |  Components    |  |   (Axios)          |   |
|  |   Pages        |  |  (Material-UI) |  |                    |   |
|  |                |  |                |  |  - JWT Storage     |   |
|  | - Login        |  | - Tables       |  |  - Auto Auth       |   |
|  | - Dashboard    |  | - Forms        |  |  - Error Handle    |   |
|  | - Workflows    |  | - Dialogs      |  |                    |   |
|  | - Connections  |  | - Steppers     |  |                    |   |
|  | - Detail Page  |  |                |  |                    |   |
|  +----------------+  +----------------+  +--------------------+   |
|                                                                   |
+-------------------------------------------------------------------+
                            |
                            | HTTPS/REST API (JSON Payload)
                            | JWT Authentication
                            v
APPLICATION TIER (Business Logic Layer)
+-------------------------------------------------------------------+
|                                                                   |
|                    FastAPI Backend Server                         |
|                                                                   |
|  +----------------+  +----------------+  +--------------------+   |
|  |                |  |                |  |                    |   |
|  |  Middleware    |  |   API          |  |   Business         |   |
|  |                |  |   Routers      |  |   Logic            |   |
|  | - CORS         |  |                |  |                    |   |
|  | - JWT Verify   |  | - /auth        |  | - DB Mgmt          |   |
|  | - Error        |  | - /server/     |  | - Workflow Exec    |   |
|  |   Handler      |  |   connections  |  | - In-Place Masking |   |
|  | - Logging      |  | - /server/     |  | - Constraint       |   |
|  |                |  |   workflows    |  |   Checking         |   |
|  |                |  | - /server/     |  | - Preview          |   |
|  |                |  |   constraints  |  |                    |   |
|  +----------------+  +----------------+  +--------------------+   |
|                                                                   |
+-------------------------------------------------------------------+
                            |
                            | SQL Queries (ODBC Protocol)
                            | Parameterized
                            v
DATA TIER (Persistence Layer)
+-------------------------------------------------------------------+
|                                                                   |
|                  Database Servers                                 |
|         (Azure SQL, PostgreSQL, Oracle, SQL Server)               |
|                                                                   |
|  +---------------------------+    +---------------------------+   |
|  |                           |    |                           |   |
|  |  Application Database     |    |   User Database           |   |
|  |  (Metadata Storage)       |    |   (In-Place Masking)      |   |
|  |                           |    |                           |   |
|  |  Tables:                  |    |  Single database where    |   |
|  |  - users                  |    |  PII masking is applied   |   |
|  |  - server_connections     |    |  directly on same table   |   |
|  |  - server_workflows       |    |                           |   |
|  |  - workflow_executions    |    |  Tables:                  |   |
|  |  - execution_logs         |    |  - User-defined tables    |   |
|  |                           |    |  - Business data          |   |
|  |  Purpose:                 |    |  - Contains PII to mask   |   |
|  |  Store system config      |    |                           |   |
|  |  and audit trail          |    |                           |   |
|  |                           |    |                           |   |
|  +---------------------------+    +---------------------------+   |
|                                                                   |
+-------------------------------------------------------------------+
"""
add_diagram(arch_diagram)

add_heading('3.2 Architecture Pattern', 2)
add_text('Pattern Type: 3-Tier Layered Architecture')
add_text('')
add_text('Why This Pattern:')
add_bullet('Separation of Concerns: Each tier has a distinct responsibility')
add_bullet('Independent Scaling: Scale presentation, logic, and data layers independently')
add_bullet('Technology Flexibility: Replace or upgrade individual tiers without affecting others')
add_bullet('Security: Multiple security checkpoints at each layer')
add_bullet('Maintainability: Changes to one layer minimally impact other layers')

doc.add_page_break()

# ==============================================================================
# SECTION 4: COMPONENT ARCHITECTURE
# ==============================================================================
print("Adding Section 4: Component Architecture...")

add_heading('4. COMPONENT ARCHITECTURE', 1)

add_heading('4.1 Frontend Component Structure', 2)

frontend_diagram = """
FRONTEND ARCHITECTURE

                         App.js (Root)
                              |
        +---------------------+---------------------+
        |                     |                     |
        v                     v                     v
  Authentication         Routing Layer         Layout Wrapper
   Management          (React Router)          (Navigation)
        |                     |                     |
        |     +---------------+---------------+     |
        |     |               |               |     |
        v     v               v               v     v
     Public Route      Protected Routes      Shared UI
        |                     |                     |
        v                     v                     v
   +---------+       +-------------------+    +----------+
   | Login   |       | Page Components   |    | AppBar   |
   | Page    |       +---------+---------+    | Drawer   |
   +---------+                 |              +----------+
                               |
                +--------------|-------------+
                |              |             |
                v              v             v
         +----------+    +----------+   +----------+
         |Dashboard |    |Workflows |   |Execution |
         |   Page   |    |   Page   |   |   Page   |
         +----------+    +----------+   +----------+
                               |
                +--------------|-------------+
                |              |             |
                v              v             v
         +----------+    +----------+   +----------+
         | Create   |    | Workflow |   |Connection|
         |Workflow  |    |  Detail  |   |   Page   |
         |   Page   |    |   Page   |   +----------+
         +----------+    +----------+
         (Multi-Step)
                |
    +-----------+-----------+
    |           |           |
    v           v           v
Step 1      Step 2      Step 3
Basic       Source      Column
Info        Selection   Mapping
"""
add_diagram(frontend_diagram)

add_heading('4.2 Frontend Component Descriptions', 2)

add_heading('App.js (Root Component)', 3)
add_bullet('Entry point of the React application')
add_bullet('Sets up routing structure')
add_bullet('Manages global state initialization')
add_bullet('Wraps application with theme provider')

add_heading('Authentication Management', 3)
add_bullet('PrivateRoute Component: Guards protected routes from unauthorized access')
add_bullet('Checks localStorage for JWT token')
add_bullet('Redirects to login page if token missing or expired')
add_bullet('Allows access to protected pages if authenticated')

add_heading('Page Components', 3)
add_text('Login Page:')
add_bullet('Username and password input fields')
add_bullet('Submit button triggers authentication API call')
add_bullet('Error message display for invalid credentials')
add_bullet('Redirects to dashboard on successful login')

add_text('Dashboard Page:')
add_bullet('Overview of system statistics (total connections, workflows, executions)')
add_bullet('Recent activity feed showing latest executions')
add_bullet('Quick action buttons to navigate to main features')
add_bullet('Summary cards with color-coded status indicators')

add_text('CreateWorkflowPage (Core Component):')
add_bullet('Multi-step wizard for workflow creation')
add_bullet('Step 1 - Basic Information: Workflow name, description')
add_bullet('Step 2 - Connection & Table Selection: Select connection, schema, and table for in-place masking')
add_bullet('Step 3 - Column Mapping (CRITICAL): Smart filtering shows only compatible PII attributes based on column data types')
add_bullet('Review and save workflow configuration')

add_text('WorkflowDetailPage:')
add_bullet('Displays comprehensive workflow information with tabbed interface')
add_bullet('Overview Tab: Workflow configuration, status, and actions')
add_bullet('Execution History Tab: Past executions with metrics and logs')
add_bullet('Preview Masking Tab: Column Mapping, Constraint Checks, Preview Masking')
add_bullet('Back button with navigation to workflows list')
add_bullet('Execute Workflow button for on-demand execution')

add_heading('4.3 Backend Component Structure', 2)

backend_diagram = """
BACKEND ARCHITECTURE

                     FastAPI Application
                              |
        +---------------------+---------------------+
        |                     |                     |
        v                     v                     v
   Middleware            API Routers          Business Services
    Layer                   Layer                   Layer
        |                     |                     |
+-------+-------+    +--------+--------+    +------+------+
|               |    |                 |    |             |
| - CORS        |    | - Auth Router   |    | - DB Manager|
| - JWT Auth    |    | - Connections   |    | - Workflow  |
| - Error       |    | - Workflows     |    |   Executor  |
|   Handler     |    | - Executions    |    | - Masking   |
| - Logging     |    | - Schemas       |    |   Engine    |
| - Rate Limit  |    | - PII Attrs     |    | - Validator |
|               |    |                 |    |             |
+-------+-------+    +--------+--------+    +------+------+
        |                     |                     |
        +---------------------+---------------------+
                              |
                              v
                        Data Access
                           Layer
                              |
                +-------------+-------------+
                |             |             |
                v             v             v
        Application DB    Source DB     Target DB
         (Metadata)       (Read PII)   (Write Masked)
"""
add_diagram(backend_diagram)

add_heading('4.4 Backend Component Descriptions', 2)

add_heading('Middleware Layer', 3)
add_bullet('CORS Middleware: Validates request origin, allows only frontend URL')
add_bullet('JWT Authentication Middleware: Decodes and validates JWT token, extracts user information')
add_bullet('Error Handler Middleware: Catches exceptions and formats error responses consistently')

add_heading('API Routers', 3)
add_bullet('Auth Router (/api/auth): Login, user profile retrieval')
add_bullet('Server Connections Router (/api/server/connections): Connection CRUD, testing, schema/table discovery')
add_bullet('Server Workflows Router (/api/server/workflows): Workflow CRUD, execution triggering, preview masking')
add_bullet('Server Constraints Router (/api/server/constraints): Check PKs, FKs, unique, check constraints, triggers, indexes')
add_bullet('Server Masking Router (/api/server/masking): Preview masking, execute in-place masking')
add_bullet('Workflows Router (/api/workflows): PII attributes retrieval')

add_heading('Business Services Layer', 3)
add_text('Database Manager Service: Establishes connections to database servers (Azure SQL, PostgreSQL, Oracle, SQL Server), manages connection pooling, executes parameterized queries, handles connection errors with retries.')
add_text('')
add_text('Workflow Executor Service: Orchestrates in-place masking - loads configuration, connects to database, applies masking transformations directly to the same table using UPDATE statements, logs results.')
add_text('')
add_text('Masking Engine Service: Implements PII masking algorithms including string masking (fake names, emails, phones), numeric masking (random numbers, ranges), date/datetime masking (shifting), boolean masking.')
add_text('')
add_text('Constraint Checker Service: Validates database constraints before masking including primary keys, foreign keys, unique constraints, check constraints, triggers, and indexes.')
add_text('')
add_text('Preview Service: Generates preview of masked data on sample records before actual execution.')

doc.add_page_break()

# ==============================================================================
# SECTION 5: DATA FLOW DIAGRAMS
# ==============================================================================
print("Adding Section 5: Data Flow Diagrams...")

add_heading('5. DATA FLOW DIAGRAMS', 1)

add_heading('5.1 User Authentication Flow', 2)

auth_flow = """
USER AUTHENTICATION FLOW

User                Frontend            Backend           Database
 |                     |                   |                  |
 |  1. Enter           |                   |                  |
 |     Credentials     |                   |                  |
 |-------------------->|                   |                  |
                       |  2. POST /api/auth/                  |
                       |     login          |                  |
                       |------------------->|                  |
                                           |  3. Query users   |
                                           |     table         |
                                           |----------------->|
                                           |  4. User record   |
                                           |<-----------------|
                                           |  5. Validate      |
                                           |     password hash |
                                           |                  |
                                           |  IF VALID:       |
                                           |  6. Generate JWT  |
                                           |     Token with:   |
                                           |     - user_id     |
                                           |     - username    |
                                           |     - role        |
                                           |     - expiration  |
                       |  7. Return Token   |                  |
                       |<-------------------|                  |
                       |  8. Store in       |                  |
                       |     localStorage   |                  |
 |  9. Redirect to     |                   |                  |
 |     Dashboard       |                   |                  |
 |<--------------------|                   |                  |
"""
add_diagram(auth_flow)

add_text('Flow Description:')
add_bullet('User Input: User enters username and password in login form')
add_bullet('API Request: Frontend sends POST request with credentials')
add_bullet('Database Query: Backend queries users table to find matching username')
add_bullet('Password Validation: Backend compares hashed passwords using bcrypt')
add_bullet('Token Generation: If valid, backend creates JWT token with user claims and 24-hour expiration')
add_bullet('Token Storage: Frontend stores token in localStorage')
add_bullet('Navigation: User redirected to protected dashboard route')
add_bullet('Subsequent Requests: All future API calls include JWT token in Authorization header')

add_text('Security Benefits:')
add_bullet('Stateless authentication (no server-side session storage)')
add_bullet('Token signature prevents tampering')
add_bullet('Automatic expiration enforces re-authentication')
add_bullet('Each API request independently authenticated')

add_heading('5.2 Workflow Creation Flow', 2)

workflow_flow = """
WORKFLOW CREATION FLOW (In-Place Masking)

User          Frontend         Backend        App DB      User DB
 |                |                |              |            |
 |  Navigate to   |                |              |            |
 |  Create Page   |                |              |            |
 |--------------->|                |              |            |
                  |  Load Initial  |              |            |
                  |  Data          |              |            |
                  |                |              |            |
                  |  GET /server/  |              |            |
                  |  connections   |              |            |
                  |--------------->|              |            |
                  |                |  Query       |            |
                  |                |  connections |            |
                  |                |------------>|            |
                  |  Connection    |  List        |            |
                  |  List          |<------------|            |
                  |<---------------|              |            |
                  |                |              |            |
 |  STEP 1:       |                |              |            |
 |  Fill Basic    |                |              |            |
 |  Info          |                |              |            |
 |--------------->|                |              |            |
                  |  Validate &    |              |            |
                  |  Move to Step2 |              |            |
                  |                |              |            |
 |  STEP 2:       |                |              |            |
 |  Select        |                |              |            |
 |  Connection,   |                |              |            |
 |  Schema, Table |                |              |            |
 |--------------->|                |              |            |
                  |  GET /server/  |              |            |
                  |  connections/  |              |            |
                  |  {id}/schemas  |              |            |
                  |--------------->|              |            |
                  |                |  Connect to  |            |
                  |                |  user DB     |            |
                  |                |--------------|----------->|
                  |                |  Schema list |            |
                  |  Schema List   |<-------------|------------|
                  |<---------------|              |            |
                  |                |              |            |
 |  STEP 3:       |                |              |            |
 |  Column        |                |              |            |
 |  Mapping       |                |              |            |
 |--------------->|                |              |            |
                  |  For each      |              |            |
                  |  column filter |              |            |
                  |  PII by type   |              |            |
                  |                |              |            |
 |  Submit        |                |              |            |
 |--------------->|                |              |            |
                  |  POST /server/ |              |            |
                  |  workflows     |              |            |
                  |--------------->|              |            |
                  |                |  INSERT      |            |
                  |                |  workflow    |            |
                  |                |------------>|            |
                  |  Success       |  Workflow ID |            |
                  |<---------------|<------------|            |
 |  Navigate to   |                |              |            |
 |  Workflows     |                |              |            |
 |<---------------|                |              |            |
"""
add_diagram(workflow_flow)

add_text('Flow Description:')
add_bullet('Initial Load: Frontend fetches server connections and categorized PII attributes from backend')
add_bullet('Step 1 - Basic Information: User provides workflow name and description')
add_bullet('Step 2 - Connection & Table Selection: User selects connection, then cascading API calls load schemas, then tables, then column metadata')
add_bullet('Step 3 - Column Mapping (Smart Filtering): For each selected column, frontend identifies data type, maps to PII category, filters dropdown')
add_bullet('Submission: Frontend sends complete configuration to backend. Backend creates workflow record with column mappings')

add_heading('5.3 Workflow Execution Flow (In-Place Masking)', 2)

execution_flow = """
IN-PLACE MASKING EXECUTION FLOW

User   Frontend   Backend     App DB       User DB
 |        |          |            |            |
 |  Click |          |            |            |
 | Execute|          |            |            |
 |------->|          |            |            |
          |  POST    |            |            |
          |  /server/|            |            |
          |  workflows            |            |
          |  /{id}/  |            |            |
          |  execute |            |            |
          |--------->|            |            |
                     |            |            |
                     |  PHASE 1: INITIALIZATION
                     |----------->|            |
                     |  Load      |            |
                     |  workflow  |            |
                     |  config    |            |
                     |  Create    |            |
                     |  exec rec  |            |
          |  exec_id |<-----------|            |
          |<---------|            |            |
 |  Show  |          |            |            |
 | Started|          |            |            |
 |<-------|          |            |            |
          |          |            |            |
          |          |  PHASE 2: IN-PLACE MASKING
          |          |----------->|----------->|
          |          |  UPDATE    |            |
          |          |  table SET |            |
          |          |  col1 =    |            |
          |          |  masked_val|            |
          |          |  COMMIT    |            |
          |          |<-----------|<-----------|
          |          |            |            |
          |          |  PHASE 3: LOGGING
          |          |----------->|            |
          |          |  Log rows  |            |
          |          |  processed |            |
          |          |<-----------|            |
          |          |            |            |
          |          |  PHASE 4: COMPLETION
          |          |----------->|            |
          |          |  Update    |            |
          |          |  exec      |            |
          |          |  status    |            |
          |          |<-----------|            |
          |          |            |            |
 |  View  |          |            |            |
 |  Logs  |          |            |            |
 |------->|          |            |            |
          |  GET     |            |            |
          |  executions           |            |
          |--------->|----------->|            |
          |  Execution            |            |
          |  history |            |            |
          |<---------|<-----------|            |
 | Display|          |            |            |
 | Results|          |            |            |
 |<-------|          |            |            |
"""
add_diagram(execution_flow)

add_heading('5.4 Smart PII Filtering Flow', 2)

pii_filter_flow = """
SMART PII ATTRIBUTE FILTERING

User                Frontend Logic           PII Categories
 |                         |                        |
 |  Load Create            |                        |
 |  Workflow Page          |                        |
 |------------------------>|  Fetch categorized     |
                           |  PII attributes        |
                           |----------------------->|
                           |  Return:               |
                           |  {                     |
                           |    string: [           |
                           |      first_name,       |
                           |      last_name]        |
                           |    numeric: [          |
                           |      random_number]    |
                           |    date: [date_shift]  |
                           |    datetime: [         |
                           |      datetime_shift]   |
                           |    boolean: [          |
                           |      random_boolean]   |
                           |  }                     |
                           |<-----------------------|
                           |  Store in state        |
                           |                        |
 |  User selects           |                        |
 |  column "age" (int)     |                        |
 |------------------------>|                        |
                           |  Smart Filter:         |
                           |  1. Get column info    |
                           |  2. Map data type:     |
                           |     "int" -> "numeric" |
                           |  3. Filter PII:        |
                           |     Return numeric     |
                           |     attributes only    |
                           |                        |
 |  Show PII dropdown:     |                        |
 |  [v] Select PII:        |                        |
 |    - random_number      |                        |
 |    - range_number       |                        |
 |    - fixed_number       |                        |
 |  NOT showing:           |                        |
 |    X first_name         |                        |
 |    X date_shift         |                        |
 |<------------------------|                        |

DATA TYPE MAPPING:
varchar/text      -> string    -> first_name, email_mask
int/numeric       -> numeric   -> random_number
date              -> date      -> date_shift
datetime          -> datetime  -> datetime_shift
bit/boolean       -> boolean   -> random_boolean
"""
add_diagram(pii_filter_flow)

add_text('Flow Description:')
add_text('Frontend automatically filters PII options based on column data type.')
add_text('')
add_text('Process:')
add_bullet('Column Selection: User checks checkbox for column (e.g., "age")')
add_bullet('Metadata Lookup: Frontend retrieves data type from column metadata')
add_bullet('Category Mapping: Frontend maps SQL type to category (int->numeric, varchar->string, date->date)')
add_bullet('PII Filtering: Frontend returns only PII attributes matching the category')
add_bullet('UI Rendering: Dropdown shows only compatible options')
add_text('')
add_text('Benefits: Prevents configuration errors at design time, improves user experience, ensures type compatibility, maintains data integrity.')

doc.add_page_break()

# ==============================================================================
# SECTION 6: SECURITY ARCHITECTURE
# ==============================================================================
print("Adding Section 6: Security Architecture...")

add_heading('6. SECURITY ARCHITECTURE', 1)

add_heading('6.1 Authentication Security Features', 2)

add_text('Password Security:')
add_bullet('Passwords hashed using bcrypt with salt (cost factor: 12 rounds)')
add_bullet('Never stored in plaintext')
add_bullet('Never transmitted except during initial login (over HTTPS)')
add_bullet('Password requirements: Minimum 8 characters (configurable)')

add_text('Token Security:')
add_bullet('JWT signed with HMAC-SHA256 algorithm')
add_bullet('Secret key stored in environment variables (never in code)')
add_bullet('Token expiration enforced (24-hour default)')
add_bullet('Automatic logout on token expiration')
add_bullet('Token includes user role for authorization')

add_text('Session Management:')
add_bullet('Stateless authentication (no server-side sessions)')
add_bullet('Token revocation via expiration')
add_bullet('Frontend clears token on logout')
add_bullet('Expired tokens automatically rejected by backend')

add_text('Benefits of RBAC:')
add_bullet('Simplified permission management')
add_bullet('Clear separation of admin and user capabilities')
add_bullet('Principle of least privilege')
add_bullet('Easy to add new roles in future')
add_bullet('Audit trail shows who had what permissions')

add_heading('6.2 Data Encryption Architecture', 2)

add_text('Encryption at Rest:')
add_text('Database connection passwords in connections.password table column are encrypted using AES-256-CBC encryption. The encryption key is stored in environment variable SECRET_ENCRYPTION_KEY, and a random 16-byte initialization vector (IV) is generated for each encryption. The encrypted ciphertext is stored in the database. When needed for connection, the password is decrypted using the key from environment and IV from cipher. The plaintext exists only in memory, is used immediately, and cleared after use. Passwords are never logged.')

add_text('Password Hashing:')
add_text('User passwords in users.password_hash are hashed using bcrypt with automatic random salt generation and cost factor of 12 rounds (2^12 iterations). The stored hash includes algorithm version, cost factor, salt, and hash output in the format $2b$12$abcd...xyz. During login, the system retrieves the stored hash, uses bcrypt.compare() to hash the input password with the same salt, and compares results.')

add_text('Encryption in Transit:')
add_text('All client-server communication uses HTTPS/TLS 1.3. The TLS handshake includes client hello with supported ciphers, server hello with chosen cipher and certificate containing public key. Client verifies certificate, generates session keys, and encrypts them with server\'s public key. Both sides derive symmetric session keys. All subsequent data is encrypted including API requests, responses, JWT tokens, and passwords during login.')

doc.add_page_break()

# ==============================================================================
# SECTION 7: DEPLOYMENT ARCHITECTURE
# ==============================================================================
print("Adding Section 7: Deployment Architecture...")

add_heading('7. DEPLOYMENT ARCHITECTURE', 1)

add_heading('7.1 Development Environment', 2)

dev_env = """
DEVELOPMENT ENVIRONMENT

Developer Workstation (Windows/Mac/Linux)
+----------------------------------------------------+
|                                                    |
|  +---------------------------------------------+  |
|  |        Frontend (Port 3000)                 |  |
|  |  --------------------------------------     |  |
|  |  React Development Server                   |  |
|  |  - npm start                                |  |
|  |  - Hot Module Replacement (HMR)             |  |
|  |  - Source maps for debugging                |  |
|  |  - React Developer Tools                    |  |
|  |  - Auto-reload on file changes              |  |
|  |  - URL: http://localhost:3000               |  |
|  +---------------------------------------------+  |
|                      |                             |
|                      | API Calls                   |
|                      v                             |
|  +---------------------------------------------+  |
|  |        Backend (Port 8000)                  |  |
|  |  --------------------------------------     |  |
|  |  FastAPI with Uvicorn                       |  |
|  |  - uvicorn main:app --reload                |  |
|  |  - Auto-reload on code changes              |  |
|  |  - OpenAPI docs: /docs                      |  |
|  |  - Debug logging enabled                    |  |
|  |  - Single worker                            |  |
|  |  - URL: http://localhost:8000               |  |
|  +---------------------------------------------+  |
|                      |                             |
|                      | SQL Queries                 |
|                      v                             |
|  +---------------------------------------------+  |
|  |      SQL Server Database                    |  |
|  |  --------------------------------------     |  |
|  |  Local Instance or Remote Dev Server        |  |
|  |  - Application DB: pii_tool_dev             |  |
|  |  - Test source/target databases             |  |
|  |  - Sample data for testing                  |  |
|  |  - Connection: localhost:1433               |  |
|  +---------------------------------------------+  |
|                                                    |
|  Configuration (.env file):                        |
|  - SECRET_KEY=dev-secret-key-12345                 |
|  - DATABASE_URL=localhost:1433                     |
|  - ENVIRONMENT=development                         |
|  - DEBUG=True                                      |
|  - LOG_LEVEL=DEBUG                                 |
|  - CORS_ORIGINS=http://localhost:3000              |
|                                                    |
+----------------------------------------------------+
"""
add_diagram(dev_env)

doc.add_page_break()

# ==============================================================================
# SECTION 8: INTEGRATION ARCHITECTURE
# ==============================================================================
print("Adding Section 8: Integration Architecture...")

add_heading('8. INTEGRATION ARCHITECTURE', 1)

add_heading('8.1 API Integration Pattern', 2)

api_integration = """
API INTEGRATION ARCHITECTURE

            +----------------+
            |  External      |
            |  Systems       |
            |  (Future)      |
            +--------+-------+
                     |
                     | REST API
                     v
+----------------------------------------------------+
|         API GATEWAY (Future)                       |
|  |- Rate limiting                                  |
|  |- API key management                             |
|  |- Request transformation                         |
|  |- Response caching                               |
+------------+---------------------------------------+
             |
             v
+----------------------------------------------------+
|     PII MASKING TOOL BACKEND API                   |
|        (FastAPI Application)                       |
+----------------------------------------------------+
|                                                    |
|  API Endpoints (RESTful):                          |
|                                                    |
|  Authentication:                                   |
|  POST   /api/auth/login                            |
|  GET    /api/auth/me                               |
|                                                    |
|  Server Connections:                               |
|  GET    /api/server/connections                    |
|  POST   /api/server/connections                    |
|  GET    /api/server/connections/{id}               |
|  DELETE /api/server/connections/{id}               |
|  POST   /api/server/connections/test               |
|                                                    |
|  Schema Discovery:                                 |
|  GET    /api/server/connections/{id}/schemas       |
|  GET    /api/server/connections/{id}/schemas/      |
|         {schema}/tables                            |
|  GET    /api/server/connections/{id}/schemas/      |
|         {schema}/tables/{table}/columns            |
|                                                    |
|  Server Workflows:                                 |
|  GET    /api/server/workflows                      |
|  POST   /api/server/workflows                      |
|  GET    /api/server/workflows/{id}                 |
|  DELETE /api/server/workflows/{id}                 |
|  POST   /api/server/workflows/{id}/execute         |
|  GET    /api/server/workflows/{id}/executions      |
|  GET    /api/server/workflows/{id}/executions/     |
|         {exec_id}/logs                             |
|                                                    |
|  Constraint Checking:                              |
|  GET    /api/server/constraints/primary-keys       |
|  GET    /api/server/constraints/foreign-keys       |
|  GET    /api/server/constraints/unique             |
|  GET    /api/server/constraints/check              |
|  GET    /api/server/constraints/triggers           |
|  GET    /api/server/constraints/indexes            |
|                                                    |
|  Masking:                                          |
|  POST   /api/server/masking/preview                |
|                                                    |
|  PII Attributes:                                   |
|  GET    /api/workflows/pii-attributes              |
|                                                    |
|  Health Check:                                     |
|  GET    /health                                    |
|                                                    |
+----------------------------------------------------+
"""
add_diagram(api_integration)

doc.add_page_break()

# ==============================================================================
# DOCUMENT END
# ==============================================================================
print("Adding Document End...")

add_heading('DOCUMENT END', 1)

add_text('This technical architecture document provides a comprehensive overview of the PII Masking Tool system design, focusing on:')
add_bullet('JWT-based authentication for stateless security')
add_bullet('Smart PII filtering based on SQL data types')
add_bullet('ETL-style masking execution with transaction safety')
add_bullet('Role-based access control for authorization')
add_bullet('Production-ready deployment patterns')

add_text('The architecture prioritizes:')
add_bullet('Data Integrity: Constraint validation before insertion')
add_bullet('Security: Multiple layers of protection (network, application, data)')
add_bullet('Maintainability: Clear component boundaries and documentation')
add_bullet('Scalability: Horizontal scaling support for frontend and backend')
add_bullet('User Experience: Intelligent UI that prevents configuration errors')

add_text('All components work together to provide a robust, secure, and user-friendly solution for PII data masking across databases.')

# Save document
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture.docx"
doc.save(output_path)

print(f"\nDocument saved successfully: {output_path}")
print("Document created with all updated sections!")
