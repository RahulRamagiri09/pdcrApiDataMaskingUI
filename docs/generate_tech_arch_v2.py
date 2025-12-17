from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

# Create a new document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(text, level=1):
    """Add a heading with specified level"""
    return doc.add_heading(text, level)

def add_text(text):
    """Add a normal paragraph"""
    return doc.add_paragraph(text)

def add_bullet(text):
    """Add a bullet point"""
    return doc.add_paragraph(text, style='List Bullet')

def add_diagram(text):
    """Add ASCII diagram in monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_table(headers, rows, header_color="003366"):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

# ==============================================================================
# TITLE PAGE
# ==============================================================================
print("Creating title page...")

title = doc.add_paragraph()
title_run = title.add_run('PII MASKING TOOL')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('ARCHITECTURE & DESIGN DOCUMENT')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(5):
    doc.add_paragraph()

info = doc.add_paragraph()
info.add_run('Version: 2.0\n\n').bold = True
info.add_run('Date: December 2025\n\n').bold = True
info.add_run('Document Type: Technical Architecture\n\n').bold = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================
print("Creating table of contents...")

doc.add_heading('TABLE OF CONTENTS', 0)
doc.add_paragraph()

toc = [
    '1. Executive Summary',
    '2. System Overview',
    '3. High-Level Architecture',
    '4. Component Architecture',
    '5. Data Flow Diagrams',
    '6. Technology Stack',
    '7. API Reference',
    '8. Security Architecture',
    '9. Deployment Architecture',
    '10. Frontend Routes'
]

for item in toc:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ==============================================================================
# SECTION 1: EXECUTIVE SUMMARY
# ==============================================================================
print("Adding Section 1: Executive Summary...")

add_heading('1. EXECUTIVE SUMMARY', 1)

add_heading('1.1 Project Overview', 2)
add_text('The PII Masking Tool is an enterprise web application designed to protect Personally Identifiable Information (PII) by performing in-place masking of sensitive data within a single database. The tool masks PII columns directly in the same table, schema, and database where the data resides.')

add_heading('1.2 Purpose', 2)
add_text('This document provides a comprehensive technical architecture overview of the PII Masking Tool, including:')
add_bullet('System architecture and component relationships')
add_bullet('Data flow through the application')
add_bullet('Technology choices and their justification')
add_bullet('Security measures and deployment strategy')

add_heading('1.3 Key Features', 2)
add_bullet('Server Connection Management: Configure and test connections to SQL Server databases (Azure SQL, PostgreSQL, Oracle, SQL Server)')
add_bullet('Schema Discovery: Browse database schemas, tables, and column metadata')
add_bullet('Intelligent Workflow Creation: Select tables and map columns to appropriate PII masking techniques based on data types')
add_bullet('In-Place PII Masking: Execute data transformation directly on the same table with constraint validation')
add_bullet('Preview Masking: Preview how masking will affect sample records before execution')
add_bullet('Constraint Checking: Validate primary keys, foreign keys, unique constraints, check constraints, triggers, and indexes before masking')
add_bullet('Execution Control: Start, stop, pause, and resume workflow executions')
add_bullet('Execution Monitoring: Track workflow execution history with success/failure metrics')
add_bullet('Role-Based Access Control: Admin, Privilege, General, and Support roles with different permission levels')
add_bullet('User Management: Admin can create users and roles')
add_bullet('Audit Trail: Complete history of who executed what and when')

add_heading('1.4 Target Users', 2)
add_bullet('Database Administrators: Manage database connections and schema exploration')
add_bullet('Data Privacy Officers: Configure PII masking workflows')
add_bullet('Data Engineers: Execute workflows to mask sensitive data in-place')
add_bullet('Compliance Teams: Monitor masking operations and audit trails')

doc.add_page_break()

# ==============================================================================
# SECTION 2: SYSTEM OVERVIEW
# ==============================================================================
print("Adding Section 2: System Overview...")

add_heading('2. SYSTEM OVERVIEW', 1)

add_heading('2.1 Business Context', 2)
add_text('Organizations need to comply with data privacy regulations (GDPR, CCPA, HIPAA) by protecting sensitive PII data within their databases. The PII Masking Tool automates the process of:')
add_bullet('Connecting to database servers (Azure SQL, PostgreSQL, Oracle, SQL Server)')
add_bullet('Discovering schemas, tables, and columns containing PII')
add_bullet('Configuring column-level masking rules based on data types')
add_bullet('Validating database constraints before masking (PKs, FKs, unique, check, triggers, indexes)')
add_bullet('Executing in-place masking transformations directly on the same table')
add_bullet('Maintaining data integrity and referential constraints')

add_heading('2.2 System Capabilities', 2)

capabilities_diagram = """
SYSTEM CAPABILITIES

SERVER CONNECTION MANAGEMENT
|- Register database server connections
|- Support multiple database types (Azure SQL, PostgreSQL, Oracle, SQL Server)
|- Test connection validity before saving
|- Secure credential storage with encryption

SCHEMA EXPLORATION
|- Browse available schemas in connected databases
|- View tables within each schema
|- Inspect column metadata (name, type, nullability)
|- Automatic data type detection

WORKFLOW CONFIGURATION
|- Create named workflows with descriptions
|- Select single connection for in-place masking
|- Select schema and table for masking
|- Configure column-level PII masking rules
|- Smart filtering: Only show compatible masking options

CONSTRAINT CHECKING
|- Check primary key constraints
|- Check foreign key constraints
|- Check unique constraints
|- Check check constraints
|- Check triggers on table
|- Check indexes on table

EXECUTION CONTROL
|- Start workflow execution
|- Stop running execution
|- Pause running execution
|- Resume paused execution
|- Monitor execution status in real-time

USER MANAGEMENT (Admin Only)
|- Create new users
|- Assign roles to users
|- Create new roles
|- View all users and roles

MONITORING & REPORTING
|- View workflow execution history
|- Track success/failure status
|- Display error messages for failed executions
|- Show row counts (processed vs masked)
|- View execution logs
|- User attribution (who executed what)

SECURITY & COMPLIANCE
|- JWT-based authentication
|- Role-based authorization (Admin/Privilege/General/Support)
|- Encrypted database credentials
|- Audit trail of all operations
|- Secure password storage (hashed)
"""
add_diagram(capabilities_diagram)

doc.add_page_break()

# ==============================================================================
# SECTION 3: HIGH-LEVEL ARCHITECTURE
# ==============================================================================
print("Adding Section 3: High-Level Architecture...")

add_heading('3. HIGH-LEVEL ARCHITECTURE', 1)

add_heading('3.1 System Architecture Diagram', 2)

arch_diagram = """
CLIENT TIER (Presentation Layer)
+-------------------------------------------------------------------+
|               React Single Page Application                        |
|  +----------------+  +----------------+  +--------------------+    |
|  |   User         |  |  Reusable      |  |   HTTP Client      |    |
|  |   Interface    |  |  Components    |  |   (Axios)          |    |
|  |   Pages        |  |  (Material-UI) |  |                    |    |
|  | - Login        |  | - Tables       |  |  - JWT Storage     |    |
|  | - Dashboard    |  | - Forms        |  |  - Auto Auth       |    |
|  | - Workflows    |  | - Dialogs      |  |  - Error Handle    |    |
|  | - Connections  |  | - Steppers     |  |                    |    |
|  +----------------+  +----------------+  +--------------------+    |
+-------------------------------------------------------------------+
                            |
                            | HTTPS/REST API (JSON Payload)
                            | JWT Authentication
                            v
APPLICATION TIER (Business Logic Layer)
+-------------------------------------------------------------------+
|                    FastAPI Backend Server                          |
|  +----------------+  +----------------+  +--------------------+    |
|  |  Middleware    |  |   API          |  |   Business         |    |
|  | - CORS         |  |   Routers      |  |   Logic            |    |
|  | - JWT Verify   |  | - /auth        |  | - DB Mgmt          |    |
|  | - Error        |  | - /datamasking/|  | - Workflow Exec    |    |
|  |   Handler      |  |   connections  |  | - In-Place Masking |    |
|  | - Logging      |  | - /datamasking/|  | - Constraint       |    |
|  |                |  |   workflows    |  |   Checking         |    |
|  +----------------+  +----------------+  +--------------------+    |
+-------------------------------------------------------------------+
                            |
                            | SQL Queries (ODBC Protocol)
                            v
DATA TIER (Persistence Layer)
+-------------------------------------------------------------------+
|                  Database Servers                                  |
|         (Azure SQL, PostgreSQL, Oracle, SQL Server)                |
|  +---------------------------+    +---------------------------+    |
|  |  Application Database     |    |   User Database           |    |
|  |  (Metadata Storage)       |    |   (In-Place Masking)      |    |
|  |  Tables:                  |    |  Single database where    |    |
|  |  - users                  |    |  PII masking is applied   |    |
|  |  - roles                  |    |  directly on same table   |    |
|  |  - server_connections     |    |                           |    |
|  |  - server_workflows       |    |                           |    |
|  |  - workflow_executions    |    |                           |    |
|  +---------------------------+    +---------------------------+    |
+-------------------------------------------------------------------+
"""
add_diagram(arch_diagram)

add_heading('3.2 Architecture Pattern', 2)
add_text('Pattern Type: 3-Tier Layered Architecture')
add_text('Why This Pattern:')
add_bullet('Separation of Concerns: Each tier has a distinct responsibility')
add_bullet('Independent Scaling: Scale presentation, logic, and data layers independently')
add_bullet('Technology Flexibility: Replace or upgrade individual tiers without affecting others')
add_bullet('Security: Multiple security checkpoints at each layer')
add_bullet('Maintainability: Changes to one layer minimally impact other layers')

doc.add_page_break()

# ==============================================================================
# SECTION 4: COMPONENT ARCHITECTURE
# ==============================================================================
print("Adding Section 4: Component Architecture...")

add_heading('4. COMPONENT ARCHITECTURE', 1)

add_heading('4.1 Frontend Source Structure', 2)

structure_diagram = """
src/
+-- components/
|   +-- common/
|   |   +-- ProtectedAction.js    # RBAC-controlled UI rendering
|   |   +-- PageHeader.js
|   +-- Layout/
|   |   +-- MainLayout.js         # Responsive layout with sidebar
|   +-- Login/
|   |   +-- Login.js              # Authentication form
|   +-- ProtectedRoute/
|   |   +-- ProtectedRoute.js     # Route guard
|   +-- Sidebar/
|   |   +-- Sidebar.js            # Navigation with role-based visibility
|   +-- ServerDashboard/
|   |   +-- ServerDashboard.js    # Main dashboard
|   +-- ServerConnections/
|   |   +-- ServerConnectionsPage.js
|   |   +-- CreateConnectionDialog.js
|   +-- ServerWorkflows/
|   |   +-- ServerWorkflowsPage.js
|   |   +-- CreateWorkflowPage.js   # 4-step workflow wizard
|   |   +-- WorkflowDetailPage.js
|   +-- UserRegistration/
|   |   +-- UserRegistration.js   # Admin: Create users
|   +-- RoleRegistration/
|       +-- RoleRegistration.js   # Admin: Create roles
+-- services/
|   +-- api.js                    # Axios instances with interceptors
+-- utils/
|   +-- auth.js                   # Authentication helpers
|   +-- rbac.js                   # Role-Based Access Control
|   +-- timeFormat.js             # Utility functions
+-- hooks/
|   +-- usePermission.js          # RBAC hooks
+-- context/
|   +-- SidebarContext.js         # Sidebar state
+-- App.js                        # Router configuration
+-- index.js                      # Entry point
"""
add_diagram(structure_diagram)

doc.add_page_break()

# ==============================================================================
# SECTION 5: DATA FLOW DIAGRAMS
# ==============================================================================
print("Adding Section 5: Data Flow Diagrams...")

add_heading('5. DATA FLOW DIAGRAMS', 1)

add_heading('5.1 User Authentication Flow', 2)

auth_flow = """
USER AUTHENTICATION FLOW

User                Frontend            Backend           Database
 |                     |                   |                  |
 |  1. Enter           |                   |                  |
 |     Credentials     |                   |                  |
 |-------------------->|                   |                  |
                       |  2. POST /api/    |                  |
                       |     auth/login    |                  |
                       |------------------->|                  |
                                           |  3. Query users  |
                                           |     table        |
                                           |----------------->|
                                           |  4. User record  |
                                           |<-----------------|
                                           |  5. Validate     |
                                           |     password     |
                                           |  6. Generate JWT |
                       |  7. Return Token  |                  |
                       |<------------------|                  |
                       |  8. Store in      |                  |
                       |     localStorage  |                  |
 |  9. Redirect to     |                   |                  |
 |     Dashboard       |                   |                  |
 |<--------------------|                   |                  |
"""
add_diagram(auth_flow)

add_heading('5.2 Workflow Execution Flow (In-Place Masking)', 2)

execution_flow = """
IN-PLACE MASKING EXECUTION FLOW

User   Frontend   Backend     App DB       User DB
 |        |          |            |            |
 |  Click |          |            |            |
 | Execute|          |            |            |
 |------->|          |            |            |
          |  POST    |            |            |
          |/datamasking/workflows/execute     |
          |--------->|            |            |
                     |  PHASE 1: INITIALIZATION |
                     |----------->|            |
                     |  Load      |            |
                     |  workflow  |            |
          |  exec_id |<-----------|            |
          |<---------|            |            |
 |  Show  |          |            |            |
 | Started|          |            |            |
 |<-------|          |            |            |
          |          |  PHASE 2: IN-PLACE MASK |
          |          |----------->|----------->|
          |          |  UPDATE    |            |
          |          |  SET col = |            |
          |          |  masked    |            |
          |          |  COMMIT    |            |
          |          |<-----------|<-----------|
          |          |  PHASE 3: LOGGING       |
          |          |----------->|            |
          |          |<-----------|            |
          |          |  PHASE 4: COMPLETION    |
          |          |----------->|            |
          |          |<-----------|            |
 |  View  |          |            |            |
 |  Logs  |          |            |            |
 |------->|          |            |            |
          |  GET     |            |            |
          | executions|           |            |
          |--------->|----------->|            |
          |  History |            |            |
          |<---------|<-----------|            |
 | Display|          |            |            |
 | Results|          |            |            |
 |<-------|          |            |            |
"""
add_diagram(execution_flow)

doc.add_page_break()

# ==============================================================================
# SECTION 6: TECHNOLOGY STACK
# ==============================================================================
print("Adding Section 6: Technology Stack...")

add_heading('6. TECHNOLOGY STACK', 1)

add_heading('6.1 Frontend Stack', 2)
frontend_headers = ['Technology', 'Version', 'Purpose']
frontend_rows = [
    ['React', '18.3.1', 'UI Framework'],
    ['React Router DOM', '6.30.1', 'Client-side routing'],
    ['Material-UI (MUI)', '7.3.2', 'UI Component library'],
    ['MUI Data Grid', '7.x', 'Advanced data tables'],
    ['Axios', '1.12.2', 'HTTP client'],
    ['Tailwind CSS', '3.4.17', 'Utility-first CSS'],
    ['Node.js', '18.16.1', 'Runtime environment'],
]
create_table(frontend_headers, frontend_rows)

doc.add_paragraph()

add_heading('6.2 Backend Stack', 2)
backend_headers = ['Technology', 'Purpose']
backend_rows = [
    ['FastAPI', 'Python web framework'],
    ['Uvicorn', 'ASGI server'],
    ['PyJWT', 'JWT token handling'],
    ['bcrypt', 'Password hashing'],
    ['pyodbc', 'Database connectivity'],
]
create_table(backend_headers, backend_rows)

doc.add_paragraph()

add_heading('6.3 Database Support', 2)
add_bullet('Azure SQL Database')
add_bullet('PostgreSQL')
add_bullet('Oracle')
add_bullet('SQL Server')

doc.add_page_break()

# ==============================================================================
# SECTION 7: API REFERENCE
# ==============================================================================
print("Adding Section 7: API Reference...")

add_heading('7. API REFERENCE', 1)

add_heading('7.1 Base Configuration', 2)
add_text('Frontend Port: 5000')
add_text('Backend API URL: http://localhost:9000')
add_text('API Base Path: /api')

doc.add_paragraph()

add_heading('7.2 Authentication APIs', 2)
auth_headers = ['Method', 'Endpoint', 'Description']
auth_rows = [
    ['POST', '/api/auth/login', 'User login'],
]
create_table(auth_headers, auth_rows)

doc.add_paragraph()

add_heading('7.3 User Management APIs (Admin Only)', 2)
user_headers = ['Method', 'Endpoint', 'Description']
user_rows = [
    ['POST', '/api/users', 'Create user'],
    ['GET', '/api/users', 'Get all users'],
    ['POST', '/api/roles', 'Create role'],
    ['GET', '/api/roles', 'Get all roles'],
]
create_table(user_headers, user_rows)

doc.add_paragraph()

add_heading('7.4 Server Connections APIs', 2)
conn_headers = ['Method', 'Endpoint', 'Description']
conn_rows = [
    ['GET', '/api/datamasking/connections', 'List all connections'],
    ['POST', '/api/datamasking/connections', 'Create connection'],
    ['POST', '/api/datamasking/connections/getById', 'Get connection details'],
    ['DELETE', '/api/datamasking/connections/delete', 'Delete connection'],
    ['POST', '/api/datamasking/connections/test', 'Test connection'],
]
create_table(conn_headers, conn_rows)

doc.add_paragraph()

add_heading('7.5 Schema Discovery APIs', 2)
schema_headers = ['Method', 'Endpoint', 'Description']
schema_rows = [
    ['POST', '/api/datamasking/connections/schemas', 'List schemas'],
    ['POST', '/api/datamasking/connections/tables', 'List tables'],
    ['POST', '/api/datamasking/connections/columns', 'List columns'],
]
create_table(schema_headers, schema_rows)

doc.add_page_break()

add_heading('7.6 Server Workflows APIs', 2)
wf_headers = ['Method', 'Endpoint', 'Description']
wf_rows = [
    ['GET', '/api/datamasking/workflows', 'List all workflows'],
    ['POST', '/api/datamasking/workflows', 'Create workflow'],
    ['POST', '/api/datamasking/workflows/getById', 'Get workflow details'],
    ['PUT', '/api/datamasking/workflows/update', 'Update workflow'],
    ['DELETE', '/api/datamasking/workflows/delete', 'Delete workflow'],
    ['POST', '/api/datamasking/workflows/executions', 'Get execution history'],
    ['GET', '/api/datamasking/workflows/pii-attributes', 'Get PII attribute types'],
]
create_table(wf_headers, wf_rows)

doc.add_paragraph()

add_heading('7.7 Execution & Masking APIs', 2)
exec_headers = ['Method', 'Endpoint', 'Description']
exec_rows = [
    ['POST', '/api/datamasking/workflows/execute', 'Execute workflow'],
    ['POST', '/api/datamasking/workflows/executions/status', 'Get execution status'],
    ['POST', '/api/datamasking/workflows/executions/stop', 'Stop execution'],
    ['POST', '/api/datamasking/workflows/executions/pause', 'Pause execution'],
    ['POST', '/api/datamasking/workflows/executions/resume', 'Resume execution'],
    ['POST', '/api/datamasking/workflows/preview', 'Preview masked data'],
    ['POST', '/api/masking/sample-data', 'Generate sample data'],
]
create_table(exec_headers, exec_rows)

doc.add_paragraph()

add_heading('7.8 Constraint Checking APIs', 2)
const_headers = ['Method', 'Endpoint', 'Description']
const_rows = [
    ['POST', '/api/datamasking/constraints/all', 'Check all constraints'],
    ['POST', '/api/datamasking/constraints/primaryKeys', 'Check primary keys'],
    ['POST', '/api/datamasking/constraints/foreignKeys', 'Check foreign keys'],
    ['POST', '/api/datamasking/constraints/unique', 'Check unique constraints'],
    ['POST', '/api/datamasking/constraints/check', 'Check check constraints'],
    ['POST', '/api/datamasking/constraints/triggers', 'Check triggers'],
    ['POST', '/api/datamasking/constraints/indexes', 'Check indexes'],
]
create_table(const_headers, const_rows)

doc.add_page_break()

# ==============================================================================
# SECTION 8: SECURITY ARCHITECTURE
# ==============================================================================
print("Adding Section 8: Security Architecture...")

add_heading('8. SECURITY ARCHITECTURE', 1)

add_heading('8.1 Authentication', 2)

add_text('Password Security:')
add_bullet('Passwords hashed using bcrypt with salt (cost factor: 12 rounds)')
add_bullet('Never stored in plaintext')
add_bullet('Never transmitted except during initial login (over HTTPS)')

add_text('JWT Token Security:')
add_bullet('JWT signed with HMAC-SHA256 algorithm')
add_bullet('Secret key stored in environment variables')
add_bullet('Token expiration enforced (24-hour default)')
add_bullet('Automatic logout on token expiration')
add_bullet('Token includes user role for authorization')

add_text('Token Storage:')
add_bullet('JWT token stored in localStorage as authToken')
add_bullet('User object stored in localStorage as user')

add_heading('8.2 Authorization (RBAC)', 2)

add_text('Four roles with different permission levels:')

rbac_headers = ['Role', 'Capabilities']
rbac_rows = [
    ['Admin', 'Full CRUD on connections, workflows, users, roles. Execute workflows.'],
    ['Privilege', 'View connections, execute workflows, control executions'],
    ['General', 'Read-only access to all features'],
    ['Support', 'Read-only access (same as General)'],
]
create_table(rbac_headers, rbac_rows)

add_text('See RBAC_Access_Matrix_v2.md for detailed permission matrix.')

add_heading('8.3 Data Encryption', 2)

add_text('At Rest:')
add_bullet('Database connection passwords encrypted with AES-256-CBC')
add_bullet('User passwords hashed with bcrypt')

add_text('In Transit:')
add_bullet('All communication over HTTPS/TLS 1.3')
add_bullet('JWT tokens encrypted in transit')

doc.add_page_break()

# ==============================================================================
# SECTION 9: DEPLOYMENT ARCHITECTURE
# ==============================================================================
print("Adding Section 9: Deployment Architecture...")

add_heading('9. DEPLOYMENT ARCHITECTURE', 1)

add_heading('9.1 Development Environment', 2)

dev_env = """
DEVELOPMENT ENVIRONMENT

Developer Workstation
+----------------------------------------------------+
|  +---------------------------------------------+   |
|  |        Frontend (Port 5000)                 |   |
|  |  React Development Server                   |   |
|  |  - npm start                                |   |
|  |  - Hot Module Replacement (HMR)             |   |
|  |  - URL: http://localhost:5000               |   |
|  +---------------------------------------------+   |
|                      |                             |
|                      | API Calls                   |
|                      v                             |
|  +---------------------------------------------+   |
|  |        Backend (Port 9000)                  |   |
|  |  FastAPI with Uvicorn                       |   |
|  |  - uvicorn main:app --reload --port 9000    |   |
|  |  - OpenAPI docs: /docs                      |   |
|  |  - URL: http://localhost:9000               |   |
|  +---------------------------------------------+   |
|                      |                             |
|                      | SQL Queries                 |
|                      v                             |
|  +---------------------------------------------+   |
|  |      Database Server                        |   |
|  |  Azure SQL / PostgreSQL / SQL Server        |   |
|  +---------------------------------------------+   |
|                                                    |
|  Configuration (.env file):                        |
|  - REACT_APP_API_URL=http://localhost:9000         |
|  - PORT=5000                                       |
+----------------------------------------------------+
"""
add_diagram(dev_env)

doc.add_page_break()

# ==============================================================================
# SECTION 10: FRONTEND ROUTES
# ==============================================================================
print("Adding Section 10: Frontend Routes...")

add_heading('10. FRONTEND ROUTES', 1)

add_heading('10.1 Route Configuration', 2)

routes_headers = ['Route', 'Component', 'Access']
routes_rows = [
    ['/login', 'Login', 'Public'],
    ['/datamasking/dashboard', 'ServerDashboard', 'Protected'],
    ['/datamasking/connections', 'ServerConnectionsPage', 'Protected'],
    ['/datamasking/workflows', 'ServerWorkflowsPage', 'Protected'],
    ['/datamasking/workflows/create', 'CreateWorkflowPage', 'Protected'],
    ['/datamasking/workflows/:id/edit', 'CreateWorkflowPage', 'Protected'],
    ['/datamasking/workflows/:id', 'WorkflowDetailPage', 'Protected'],
    ['/register-user', 'UserRegistration', 'Protected (Admin)'],
    ['/register-role', 'RoleRegistration', 'Protected (Admin)'],
    ['/', 'Redirect to /login', '-'],
    ['*', 'Redirect to /login', '-'],
]
create_table(routes_headers, routes_rows)

doc.add_paragraph()

add_heading('10.2 Navigation Structure', 2)
add_text('Sidebar Menu Items:')
add_bullet('Dashboard')
add_bullet('Connections')
add_bullet('Workflows')
add_bullet('Register User (Admin only)')
add_bullet('Register Role (Admin only)')

doc.add_page_break()

# ==============================================================================
# CHANGES FROM V1
# ==============================================================================
print("Adding Changes from V1...")

add_heading('Changes from Version 1.0', 1)

changes_headers = ['Area', 'v1.0', 'v2.0']
changes_rows = [
    ['API Base Path', '/api/server/...', '/api/datamasking/...'],
    ['API Pattern', 'Path parameters', 'POST with request body'],
    ['Frontend Port', '3000', '5000'],
    ['Backend Port', '8000', '9000'],
    ['Execution Control', 'Not documented', 'Stop/Pause/Resume'],
    ['User Management', 'Not documented', 'Create users/roles'],
    ['RBAC Roles', 'Admin/User', 'Admin/Privilege/General/Support'],
]
create_table(changes_headers, changes_rows)

doc.add_paragraph()
doc.add_paragraph()

# ==============================================================================
# DOCUMENT END
# ==============================================================================
print("Adding Document End...")

add_heading('DOCUMENT END', 1)

add_text('This technical architecture document provides a comprehensive overview of the PII Masking Tool system design for Version 2.0, including:')
add_bullet('JWT-based authentication for stateless security')
add_bullet('4-role RBAC system (Admin, Privilege, General, Support)')
add_bullet('Smart PII filtering based on SQL data types')
add_bullet('In-place masking execution with transaction safety')
add_bullet('Execution control (start/stop/pause/resume)')
add_bullet('User management for admins')

# Save document
output_path = r"c:\Users\rahul.ramagiri\Documents\poc_pii_ui\pii-masking-tool\docs\Technical_Architecture_v2.docx"
doc.save(output_path)

print(f"\nDocument saved successfully: {output_path}")
print("Technical Architecture v2 DOCX created!")
