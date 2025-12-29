from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

# Create a new document
doc = Document()

# Set up document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(text, level=1):
    """Add a heading with specified level"""
    heading = doc.add_heading(text, level)
    return heading

def add_text(text):
    """Add a normal paragraph"""
    return doc.add_paragraph(text)

def add_bullet(text):
    """Add a bullet point"""
    return doc.add_paragraph(text, style='List Bullet')

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_table(headers, rows, header_color="003366"):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def add_code_block(text):
    """Add a code block (monospace font)"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return para

# ==============================================================================
# TITLE PAGE
# ==============================================================================
print("Creating title page...")

title = doc.add_paragraph()
title_run = title.add_run('PII MASKING TOOL')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('ARCHITECTURE & DESIGN DOCUMENT')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(5):
    doc.add_paragraph()

info = doc.add_paragraph()
info.add_run('Version: 3.0\n\n').bold = True
info.add_run('Date: December 2025\n\n').bold = True
info.add_run('Document Type: Technical Architecture\n\n').bold = True
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================
print("Adding Table of Contents...")

add_heading('TABLE OF CONTENTS', 1)

toc_items = [
    '1. Executive Summary',
    '2. System Overview',
    '3. High-Level Architecture',
    '4. Component Architecture',
    '5. Data Flow Diagrams',
    '6. Technology Stack',
    '7. API Reference',
    '8. Security Architecture',
    '9. Deployment Architecture',
    '10. Frontend Routes',
]

for item in toc_items:
    add_text(item)

doc.add_page_break()

# ==============================================================================
# EXECUTIVE SUMMARY
# ==============================================================================
print("Adding Executive Summary...")

add_heading('1. EXECUTIVE SUMMARY', 1)

add_heading('1.1 Project Overview', 2)
add_text('The PII Masking Tool is an enterprise web application designed to protect Personally Identifiable Information (PII) by performing in-place masking of sensitive data within a single database. The tool masks PII columns directly in the same table, schema, and database where the data resides.')

add_heading('1.2 Purpose', 2)
add_text('This document provides a comprehensive technical architecture overview of the PII Masking Tool, including:')
add_bullet('System architecture and component relationships')
add_bullet('Data flow through the application')
add_bullet('Technology choices and their justification')
add_bullet('Security measures and deployment strategy')

add_heading('1.3 Key Features', 2)
add_bullet('Server Connection Management: Configure and test connections to SQL Server databases')
add_bullet('Schema Discovery: Browse database schemas, tables, and column metadata')
add_bullet('Intelligent Workflow Creation: Select tables and map columns to PII masking techniques')
add_bullet('In-Place PII Masking: Execute data transformation directly on the same table')
add_bullet('Preview Masking: Preview how masking will affect sample records')
add_bullet('Constraint Checking: Validate constraints before masking')
add_bullet('Execution Control: Start, stop, pause, and resume workflow executions')
add_bullet('Role-Based Access Control: Admin, Privilege, General, and Support roles')
add_bullet('User Management: Admin can create users and roles')
add_bullet('API Encryption: Optional AES-256 encryption for all API payloads')

add_heading('1.4 Target Users', 2)
add_bullet('Database Administrators: Manage database connections and schema exploration')
add_bullet('Data Privacy Officers: Configure PII masking workflows')
add_bullet('Data Engineers: Execute workflows to mask sensitive data')
add_bullet('Compliance Teams: Monitor masking operations and audit trails')

doc.add_page_break()

# ==============================================================================
# TECHNOLOGY STACK
# ==============================================================================
print("Adding Technology Stack...")

add_heading('6. TECHNOLOGY STACK', 1)

add_heading('6.1 Frontend Stack', 2)
frontend_headers = ['Technology', 'Version', 'Purpose']
frontend_rows = [
    ['React', '18.3.1', 'UI Framework'],
    ['React Router DOM', '6.30.1', 'Client-side routing'],
    ['Material-UI (MUI)', '7.3.2', 'UI Component library'],
    ['MUI Data Grid', '7.x', 'Advanced data tables'],
    ['Axios', '1.12.2', 'HTTP client'],
    ['Tailwind CSS', '3.4.17', 'Utility-first CSS'],
    ['CryptoJS', '4.2.0', 'AES encryption'],
    ['Node.js', '18.16.1', 'Runtime environment'],
]
create_table(frontend_headers, frontend_rows)
doc.add_paragraph()

add_heading('6.2 Backend Stack', 2)
backend_headers = ['Technology', 'Purpose']
backend_rows = [
    ['FastAPI', 'Python web framework'],
    ['Uvicorn', 'ASGI server'],
    ['PyJWT', 'JWT token handling'],
    ['bcrypt', 'Password hashing'],
    ['pyodbc', 'Database connectivity'],
    ['PyCryptodome', 'AES encryption'],
]
create_table(backend_headers, backend_rows)
doc.add_paragraph()

add_heading('6.3 Database Support', 2)
add_bullet('Azure SQL Database')
add_bullet('PostgreSQL')
add_bullet('Oracle')
add_bullet('SQL Server')

doc.add_page_break()

# ==============================================================================
# API REFERENCE
# ==============================================================================
print("Adding API Reference...")

add_heading('7. API REFERENCE', 1)

add_heading('7.1 Base Configuration', 2)
add_text('Frontend Port: 5000')
add_text('Backend API URL: http://localhost:9000')
add_text('API Base Path: /api')
doc.add_paragraph()

add_heading('7.2 Authentication APIs', 2)
auth_headers = ['Method', 'Endpoint', 'Description']
auth_rows = [
    ['POST', '/api/auth/login', 'User login'],
]
create_table(auth_headers, auth_rows)
doc.add_paragraph()

add_heading('7.3 User Management APIs (Admin Only)', 2)
user_headers = ['Method', 'Endpoint', 'Description']
user_rows = [
    ['POST', '/api/users', 'Create user'],
    ['GET', '/api/users', 'Get all users'],
    ['POST', '/api/roles', 'Create role'],
    ['GET', '/api/roles', 'Get all roles'],
]
create_table(user_headers, user_rows)
doc.add_paragraph()

add_heading('7.4 Server Connections APIs', 2)
conn_headers = ['Method', 'Endpoint', 'Description']
conn_rows = [
    ['GET', '/api/datamasking/connections', 'List all connections'],
    ['POST', '/api/datamasking/connections', 'Create connection'],
    ['POST', '/api/datamasking/connections/getById', 'Get connection details'],
    ['DELETE', '/api/datamasking/connections/delete', 'Delete connection'],
    ['POST', '/api/datamasking/connections/test', 'Test connection'],
]
create_table(conn_headers, conn_rows)
doc.add_paragraph()

add_heading('7.5 Schema Discovery APIs', 2)
schema_headers = ['Method', 'Endpoint', 'Description']
schema_rows = [
    ['POST', '/api/datamasking/connections/schemas', 'List schemas'],
    ['POST', '/api/datamasking/connections/tables', 'List tables'],
    ['POST', '/api/datamasking/connections/columns', 'List columns'],
]
create_table(schema_headers, schema_rows)
doc.add_paragraph()

add_heading('7.6 Server Workflows APIs', 2)
wf_headers = ['Method', 'Endpoint', 'Description']
wf_rows = [
    ['GET', '/api/datamasking/workflows', 'List all workflows'],
    ['POST', '/api/datamasking/workflows', 'Create workflow'],
    ['POST', '/api/datamasking/workflows/getById', 'Get workflow details'],
    ['PUT', '/api/datamasking/workflows/update', 'Update workflow'],
    ['DELETE', '/api/datamasking/workflows/delete', 'Delete workflow'],
    ['POST', '/api/datamasking/workflows/executions', 'Get execution history'],
    ['GET', '/api/datamasking/workflows/pii-attributes', 'Get PII attribute types'],
]
create_table(wf_headers, wf_rows)

doc.add_page_break()

add_heading('7.7 Execution & Masking APIs', 2)
exec_headers = ['Method', 'Endpoint', 'Description']
exec_rows = [
    ['POST', '/api/datamasking/workflows/execute', 'Execute workflow'],
    ['POST', '/api/datamasking/workflows/executions/status', 'Get execution status'],
    ['POST', '/api/datamasking/workflows/executions/stop', 'Stop execution'],
    ['POST', '/api/datamasking/workflows/executions/pause', 'Pause execution'],
    ['POST', '/api/datamasking/workflows/executions/resume', 'Resume execution'],
    ['POST', '/api/datamasking/workflows/preview', 'Preview masked data'],
    ['POST', '/api/masking/sample-data', 'Generate sample data'],
    ['POST', '/api/datamasking/masking/validate-workflow', 'Validate workflow'],
]
create_table(exec_headers, exec_rows)
doc.add_paragraph()

add_heading('7.8 Constraint Checking APIs', 2)
const_headers = ['Method', 'Endpoint', 'Description']
const_rows = [
    ['POST', '/api/datamasking/constraints/all', 'Check all constraints'],
    ['POST', '/api/datamasking/constraints/primaryKeys', 'Check primary keys'],
    ['POST', '/api/datamasking/constraints/foreignKeys', 'Check foreign keys'],
    ['POST', '/api/datamasking/constraints/unique', 'Check unique constraints'],
    ['POST', '/api/datamasking/constraints/check', 'Check check constraints'],
    ['POST', '/api/datamasking/constraints/triggers', 'Check triggers'],
    ['POST', '/api/datamasking/constraints/indexes', 'Check indexes'],
]
create_table(const_headers, const_rows)

doc.add_page_break()

# ==============================================================================
# SECURITY ARCHITECTURE
# ==============================================================================
print("Adding Security Architecture...")

add_heading('8. SECURITY ARCHITECTURE', 1)

add_heading('8.1 Authentication', 2)
add_text('Password Security:')
add_bullet('Passwords hashed using bcrypt with salt (cost factor: 12 rounds)')
add_bullet('Never stored in plaintext')
add_bullet('Never transmitted except during initial login (over HTTPS)')
doc.add_paragraph()

add_text('JWT Token Security:')
add_bullet('JWT signed with HMAC-SHA256 algorithm')
add_bullet('Secret key stored in environment variables')
add_bullet('Token expiration enforced (24-hour default)')
add_bullet('Token includes user role for authorization')
doc.add_paragraph()

add_text('Token Storage:')
add_bullet('JWT token stored in localStorage as authToken')
add_bullet('User object stored in localStorage as user')

add_heading('8.2 Authorization (RBAC)', 2)
add_text('Four roles with different permission levels:')
doc.add_paragraph()
role_headers = ['Role', 'Capabilities']
role_rows = [
    ['Admin', 'Full CRUD on connections, workflows, users, roles. Execute workflows.'],
    ['Privilege', 'View connections, execute workflows, control executions'],
    ['General', 'Read-only access to all features'],
    ['Support', 'Read-only access (same as General)'],
]
create_table(role_headers, role_rows)
doc.add_paragraph()
add_text('See RBAC_Access_Matrix_v3.md for detailed permission matrix.')

add_heading('8.3 Data Encryption', 2)
add_text('At Rest:')
add_bullet('Database connection passwords encrypted with AES-256-CBC')
add_bullet('User passwords hashed with bcrypt')
doc.add_paragraph()

add_text('In Transit:')
add_bullet('All communication over HTTPS/TLS 1.3')
add_bullet('JWT tokens encrypted in transit')
add_bullet('Optional API payload encryption (see 8.4)')

doc.add_page_break()

add_heading('8.4 API Request/Response Encryption', 2)
add_text('Overview: Optional AES-256 encryption for all API request and response payloads provides an additional layer of security beyond HTTPS.')
doc.add_paragraph()

add_text('Configuration (.env):')
add_code_block('''REACT_APP_ENCRYPTION_ENABLED=true
REACT_APP_ENCRYPTION_KEY=MySecure32CharacterKeyHere123!''')
doc.add_paragraph()

enc_headers = ['Setting', 'Description']
enc_rows = [
    ['REACT_APP_ENCRYPTION_ENABLED', 'Enable/disable encryption (true/false)'],
    ['REACT_APP_ENCRYPTION_KEY', '32-character key for AES-256 encryption'],
]
create_table(enc_headers, enc_rows)
doc.add_paragraph()

add_text('How It Works:')
add_bullet('Request Encryption: Axios interceptor encrypts outgoing data as { encrypted: "..." }')
add_bullet('Response Decryption: Axios interceptor decrypts incoming data')
add_bullet('Error Response Decryption: Error responses are also decrypted before error handling')
doc.add_paragraph()

add_text('Implementation Files:')
add_bullet('src/utils/encryption.js - Encryption/decryption functions')
add_bullet('src/services/api.js - Axios interceptors with encryption')
doc.add_paragraph()

add_text('Security Notes:')
add_bullet('The same encryption key must be configured on both frontend and backend')
add_bullet('Change the default key in production environments')
add_bullet('Key should be stored securely and not committed to version control')

doc.add_page_break()

# ==============================================================================
# DEPLOYMENT ARCHITECTURE
# ==============================================================================
print("Adding Deployment Architecture...")

add_heading('9. DEPLOYMENT ARCHITECTURE', 1)

add_heading('9.1 Development Environment', 2)
add_text('Frontend: Port 5000 (React Development Server)')
add_text('Backend: Port 9000 (FastAPI with Uvicorn)')
add_text('Database: Azure SQL / PostgreSQL / SQL Server')

add_heading('9.2 Environment Configuration', 2)
add_text('.env file:')
add_code_block('''# API Configuration
REACT_APP_API_URL=http://localhost:9000
PORT=5000

# Encryption Settings
# Set to 'true' to enable encryption/decryption, 'false' to disable
REACT_APP_ENCRYPTION_ENABLED=true

# Encryption Key (32 characters for AES-256)
# IMPORTANT: Change this key in production and keep it secret!
# The same key must be used on the backend for decryption
REACT_APP_ENCRYPTION_KEY=MySecure32CharacterKeyHere123!''')

doc.add_page_break()

# ==============================================================================
# FRONTEND ROUTES
# ==============================================================================
print("Adding Frontend Routes...")

add_heading('10. FRONTEND ROUTES', 1)

add_heading('10.1 Route Configuration', 2)
route_headers = ['Route', 'Component', 'Access']
route_rows = [
    ['/login', 'Login', 'Public'],
    ['/datamasking/dashboard', 'ServerDashboard', 'Protected'],
    ['/datamasking/connections', 'ServerConnectionsPage', 'Protected'],
    ['/datamasking/workflows', 'ServerWorkflowsPage', 'Protected'],
    ['/datamasking/workflows/create', 'CreateWorkflowPage', 'Protected (workflow.create)'],
    ['/datamasking/workflows/:id/edit', 'CreateWorkflowPage', 'Protected (workflow.update)'],
    ['/datamasking/workflows/:id', 'WorkflowDetailPage', 'Protected'],
    ['/registeruser', 'UserRegistration', 'Protected (Admin)'],
    ['/registerrole', 'RoleRegistration', 'Protected (Admin)'],
    ['/', 'Redirect to /login', '-'],
    ['*', 'Redirect to /login', '-'],
]
create_table(route_headers, route_rows)
doc.add_paragraph()

add_heading('10.2 Navigation Structure', 2)
add_text('Sidebar Menu Items:')
add_bullet('Dashboard')
add_bullet('Connections')
add_bullet('Workflows')
add_bullet('Register User (Admin only)')
add_bullet('Register Role (Admin only)')

doc.add_page_break()

# ==============================================================================
# CHANGES FROM V2
# ==============================================================================
print("Adding Changes from V2...")

add_heading('Changes from Version 2.0', 1)

changes_headers = ['Area', 'v2.0', 'v3.0']
changes_rows = [
    ['Route /register-user', '/register-user (incorrect)', '/registeruser (fixed)'],
    ['Route /register-role', '/register-role (incorrect)', '/registerrole (fixed)'],
    ['API Encryption', 'Not documented', 'Fully documented (Section 8.4)'],
    ['.env Configuration', 'Basic (2 settings)', 'Extended (4 settings with encryption)'],
    ['Technology Stack', 'Missing CryptoJS', 'Added CryptoJS 4.2.0'],
    ['Frontend Structure', 'Missing encryption.js', 'Added encryption.js'],
]
create_table(changes_headers, changes_rows)

doc.add_paragraph()
doc.add_paragraph()

# Add document end
end_text = doc.add_paragraph()
end_text.add_run('DOCUMENT END').bold = True
end_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
output_path = r"c:\Users\rahul.ramagiri\Documents\git\pdcrApiDataMaskingUI\pdcrApiDataMaskingUI\docs\Technical_Architecture_v3.docx"
doc.save(output_path)

print(f"\nDocument saved successfully: {output_path}")
print("Technical Architecture v3 DOCX created!")
